# -*- coding: utf-8 -*-
"""
Czech DOCX Anonymizer – Complete v7.0
- Načítá jména z JSON knihovny (cz_names.v1.json)
- Kompletní anonymizace podle GDPR
- Vylepšená detekce adres, osob, kontaktů
Výstupy: <basename>_anon.docx / _map.json / _map.txt
"""

import sys, re, json, unicodedata
from typing import Optional, Set
from pathlib import Path
from collections import defaultdict, OrderedDict
from docx import Document
from datetime import datetime

# =============== Globální proměnné ===============
CZECH_FIRST_NAMES = set()

# =============== Načítání knihovny jmen ===============
def load_names_library(json_path: str = "cz_names.v1.json") -> Set[str]:
    """Načte česká jména z JSON souboru."""
    try:
        script_dir = Path(__file__).parent if '__file__' in globals() else Path.cwd()
        json_file = script_dir / json_path

        if not json_file.exists():
            json_file = Path.cwd() / json_path

        if not json_file.exists():
            print(f"⚠️  Varování: {json_path} nenalezen, používám prázdnou knihovnu!")
            return set()

        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
            names = set()
            if isinstance(data, dict):
                # Nová struktura: {"firstnames": {"M": [...], "F": [...], "U": [...]}}
                if 'firstnames' in data:
                    firstnames = data['firstnames']
                    if isinstance(firstnames, dict):
                        for gender_key in ['M', 'F', 'U']:
                            if gender_key in firstnames:
                                names.update(firstnames[gender_key])
                # Stará struktura: {"male": [...], "female": [...]}
                else:
                    names.update(data.get('male', []))
                    names.update(data.get('female', []))
            elif isinstance(data, list):
                names.update(data)

            # Převod na lowercase pro jednodušší porovnávání
            names = {name.lower() for name in names}
            print(f"✓ Načteno {len(names)} jmen z knihovny")
            return names
    except Exception as e:
        print(f"⚠️  Chyba při načítání {json_path}: {e}")
        return set()

# =============== Varianty pro nahrazování ===============
def variants_for_first(first: str) -> set:
    """Generuje základní pádové varianty křestního jména (optimalizováno pro výkon)."""
    f = first.strip()
    if not f: return {''}

    V = {f, f.lower(), f.capitalize()}
    low = f.lower()

    # Ženská jména na -a
    if low.endswith('a'):
        stem = f[:-1]
        # 7 pádů: nominativ, genitiv, dativ, akuzativ, vokativ, lokál, instrumentál
        V |= {stem+'y', stem+'e', stem+'ě', stem+'u', stem+'ou', stem+'o'}
    else:
        # Mužská jména
        V |= {f+'a', f+'ovi', f+'e', f+'em', f+'u', f+'om'}
        # Speciální případy
        if low.endswith('ek'): V.add(f[:-2] + 'ka')
        if low.endswith('el'): V.add(f[:-2] + 'la')
        if low.endswith('ec'): V.add(f[:-2] + 'ce')

    # Bez diakritiky
    ascii_variants = {unicodedata.normalize('NFKD', v).encode('ascii','ignore').decode('ascii') for v in list(V)}

    # DŮLEŽITÉ: Aplikuj normalizaci variant i na ASCII verze!
    # Julia → Julie, Maria → Marie, atd.
    name_variants = {
        'julia': 'julie', 'maria': 'marie', 'alica': 'alice',
        'beatrica': 'beatrice', 'kornelia': 'kornelie', 'rosalia': 'rosalie',
        'nadia': 'nadie', 'silvia': 'silvie', 'elea': 'ela',
        'aurelia': 'aurelie', 'terezia': 'terezie', 'otilia': 'otilie',
        'sofia': 'sofie', 'valeria': 'valerie', 'amalia': 'amalie',
        'antonia': 'antonie', 'melania': 'melanie',
    }
    normalized_ascii = set()
    for av in ascii_variants:
        av_lo = av.lower()
        if av_lo in name_variants:
            normalized_ascii.add(name_variants[av_lo])
            normalized_ascii.add(name_variants[av_lo].capitalize())
        else:
            normalized_ascii.add(av)

    V |= normalized_ascii
    return V

def variants_for_surname(surname: str) -> set:
    """Generuje základní pádové varianty příjmení (optimalizováno pro výkon)."""
    s = surname.strip()
    if not s: return {''}

    out = {s, s.lower(), s.capitalize()}
    low = s.lower()

    # Ženská příjmení na -ová
    if low.endswith('ová'):
        base = s[:-1]
        out |= {s, base+'é', base+'ou'}
        return out

    # Přídavná jména -ský, -cký, -ý
    if low.endswith(('ský','cký','ý')):
        if low.endswith(('ský','cký')):
            stem = s[:-2]
        else:
            stem = s[:-1]
        out |= {stem+'ý', stem+'ého', stem+'ému', stem+'ým', stem+'ém'}
        out |= {stem+'á', stem+'é', stem+'ou'}
        return out

    # Ženská na -á
    if low.endswith('á'):
        stem = s[:-1]
        out |= {s, stem+'é', stem+'ou'}
        return out

    # Speciální případy
    if low.endswith('ek') and len(s) >= 3:
        stem_k = s[:-2] + 'k'
        out |= {s, stem_k+'a', stem_k+'ovi', stem_k+'em', stem_k+'u', stem_k+'e'}
        return out

    if low.endswith('el') and len(s) >= 3:
        stem_l = s[:-2] + 'l'
        out |= {s, stem_l+'a', stem_l+'ovi', stem_l+'em', stem_l+'u'}
        return out

    if low.endswith('ec') and len(s) >= 3:
        stem_c = s[:-2] + 'c'
        out |= {s, stem_c+'e', stem_c+'i', stem_c+'em', stem_c+'u'}
        return out

    # Standardní mužská příjmení
    out |= {s+'a', s+'ovi', s+'e', s+'em', s+'u', s+'ům', s+'em'}
    # Množné číslo: u Nováků
    out |= {s+'ů', s+'ům'}

    return out

# =============== Inference funkcí ===============
def _male_genitive_to_nominative(obs: str) -> Optional[str]:
    """Převede pozorovaný tvar (např. genitiv) na nominativ pro mužská jména.

    DŮLEŽITÉ: Testuje koncovky s PRIORITOU - nejdřív -u, pak -a.
    Důvod: "Petra" může být genitiv od "Petr" (správně) nebo "Petro" (chybně).
    """
    lo = obs.lower()

    # FIRST: Hardcoded list of common feminine names that should NEVER be converted
    # This is necessary because the name library is incomplete (missing Martina, etc.)
    common_feminine_names = {
        'martina', 'jana', 'petra', 'eva', 'anna', 'marie', 'lenka', 'kateřina',
        'alena', 'hana', 'lucie', 'veronika', 'monika', 'jitka', 'zuzana', 'ivana',
        'tereza', 'barbora', 'andrea', 'michaela', 'simona', 'nikola', 'pavla',
        'daniela', 'alexandra', 'kristýna', 'markéta', 'renata', 'šárka', 'karolína',
        'krista', 'beata'
    }

    if lo in common_feminine_names:
        return None  # Don't convert feminine names to masculine

    # Also check library if available
    if lo in CZECH_FIRST_NAMES and lo.endswith('a'):
        return None  # Don't convert, let the caller handle it

    cands = []

    # PRIORITA 1: Speciální případy -ka → -ek, -la → -el, -ce → -ec
    if lo.endswith('ka') and len(obs) > 2:
        cand = obs[:-2] + 'ek'
        if cand.lower() in CZECH_FIRST_NAMES:
            return cand.capitalize()
        cands.append(cand)

    if lo.endswith('la') and len(obs) > 2:
        cand = obs[:-2] + 'el'
        if cand.lower() in CZECH_FIRST_NAMES:
            return cand.capitalize()
        cands.append(cand)

    if lo.endswith('ce') and len(obs) > 2:
        cand = obs[:-2] + 'ec'
        if cand.lower() in CZECH_FIRST_NAMES:
            return cand.capitalize()
        cands.append(cand)

    # PRIORITA 2: Dativ -ovi → remove (Petrovi → Petr, Pavlovi → Pavel)
    if lo.endswith('ovi') and len(obs) > 3:
        cand = obs[:-3]
        if cand.lower() in CZECH_FIRST_NAMES:
            return cand.capitalize()

        # Zkus vložné 'e' (Pavlovi → Pavel, Lukášovi → Lukáš)
        if len(cand) >= 3:
            vowels = 'aeiouyáéěíóúůý'
            last_char = cand[-1]
            if last_char.lower() not in vowels:
                cand_with_e = cand[:-1] + 'e' + last_char
                if cand_with_e.lower() in CZECH_FIRST_NAMES:
                    return cand_with_e.capitalize()

        # Cizí jména končící na -o: Hugovi → Hug → Hugo, Radkovi → Radk → Radko
        if len(cand) >= 2 and cand[-1].lower() not in 'aeiouyáéěíóúůý':
            cand_with_o = cand + 'o'
            foreign_o_names = {'hug', 'iv', 'vit', 'le', 'brun', 'marc', 'dieg', 'radk', 'vald'}
            if cand_with_o.lower() in CZECH_FIRST_NAMES or cand.lower() in foreign_o_names:
                return cand_with_o.capitalize()

        # Jména končící na -a: Oldovi → Old → Olda
        if len(cand) >= 2:
            cand_with_a = cand + 'a'
            common_a_names = {'old', 'jir', 'kub', 'ondr'}  # Olda, Jirka, Kuba, Ondra
            if cand_with_a.lower() in CZECH_FIRST_NAMES or cand.lower() in common_a_names:
                return cand_with_a.capitalize()

        # Jména na -ek: Markovi → Mark → Marek (vložit 'e' před 'k')
        # Podobně: Jankovi → Jank → Janek
        if len(cand) >= 2 and cand[-1].lower() not in 'aeiouyáéěíóúůý':
            # Vlož 'e' před poslední souhlásku
            cand_with_e_before_last = cand[:-1] + 'e' + cand[-1]
            common_ek_stems = {'mark', 'jank', 'zdeňk', 'čeňk', 'tomášk', 'vojtěšk'}
            if cand_with_e_before_last.lower() in CZECH_FIRST_NAMES or cand.lower() in common_ek_stems:
                return cand_with_e_before_last.capitalize()

        cands.append(cand)

    # PRIORITA 3: Instrumentál -em → remove (Petrem → Petr, Pavlem → Pavel)
    if lo.endswith('em') and len(obs) > 2:
        cand = obs[:-2]
        if cand.lower() in CZECH_FIRST_NAMES:
            return cand.capitalize()

        # Zkus vložné 'e' (Pavlem → Pavel, Lukášem → Lukáš)
        # Když "Pavl" není v knihovně, zkus "Pavel" (vložení 'e' před poslední souhlásku)
        if len(cand) >= 3:
            vowels = 'aeiouyáéěíóúůý'
            last_char = cand[-1]
            if last_char.lower() not in vowels:  # Poslední znak je souhláska
                # Vlož 'e' před poslední souhlásku: Pavl → Pav + e + l = Pavel
                cand_with_e = cand[:-1] + 'e' + last_char
                if cand_with_e.lower() in CZECH_FIRST_NAMES:
                    return cand_with_e.capitalize()

        # Cizí jména končící na -o: Hugem → Hug → Hugo
        if len(cand) >= 2 and cand[-1].lower() not in 'aeiouyáéěíóúůý':
            cand_with_o = cand + 'o'
            foreign_o_names = {'hug', 'iv', 'vit', 'le', 'brun', 'marc', 'dieg'}
            if cand_with_o.lower() in CZECH_FIRST_NAMES or cand.lower() in foreign_o_names:
                return cand_with_o.capitalize()

        cands.append(cand)

    # PRIORITA 4: Dativ/Akuzativ -u → remove (Petru → Petr, Karlu → Karel) - PŘED -a!
    # Důležité: testovat PŘED -a, protože "Petra" může být "Petr" + -a
    if lo.endswith('u') and len(obs) > 1:
        cand = obs[:-1]
        if cand.lower() in CZECH_FIRST_NAMES:
            return cand.capitalize()

        # Zkus vložné 'e' (Karlu → Karl → Karel, Pavlu → Pavl → Pavel)
        if len(cand) >= 3:
            vowels = 'aeiouyáéěíóúůý'
            last_char = cand[-1]
            if last_char.lower() not in vowels:
                cand_with_e = cand[:-1] + 'e' + last_char
                # Preferuj formu s vložným 'e' i když není v knihovně
                # Karel, Pavel, Havel jsou běžná jména která mohou chybět v knihovně
                if cand_with_e.lower() in CZECH_FIRST_NAMES:
                    return cand_with_e.capitalize()
                # Known male names with inserted 'e' that might be missing from library
                common_e_names = {'karel', 'pavel', 'havel', 'marcel'}
                if cand_with_e.lower() in common_e_names:
                    return cand_with_e.capitalize()

        cands.append(cand)

    # PRIORITA 4.5: Genitiv ženských jmen -y → remove (Nikoly → Nikola, Radky → Radka)
    # MUSÍ být PŘED -a, protože "Nikoly" obsahuje 'a' uvnitř
    if lo.endswith('y') and len(obs) > 2:
        cand = obs[:-1] + 'a'  # Nikoly → Nikola, Radky → Radka
        if cand.lower() in CZECH_FIRST_NAMES:
            return cand.capitalize()
        cands.append(cand)

    # PRIORITA 5: Genitiv/Akuzativ -a → remove (Petra → Petr, ale i Jana → Jan)
    if lo.endswith('a') and len(obs) > 1:
        cand = obs[:-1]
        if cand.lower() in CZECH_FIRST_NAMES:
            return cand.capitalize()

        # Zkus vložné 'e' (Pavla → Pavel, Lukáša → Lukáš)
        if len(cand) >= 3:
            vowels = 'aeiouyáéěíóúůý'
            last_char = cand[-1]
            if last_char.lower() not in vowels:
                cand_with_e = cand[:-1] + 'e' + last_char
                if cand_with_e.lower() in CZECH_FIRST_NAMES:
                    return cand_with_e.capitalize()

        # Cizí jména končící na -o: Huga → Hug → Hugo, Iva → Iv → Ivo
        if len(cand) >= 2 and cand[-1].lower() not in 'aeiouyáéěíóúůý':
            cand_with_o = cand + 'o'
            foreign_o_names = {'hug', 'iv', 'vit', 'le', 'brun', 'marc', 'dieg'}
            if cand_with_o.lower() in CZECH_FIRST_NAMES or cand.lower() in foreign_o_names:
                return cand_with_o.capitalize()

        cands.append(cand)

    # Vrať první kandidát (pokud existuje)
    return cands[0].capitalize() if cands else None

def normalize_name_variant(obs: str) -> str:
    """Aplikuje pouze normalizaci variant jmen (Julia→Julie) bez pádové inference."""
    lo = obs.lower()

    name_variants = {
        # Mužská jména (včetně genitivů)
        'karl': 'karel',
        'karla': 'karel',  # Genitiv od Karel
        'karlo': 'karel',
        'mark': 'marek',  # Mark je v knihovně, ale normalizuj na Marek
        'stanislava': 'stanislav',  # Genitiv od Stanislav
        'zdeňka': 'zdeněk',  # Genitiv od Zdeněk
        'čeňka': 'čeněk',  # Genitiv od Čeněk
        'františka': 'františek',  # Genitiv od František
        # Dativy mužských jmen (pokud jsou v knihovně jako samostatná jména)
        'petru': 'petr',  # Dativ od Petr (pokud je petru v knihovně)
        'filipu': 'filip',  # Dativ od Filip
        'pavlu': 'pavel',  # Dativ od Pavel
        'tomášu': 'tomáš',  # Dativ od Tomáš
        'lukášu': 'lukáš',  # Dativ od Lukáš
        'jakubu': 'jakub',  # Dativ od Jakub
        # Dativy mužských jmen -ovi (pokud jsou v knihovně jako samostatná jména)
        'albínovi': 'albín',  # Dativ od Albín
        'radomírovi': 'radomír',  # Dativ od Radomír
        'jaroslavovi': 'jaroslav',  # Dativ od Jaroslav
        'stanislavovi': 'stanislav',  # Dativ od Stanislav
        'miroslavovi': 'miroslav',  # Dativ od Miroslav
        'václavovi': 'václav',  # Dativ od Václav
        'ladislavovi': 'ladislav',  # Dativ od Ladislav
        # Varianty jmen které jsou v knihovně ale jsou tvary jiných jmen
        'alexandru': 'alexandr',  # Alexandru je dativ od Alexandr (nebo variant)
        'renem': 'rene',  # Instrumentál od Rene
        'renemu': 'rene',  # Dativ od Rene
        # Ženská jména -ia → -ie
        'maria': 'marie',
        'julia': 'julie',
        'sofia': 'sofie',
        'valeria': 'valerie',
        'amália': 'amálie',
        'antonia': 'antonie',
        'melania': 'melanie',
        'alica': 'alice',
        'beatrica': 'beatrice',
        'beatricía': 'beatrice',  # Přidáno: Beatricía→Beatrice
        'kornelia': 'kornelie',
        'rosalia': 'rosalie',
        'nadia': 'nadie',
        'silvia': 'silvie',
        'elea': 'ela',
        'aurelia': 'aurelie',
        'terezia': 'terezie',
        'lívia': 'lívia',  # Explicitně - není v knihovně, nechť zůstane
        'lívie': 'lívia',  # Lívie je v knihovně, ale normalizuj na Lívia
        'otilia': 'otilie',
        # Zkrácené varianty
        'ela': 'ela',  # Explicitně aby se nesloučila s Elena
        'ele': 'ela',  # Dativ od Ela
        'elen': 'elena',  # Genitiv od Elena (Elena → Elen)
        'eleně': 'elena',  # Lokál od Elena
        'stela': 'stela',  # Explicitně
        'stele': 'stela',  # Dativ od Stela
        'hedvika': 'hedvika',  # Explicitně
        'hedvice': 'hedvika',  # Dativ od Hedvika
        'amira': 'amira',  # Explicitně
        'amiře': 'amira',  # Dativ/Lokál od Amira
        # Genitive forms that should normalize to nominative
        'hany': 'hana',  # Genitiv od Hana
        'jany': 'jana',  # Genitiv od Jana
        'evy': 'eva',  # Genitiv od Eva
        'dany': 'dana',  # Genitiv od Dana
        'gabriela': 'gabriela',  # Explicitně
        'gabriele': 'gabriela',  # Dativ od Gabriela
        'andrea': 'andrea',  # Explicitně
        'andree': 'andrea',  # Dativ od Andrea
        'mila': 'mila',  # Explicitně
        'mile': 'mila',  # Dativ od Mila
        # Cizí jména - zkrácené formy
        'marc': 'marco',  # Marc → Marco
        'le': 'leo',  # Le → Leo
        'lea': 'leo',  # Lea → Leo (ženská forma/genitiv)
    }
    if lo in name_variants:
        return name_variants[lo].capitalize()
    return obs.capitalize()

def infer_first_name_nominative(obs: str) -> str:
    """Odhadne nominativ křestního jména z pozorovaného tvaru.

    Pokrývá všechny české pády + speciální vzory (ice→ika, ře→ra, apod.).
    """
    lo = obs.lower()

    # NORMALIZACE VARIANT - ÚPLNĚ PRVNÍ!
    # Aplikuj normalizaci ještě PŘED jakoukoliv logikou
    # Důvod: Julia a Julie jsou OBĚ v knihovně, musíme unifikovat
    name_variants = {
        # Mužská jména (včetně genitivů)
        # POUZE jednoznačné mapování! Nepoužívat pro jména která mohou být i ženská.
        'karl': 'karel',
        'karla': 'karel',  # Genitiv od Karel
        'karlo': 'karel',
        'mark': 'marek',  # Mark je v knihovně, ale normalizuj na Marek
        'stanislava': 'stanislav',  # Genitiv od Stanislav
        'zdeňka': 'zdeněk',  # Genitiv od Zdeněk
        'čeňka': 'čeněk',  # Genitiv od Čeněk
        'františka': 'františek',  # Genitiv od František
        # Dativy mužských jmen -ovi (pokud jsou v knihovně jako samostatná jména)
        'albínovi': 'albín',  # Dativ od Albín
        'radomírovi': 'radomír',  # Dativ od Radomír
        'jaroslavovi': 'jaroslav',  # Dativ od Jaroslav
        'stanislavovi': 'stanislav',  # Dativ od Stanislav
        'miroslavovi': 'miroslav',  # Dativ od Miroslav
        'václavovi': 'václav',  # Dativ od Václav
        'ladislavovi': 'ladislav',  # Dativ od Ladislav
        # NEPOUŽÍVAT: 'radka' (může být Radko nebo Radek!)
        # NEPOUŽÍVAT: 'janka' (může být Janek nebo ženské jméno Janka!)
        # NEPOUŽÍVAT: 'marka' (může být Marek nebo ženské jméno Marka!)
        # Ženská jména -ia → -ie
        'maria': 'marie',
        'julia': 'julie',
        'sofia': 'sofie',
        'valeria': 'valerie',
        'amália': 'amálie',
        'antonia': 'antonie',
        'melania': 'melanie',
        'alica': 'alice',
        'beatrica': 'beatrice',
        'beatricía': 'beatrice',  # Přidáno: Beatricía→Beatrice
        'kornelia': 'kornelie',
        'rosalia': 'rosalie',
        'nadia': 'nadie',
        'silvia': 'silvie',
        'elea': 'ela',
        'aurelia': 'aurelie',
        'terezia': 'terezie',
        'lívia': 'lívia',  # Explicitně - není v knihovně, nechť zůstane
        'lívie': 'lívia',  # Lívie je v knihovně, ale normalizuj na Lívia
        # Dativ ženských jmen na -e → -a
        'adéle': 'adéla',
        'tereze': 'tereza',
        'lenky': 'lenka', # Should be 'lence'
        # Genitivy ženských jmen na -y → -a
        'dany': 'dana',  # Genitiv od Dana
        'hany': 'hana',  # Genitiv od Hana
        'jany': 'jana',  # Genitiv od Jana
        'evy': 'eva',  # Genitiv od Eva
        'petra': 'petr',  # Could be genitive of Petr OR female name Petra - ambiguous
        'petře': 'petra',
        'alici': 'alice',  # Dativ od Alice
        'alicí': 'alice',  # Instrumentál od Alice
        # Různé varianty
        'otilia': 'otilie',
        # Zkrácené varianty
        'ela': 'ela',  # Explicitně aby se nesloučila s Elena
        'ele': 'ela',  # Dativ od Ela
        'elen': 'elena',  # Genitiv od Elena (Elena → Elen)
        'eleně': 'elena',  # Lokál od Elena
        'stela': 'stela',  # Explicitně
        'stele': 'stela',  # Dativ od Stela
        'hedvika': 'hedvika',  # Explicitně
        'hedvice': 'hedvika',  # Dativ od Hedvika
        'amira': 'amira',  # Explicitně
        'amiře': 'amira',  # Dativ/Lokál od Amira
        'gabriela': 'gabriela',  # Explicitně
        'gabriele': 'gabriela',  # Dativ od Gabriela
        'andrea': 'andrea',  # Explicitně
        'andree': 'andrea',  # Dativ od Andrea
        'mila': 'mila',  # Explicitně
        'mile': 'mila',  # Dativ od Mila
        # Cizí jména - zkrácené formy
        'marc': 'marco',  # Marc → Marco
        'le': 'leo',  # Le → Leo
        'lea': 'leo',  # Lea → Leo (ženská forma/genitiv)
        # Varianty jmen které jsou v knihovně ale jsou tvary jiných jmen
        'alexandru': 'alexandr',  # Alexandru je dativ od Alexandr (nebo variant)
        'renem': 'rene',  # Instrumentál od Rene (reno je v knihovně, ale prioritizuj rene)
        'renemu': 'rene',  # Dativ od Rene
    }
    if lo in name_variants:
        # VŽDY normalizuj, i když je v knihovně
        normalized = name_variants[lo].capitalize()
        # Pokud mapování ukazuje na sebe sama (lívia→lívia), je to nominativ - vrať to
        if name_variants[lo] == lo:
            return normalized
        obs = normalized
        lo = obs.lower()

    # SPECIÁLNÍ PŘÍPAD: Roberta může být genitiv od Robert
    # Musí být PŘED kontrolou knihovny, protože Roberta je v knihovně jako ženské jméno!
    if lo == 'roberta':
        # Preferujeme Robert (mužské jméno), protože Roberta je častěji genitiv než samostatné jméno
        return 'Robert'

    # DŮLEŽITÉ: Kontrola, zda už je v nominativu (v knihovně jmen)
    if lo in CZECH_FIRST_NAMES:
        return obs.capitalize()

    # SPECIÁLNÍ: Jména na -ňa/-ťa jsou pravděpodobně nominativ (Stáňa, Káťa)
    if lo.endswith(('ňa', 'ťa')) and len(obs) >= 4:
        return obs.capitalize()

    # SPECIÁLNÍ VZORY - PRIORITA (před obecnými pravidly)

    # 1. ice → ika (Anice → Anika, Clarice → Clarika)
    if lo.endswith('ice') and len(obs) > 3:
        stem_ika = obs[:-3] + 'ika'
        if stem_ika.lower() in CZECH_FIRST_NAMES:
            return stem_ika.capitalize()

    # 2. ře → ra (Barbaře → Barbara, Saře → Sara)
    if lo.endswith('ře') and len(obs) > 2:
        stem_ra = obs[:-2] + 'ra'
        if stem_ra.lower() in CZECH_FIRST_NAMES:
            return stem_ra.capitalize()

    # 3. Zkrácená jména (Han → Hana, Mart → Marta, ale NE David → Davida)
    # POUZE pro krátká jména (max 4 znaky) aby se předešlo chybám jako David → Davida
    if len(obs) <= 4:
        # Priorita: nejdřív zkus +ina (pro Mart → Martina), pak +a
        if lo + 'ina' in CZECH_FIRST_NAMES:
            return (obs + 'ina').capitalize()
        if lo + 'a' in CZECH_FIRST_NAMES:
            return (obs + 'a').capitalize()

    # ŽENSKÁ JMÉNA - pádové varianty

    # Dativ: -ce → -ka (Lence → Lenka, Radce → Radka, Elišce → Eliška)
    if lo.endswith('ce') and len(obs) > 2:
        stem = obs[:-2]
        if (stem + 'ka').lower() in CZECH_FIRST_NAMES:
            return (stem + 'ka').capitalize()

    # SPECIÁLNÍ: Genitiv bez -y (Elen od Elena, Irén od Irena)
    if lo.endswith('n') and len(obs) > 2:
        stem = obs[:-1]  # Elen → Ele
        if (stem + 'a').lower() in CZECH_FIRST_NAMES:  # Elena
            return (stem + 'a').capitalize()

    # SPECIÁLNÍ: Akuzativ -ii od jmen na -ie (Natálii → Natálie, Julii → Julie)
    if lo.endswith('ii') and len(obs) > 3:
        stem_ie = obs[:-2] + 'ie'
        if stem_ie.lower() in CZECH_FIRST_NAMES:
            return stem_ie.capitalize()
        # Common -ie names
        common_ie_names = {'natálie', 'julie', 'rosalie', 'aurélie', 'amélie', 'otilie', 'kornelie'}
        if stem_ie.lower() in common_ie_names:
            return stem_ie.capitalize()

    # Genitiv/Dativ/Lokál: -y/-ě/-e → -a
    # Speciální: Bei → Bea (genitiv od Bea), Helze → Helga (palatalizace)
    if lo.endswith(('y', 'ě', 'e', 'i')):
        stem = obs[:-1]

        # SPECIÁLNÍ: Palatalizace -ze → -ga (Helze → Helga, Olze → Olga)
        if lo.endswith('ze') and len(stem) >= 3:
            stem_ga = stem[:-1] + 'ga'
            if stem_ga.lower() in CZECH_FIRST_NAMES:
                return stem_ga.capitalize()
            # Common names with palatalization
            common_ga_names = {'helga', 'olga', 'inga'}
            if stem_ga.lower() in common_ga_names:
                return stem_ga.capitalize()

        # Pro -e: zkus nejprve stem bez změny (Denise → Denis, Aleše → Aleš)
        # Pak zkus stem+a (Adéle → Adéla, Tereze → Tereza)
        if lo.endswith('e') and len(stem) >= 2:
            if stem.lower() in CZECH_FIRST_NAMES:
                return stem.capitalize()
            # Zkus stem + a pro ženská jména (Adéle → Adéla)
            if (stem + 'a').lower() in CZECH_FIRST_NAMES:
                return (stem + 'a').capitalize()
            # Common names ending in -a
            common_e_to_a = {'adéla', 'tereza', 'lenka', 'petra', 'eliška', 'aneta', 'veronika', 'monika'}
            if (stem + 'a').lower() in common_e_to_a:
                return (stem + 'a').capitalize()

        # Pro -i: zkus nejprve stem (MALE names - Aleši → Aleš, Petrovi → Petr)
        # Pak zkus stem+a (FEMALE names - Hany → Hana)
        if lo.endswith('i') and len(stem) >= 2:
            # FIRST: Check if stem is a male name (Aleši → Aleš)
            if stem.lower() in CZECH_FIRST_NAMES:
                return stem.capitalize()

        # Pro -y/-ě/-i: zkus stem+a (Boženy → Božena, Žaniny → Žanina, Žanině → Žanina, Hany → Hana)
        if (stem + 'a').lower() in CZECH_FIRST_NAMES:
            return (stem + 'a').capitalize()

        # SPECIÁLNÍ: Palatalizace n → ň v dativu (Stáně → Stáňa, ne Stána)
        # MUSÍ BÝT PŘED common_a_names, protože preferujeme variantu s 'ň'
        # Když vidíme -ně a stem končí na 'n', zkus nahradit 'n' za 'ň'
        common_a_names = {'hana', 'jana', 'anna', 'eva', 'dana', 'žanina', 'kristýna', 'karolína', 'justýna', 'martina', 'regina', 'paulína', 'stáňa', 'stána', 'máňa', 'mána', 'táňa', 'tána'}
        if lo.endswith('ně') and len(stem) >= 2 and stem[-1].lower() == 'n':
            stem_nha = stem[:-1] + 'ň' + 'a'
            if stem_nha.lower() in common_a_names or stem_nha.lower() in CZECH_FIRST_NAMES:
                return stem_nha.capitalize()

        # Common female names ending in -a/-ina/-ína that might be missing from library
        # Hany → Hana, Žaniny → Žanina, Kristýny → Kristýna, Stáně → Stána/Stáňa
        if (stem + 'a').lower() in common_a_names:
            return (stem + 'a').capitalize()

        # Pro -i zkus také bez změny (pokud je to už v nominativu)
        if lo.endswith('i') and len(stem) >= 2:
            # Bei → Bea, ale také kontrola zda není už nominativ (Eli zůstává Eli)
            if stem.lower() in CZECH_FIRST_NAMES:
                return stem.capitalize()

    # Dativ: -u → remove for MALE names first (Danielu → Daniel), then try -u → -a for FEMALE names (Hanu → Hana)
    if lo.endswith('u') and len(obs) > 1:
        stem = obs[:-1]

        # FIRST: Check if stem is a male name (Danielu → Daniel, Robertu → Robert, Michalu → Michal)
        if stem.lower() in CZECH_FIRST_NAMES:
            return stem.capitalize()

        # SECOND: Check if stem + 'a' is a female name (Hanu → Hana, Danu → Dana)
        if (stem + 'a').lower() in CZECH_FIRST_NAMES:
            return (stem + 'a').capitalize()

    # Dativ/Genitiv: -iie → -ia (Líviie → Lívia)
    if lo.endswith('iie') and len(obs) > 3:
        stem = obs[:-3]
        # Zkus -ia (Líviie → Lívia)
        if (stem + 'ia').lower() in CZECH_FIRST_NAMES:
            return (stem + 'ia').capitalize()
        # Common names ending in -ia
        common_ia_names = {'lívia', 'júlia', 'emília', 'cecília'}
        if (stem + 'ia').lower() in common_ia_names:
            return (stem + 'ia').capitalize()

    # Dativ/Lokál: -ii/-ií → -ie/-ia (Lucii → Lucie, Marii → Marie, Julii → Julie, Lívií → Lívia)
    if lo.endswith(('ii', 'ií')) and len(obs) > 2:
        stem = obs[:-2]
        # Zkus -ie (Lucii → Lucie, Marii → Marie)
        if (stem + 'ie').lower() in CZECH_FIRST_NAMES:
            return (stem + 'ie').capitalize()
        # Zkus -ia (Julii → Julia, Lívií → Lívia, méně časté)
        if (stem + 'ia').lower() in CZECH_FIRST_NAMES:
            return (stem + 'ia').capitalize()
        # Common names ending in -ia
        common_ia_names = {'lívia', 'júlia', 'emília', 'cecília'}
        if (stem + 'ia').lower() in common_ia_names:
            return (stem + 'ia').capitalize()

    # Instrumentál: -ici → -ice (Alici → Alice, Beatrici → Beatrice)
    if lo.endswith('ici') and len(obs) > 3:
        stem = obs[:-3]
        if (stem + 'ice').lower() in CZECH_FIRST_NAMES:
            return (stem + 'ice').capitalize()

    # Instrumentál: -í → -e nebo -ice (Alicí → Alice, Marií → Marie)
    # Ženská jména na -ice mají instrumentál -icí (Alice → Alicí)
    # Ženská jména na -ie mají instrumentál -ií (Marie → Marií)
    if lo.endswith('í') and len(obs) > 2:
        stem = obs[:-1]

        # Zkus stem + 'e' (Alicí → Alic → Alice)
        if (stem + 'e').lower() in CZECH_FIRST_NAMES:
            return (stem + 'e').capitalize()

        # Zkus stem (pro případ že je to už základní tvar - Jiří)
        if stem.lower() in CZECH_FIRST_NAMES:
            return stem.capitalize()

    # Instrumentál: -ou → -a (Hanou → Hana, Kristou → Krista)
    if lo.endswith('ou') and len(obs) > 2:
        stem = obs[:-2]
        if (stem + 'a').lower() in CZECH_FIRST_NAMES:
            return (stem + 'a').capitalize()
        # Common female names ending in -a
        common_ou_names = {'hana', 'jana', 'anna', 'eva', 'dana', 'krista', 'beata'}
        if (stem + 'a').lower() in common_ou_names:
            return (stem + 'a').capitalize()

    # MUŽSKÁ JMÉNA - genitiv/dativ/instrumentál
    male_nom = _male_genitive_to_nominative(obs)
    if male_nom:
        return male_nom

    # POSSESSIVE FORMS - Petřin → Petra, Janin → Jana
    if lo.endswith('in') and len(obs) > 2:
        stem = obs[:-2]
        # Zkus ženskou variantu (Petřin → Petra)
        if (stem + 'a').lower() in CZECH_FIRST_NAMES:
            return (stem + 'a').capitalize()

    # Pokud nic nepomohlo, vrať původní tvar s velkým písmenem
    return obs.capitalize()

def infer_surname_nominative(obs: str) -> str:
    """Odhadne nominativ příjmení z pozorovaného tvaru.

    Pokrývá:
    - Ženská příjmení (-ová, -á)
    - Přídavná jména (-ský, -cký, -ý)
    - Vložné 'e' (Havl → Havel, Petr → Petra)
    - Zvířecí příjmení (Liška, Vrba)
    - Všechny pády
    """
    lo = obs.lower()

    # ========== ŽENSKÁ PŘÍJMENÍ ==========

    # -é → -á (genitiv/dativ/lokál žen: Pokorné → Pokorná, Houfové → Houfová)
    if lo.endswith('é') and len(obs) > 3:
        # SPECIÁLNÍ: "-ské" může být genitiv od "-ská" (Panské → Panská)
        # nebo přídavné jméno (Novákské zůstává)
        # Heuristika: pokud je krátké (max 8 znaků) a nezačíná velkým písmenem uvnitř,
        # je to pravděpodobně genitiv příjmení
        if lo.endswith('ské'):
            # Pokud je to krátké slovo bez velkých písmen uprostřed → genitiv příjmení
            if len(obs) <= 10:  # Krátké příjmení
                return obs[:-1] + 'á'  # -ské → -ská (Panské → Panská, Horské → Horská)
        # Pro ostatní -é (ne -ské/-cké)
        elif not lo.endswith('cké'):
            return obs[:-1] + 'á'

    # -ou → může být -a (mužské příjmení), -á (žena) nebo -ý (muž)
    if lo.endswith('ou') and len(obs) > 3:
        # Kontrola, že není -skou/-ckou (přídavné jméno)
        if not lo.endswith(('skou', 'ckou')):
            # Heuristika: pokud základ končí na souhlásku + typický vzor
            base = obs[:-2]

            # PRIORITA 1: Mužská příjmení končící na -a (Jura, Skála, Liška, Vrba, Svoboda)
            # Instrumentál: -ou → -a (Jurou → Jura, Skálou → Skála)
            masculine_a_stems = {'jur', 'skál', 'lišk', 'vrb', 'svobod', 'háb', 'kár',
                                 'forejt', 'korbel', 'machač', 'sedlač', 'ouhel',
                                 'hrab', 'kub', 'kunc', 'másl', 'slíž'}
            if base.lower() in masculine_a_stems or base.lower().endswith(('čk', 'šk', 'nk')):
                return base + 'a'

            # PRIORITA 2: Mužská přídavná jména (Vráný, Novotný)
            # Pro příjmení jako "Vránou" → může být "Vráný" (muž) nebo "Vráná" (žena)
            if base.lower().endswith(('vrán', 'novot', 'malý', 'černý', 'bilý', 'vesel')):
                return base + 'ý'

            # PRIORITA 3: Ženský tvar (defaultní)
            return base + 'á'

    # ========== PŘÍDAVNÁ JMÉNA (-ský, -cký, -ý) ==========

    if lo.endswith(('ého', 'ému', 'ým', 'ém')):
        # Všechny tyto koncovky → -ý
        if lo.endswith('ého'):
            return obs[:-3] + 'ý'
        elif lo.endswith('ému'):
            return obs[:-3] + 'ý'
        elif lo.endswith('ým'):
            return obs[:-2] + 'ý'
        elif lo.endswith('ém'):
            return obs[:-2] + 'ý'

    # -skou/-ckou → -ská/-cká (ženská přídavná)
    if lo.endswith(('skou', 'ckou')):
        return obs[:-3] + 'ká'  # Otradovskou → Otradovská, ne Otradovsá!

    # ========== VLOŽNÉ 'E' - Havl/Havla → Havel, Petr/Petra → Petra ==========

    # Seznam kmenů vyžadujících vložné e
    vlozne_e_stems = {
        'havl', 'sed', 'pes', 'koz', 'voj', 'pav', 'petr', 'jos',
        'alex', 'filip', 'luk', 'mar', 'dan', 'tom', 'jar', 'stef'
    }

    # Detekce: pokud lo končí na 'a' a stem (bez -a) je ve vlozne_e_stems
    if lo.endswith('a') and len(obs) > 2:
        stem_without_a = lo[:-1]  # např. "havl" z "havla"
        if stem_without_a in vlozne_e_stems:
            # Vlož 'e' před poslední souhlásku
            # Havla → Havl → Havel (vložit 'e' před 'l')
            stem = obs[:-1]  # "Havla" → "Havl"
            return stem[:-1] + 'e' + stem[-1]  # "Hav" + "e" + "l" → "Havel"

    # Lepší implementace vložného e:
    # Pokud stem končí na souhlásku-souhlásku, vlož 'e' mezi ně
    if lo.endswith('a') and len(obs) > 3:
        stem = obs[:-1]  # např. "Havl" z "Havla"
        stem_lo = stem.lower()
        # Kontrola, jestli předposlední a poslední znak jsou souhlásky
        consonants = 'bcčdďfghjklmnňpqrřsštťvwxzž'
        if len(stem) >= 2 and stem_lo[-2] in consonants and stem_lo[-1] in consonants:
            # Zkontroluj známé případy
            if stem_lo in vlozne_e_stems or stem_lo + 'e' + 'l' in ['havel', 'pavel']:
                # Vlož 'e' mezi poslední dvě souhlásky
                return stem[:-1] + 'e' + stem[-1]

    # ========== ZVÍŘECÍ PŘÍJMENÍ ==========
    # Liška, Vrba, Ryba, Kočka, Panda, atd. - končí na -a v nominativu!
    # NEODSTRAŇUJ -a, pokud je to zvířecí nebo rostlinné příjmení

    animal_plant_surnames = {
        'liška', 'vrba', 'ryba', 'kočka', 'panda', 'veverka',
        'sova', 'holub', 'vraná', 'zajíc', 'koza', 'ovečka',
        'bříza', 'dub', 'jeřábek', 'jílková', 'kaštanka'
    }

    if lo in animal_plant_surnames:
        return obs  # Zachovej nominativ

    # ========== SPECIÁLNÍ PŘÍPADY: -ka, -la, -ce → -ek, -el, -ec ==========

    # Ale POUZE pokud to NENÍ běžné příjmení končící na -a v nominativu
    common_surnames_a = {
        'svoboda', 'skála', 'hora', 'kula', 'hala', 'krejča',
        'liška', 'vrba', 'ryba', 'kočka', 'sluka', 'janda',
        'procházka', 'blaha', 'kafka', 'smetana', 'brabec',
        'kuřátka', 'kubíčka', 'marečka', 'vašíčka',
        # Další příjmení končící na -ka v nominativu
        'kuba', 'červinka', 'hromádka', 'horčička', 'straka', 'paseka',
        'krupička', 'koudelka', 'řezníčka', 'urbanek',  # Přidáno
        'pavelka', 'popelka', 'hruška', 'kotek'  # Přidáno pro smlouva22
    }

    if lo.endswith('ka') and len(obs) > 3 and lo not in common_surnames_a:
        # NOVÁ LOGIKA: Rozlišit dvě situace:
        # 1. "Hájka" (genitiv od "Hájek") → base bez -ka + ek = "Hájek"
        # 2. "Dvořáka" (genitiv od "Dvořák") → base bez -a = "Dvořák"
        base_without_ka = obs[:-2]  # "Hájka" → "Háj", "Dvořáka" → "Dvořá"
        base_with_ek = base_without_ka + 'ek'

        # Heuristika: pokud base (bez -ka) je krátký (2-5 znaků) a nekončí na samohlásku,
        # je to pravděpodobně -ek příjmení
        base_lo_ka = base_without_ka.lower()
        vowels = ('a', 'á', 'e', 'é', 'ě', 'i', 'í', 'o', 'ó', 'u', 'ú', 'ů', 'y', 'ý')
        if 2 <= len(base_without_ka) <= 5 and not base_lo_ka.endswith(vowels):
            return base_with_ek  # "Hájka" → "Hájek"
        else:
            # Jinak genitiv od příjmení končícího na souhlásku (dlouhé base nebo končí na samohlásku)
            # "Dvořáka" → "Dvořák" (odstranění jen -a)
            return obs[:-1]

    if lo.endswith('la') and len(obs) > 3 and lo not in common_surnames_a:
        # Rozlišit: Havla → Havel vs Šídla → Šídlo
        base_without_a = obs[:-1]  # Šídla → Šídl
        base_lo = base_without_a.lower()

        # Pokud base končí na 'dl', 'zl', 'tl', 'nl' → může být -lo příjmení
        # Šídlo → Šídla → vrátit Šídlo
        # Mazlo → Mazla → vrátit Mazlo
        if base_lo.endswith(('dl', 'zl', 'tl', 'nl', 'sl', 'cl')):
            return base_without_a + 'o'  # Šídla → Šídl → Šídlo
        else:
            # Havel → Havla → vrátit Havel
            base_without_la = obs[:-2]
            return base_without_la + 'el'

    if lo.endswith('ce') and len(obs) > 3:
        # Two patterns:
        # 1. Němec → Němce (genitiv with vložné 'e') → Němce → Němec
        # 2. Švec → Švece (genitiv without vložné 'e') → Švece → Švec

        # Known surnames without vložné 'e' (just remove -e)
        simple_c_surnames = {'švec', 'rybec', 'chlebec', 'holec'}
        base_with_c = obs[:-1].lower()  # "Švece" → "švec"

        if base_with_c in simple_c_surnames:
            return obs[:-1]  # Švece → Švec
        else:
            return obs[:-2] + 'ec'  # Němce → Němec

    # Genitiv s vložným 'e' - obecné: Holase → Holas, Šídla → Šídel
    if lo.endswith('se') and len(obs) > 3:
        # Holase → Holas (odebrat -e)
        return obs[:-1]

    # Šídla, Havla apod. už řešeno výš (řádek 492)

    # ========== DATIV: -ovi → REMOVE ==========

    if lo.endswith('ovi') and len(obs) > 5:
        # Novákovi → Novák
        stem = obs[:-3]
        stem_lo = stem.lower()

        # KONTROLA: Některé kmeny potřebují -a (Pasekovi → Paseka, ne Pasek)
        surname_stems_needing_a = {
            'kub', 'červink', 'hromádk', 'horčičk', 'strak',
            'kuč', 'bárt', 'procházk', 'klím', 'svobod', 'pasek'
        }

        # KONTROLA: Kmeny končící na -dl, -zl, -tl potřebují -o (Šídlovi → Šídlo)
        if stem_lo.endswith(('dl', 'zl', 'tl', 'nl', 'sl', 'cl')):
            return stem + 'o'  # Šídlovi → Šídlo
        elif stem_lo in surname_stems_needing_a:
            return stem + 'a'  # Pasekovi → Paseka
        # NOVÉ: Zkontroluj jestli stem + 'a' je známé příjmení na -ka (Veverkovi → Veverka)
        elif (stem_lo + 'a') in animal_plant_surnames or (stem_lo + 'a') in common_surnames_a:
            return stem + 'a'  # Veverkovi → Veverka
        # NOVÉ: Zkontroluj jestli stem je v vlozne_e_stems (Havlovi → Havel)
        elif stem_lo in vlozne_e_stems:
            # Vlož 'e' před poslední souhlásku
            return stem[:-1] + 'e' + stem[-1]  # Havlovi → Havl → Havel
        # NOVÉ: Zkontroluj jestli stem + 'ek' dává smysl (Hájkovi → Hájek)
        # Nebo vlož 'e' pokud končí na souhlásku-souhlásku (Blažkovi → Blažek)
        # DŮLEŽITÉ: Aplikuj POUZE na kmeny kde chybí vložné 'e'!
        elif stem_lo[-1] not in 'aeiouyáéěíóúůý' and 2 <= len(stem_lo) <= 5:
            # Specifické vzory které vyžadují vložné 'e'
            # Příklady: Hájk → Hájek, Blažk → Blažek, Krčk → Krček
            consonants = 'bcčdďfghjklmnňpqrřsštťvwxzž'

            # Zkontroluj, jestli končí na specifickou kombinaci souhlásek která potřebuje 'e'
            if len(stem) >= 2 and stem_lo[-2] in consonants and stem_lo[-1] in consonants:
                # Specifické vzory které VŽDY potřebují vložné 'e': -jk, -žk, -čk, -rk, -šk
                needs_e_patterns = ['jk', 'žk', 'čk', 'rk', 'šk', 'tk', 'dk', 'ck', 'nk']
                last_two = stem_lo[-2:]

                if last_two in needs_e_patterns:
                    # Vložné 'e' je potřeba
                    return stem[:-1] + 'e' + stem[-1]  # Hájkovi → Hájk → Hájek, Blažkovi → Blažk → Blažek
                else:
                    # Jiné kombinace (lf, rt, etc.) → pravděpodobně úplné jméno
                    return stem  # Volfovi → Volf, Šubrtovi → Šubrt
            else:
                # Nekončí na dvě souhlásky → vrať stem
                return stem  # Řehákovi → Řehák
        else:
            # Pro delší kmeny (6+ znaků) nebo kmeny končící na samohlásku, vrať stem
            return stem  # Novákovi → Novák, Špinarovi → Špinar

    # ========== INSTRUMENTÁL: -ím → -í ==========
    # Krejčím → Krejčí, Košíkem → Košík
    # Příjmení končící na -í/-ý mají instrumentál -ím/-ým
    if lo.endswith('ím') and len(obs) > 3:
        # Krejčím → Krejčí
        return obs[:-1]  # Odstranň 'm'

    # ========== INSTRUMENTÁL: -em → REMOVE ==========

    # Ale POUZE pokud to není součást příjmení (Šembera, Chlumec, atd.)
    if lo.endswith('em') and len(obs) > 4:
        # Speciální případ: -alem, -elem, -olem → pravděpodobně instrumentál od -al, -el, -ol
        if lo.endswith(('alem', 'elem', 'olem', 'ilem')):
            # Doležalem → Doležal, Havlem → Havel, Kokolem → Kokol
            # Ale Havel má výjimku výš (řádek 383-385), takže tady řešíme jen -al/-ol/-il
            if lo.endswith('alem'):
                return obs[:-2]  # Doležalem → Doležal
            elif lo.endswith('elem'):
                return obs[:-2]  # ?elem → ?el (edge case)
            elif lo.endswith('olem'):
                return obs[:-2]  # ?olem → ?ol
            elif lo.endswith('ilem'):
                return obs[:-2]  # ?ilem → ?il

        # Speciální případ: -kem → -ek nebo -ík
        # Musíme rozlišit:
        # 1) -íkem → -ík (Kubíkem → Kubík, Novíkem → Novík)
        # 2) -áškem, -ékem, atd. → -ášek, -ének (Práškem → Prášek, Štefánkem → Štefánek)
        elif lo.endswith('kem') and len(obs) > 5:
            # Pokud před -kem je -í, je to pravděpodobně -ík v nominativu
            if lo.endswith('íkem'):
                return obs[:-2]  # Kubíkem → Kubík (odstranit -em)
            else:
                # Jinak použij pravidlo -kem → -ek
                return obs[:-3] + 'ek'  # Práškem → Prášek, Štefánkem → Štefánek

        # Kontrola: není -bem, -dem (součást některých příjmení)
        # POZNÁMKA: -lem a -rem jsou řešeny výše (řádky 431-441)
        # DŮLEŽITÉ: -cem, -chem, -šem, -gem NEJSOU v blacklistu - jsou to běžné instrumentální koncovky
        # Příklady: Švecem → Švec, Hrubešem → Hrubeš, Vlachem → Vlach
        elif not lo.endswith(('bem', 'dem')):
            # Novákem → Novák, Procházkou → Procházka
            # Kratochvílem → Kratochvíl, Králem → Král
            # Švecem → Švec, Hrubešem → Hrubeš
            return obs[:-2]

    # ========== GENITIV: -a pro příjmení končící na -lo ==========
    # Šídlo → Šídla (genitiv) → vrátit Šídlo
    # Mazlo → Mazla (genitiv) → vrátit Mazlo
    if lo.endswith('dla') or lo.endswith('tla') or lo.endswith('zla') or lo.endswith('nla') or lo.endswith('sla') or lo.endswith('cla'):
        # Šídla → Šídlo (odstranit 'a', přidat 'o')
        return obs[:-1] + 'o'

    # ========== GENITIV: -a → NEODSTRAŇUJ! ==========
    # Mnoho příjmení končí na -a v nominativu (Svoboda, Skála, Liška, atd.)
    # Příliš riskantní, necháme to být

    # ========== GENITIV MNOŽNÉHO ČÍSLA: -ů → remove ==========
    # Šustrů (u Šustrů = u rodiny Šustr) → Šustr nebo Šustrová
    # Pro ženské příjmení potřebujeme přidat -ová
    if lo.endswith('ů') and len(obs) > 2:
        # Šustrů → Šustr (základní tvar)
        # ALE: Některá příjmení končí na -a: Kubů → Kuba, ne Kub
        stem = obs[:-1]

        # Seznam kmenů příjmení která mají nominativ na -a (stem bez -a!)
        # Kubů → stem "Kub" → zkontroluj že je v seznamu → vrať "Kuba"
        surname_stems_needing_a = {
            'kub', 'červink', 'hromádk', 'horčičk', 'strak',
            'kuč', 'bárt', 'procházk', 'klím', 'svobod'
        }

        if stem.lower() in surname_stems_needing_a:
            return stem + 'a'  # Kubů → Kuba
        else:
            return stem  # Šustrů → Šustr

    # ========== DATIV/LOKÁL: -ři, -ře → odstranit 'i' nebo 'e' ==========
    # Šindeláři → Šindelář, Šindeláře → Šindelář
    # Vrátka → Vrát (ne), Šídla → Šídel (dřív), ale Klíči → Klíč, Košíři → Košíř
    if lo.endswith('ři') and len(obs) > 3:
        return obs[:-1]  # Šindeláři → Šindelář, Klíči → Klíč

    if lo.endswith('ře') and len(obs) > 3:
        return obs[:-1]  # Šindeláře → Šindelář, Košíře → Košíř

    # ========== GENITIV: -še, -že, -če → -š, -ž, -č ==========
    # Bartoše → Bartoš, Čapkože → Čapko ž, Vachouškě → Vachoušek
    if lo.endswith(('še', 'že', 'če')) and len(obs) > 3:
        return obs[:-1]  # Odstraň 'e': Bartoše → Bartoš

    # ========== VOKATIV: -le, -re, -ne, -de, -te, -ke → odstranit -e ==========
    # Chýle → Chýl (vokativ), Havle → Havel (ale už řešeno výš)
    # Musíme být opatrní - některá příjmení končí na -e v nominativu (Hrabě)
    if lo.endswith('e') and len(obs) > 2:
        base = obs[:-1]
        base_lo = base.lower()

        # Kontrola: Jestli base končí na souhlásku + typické vokativní vzory
        # Pro příjmení jako Chýl, Král, Srp končící na souhlásku
        if base_lo.endswith(('l', 'r', 'n', 'd', 't', 'k', 'p', 's', 'm', 'v')):
            # VÝJIMKY: příjmení která skutečně končí na -e v nominativu
            if base_lo not in {'hrab', 'knež', 'pán'}:  # Hrabě, Kněže, Páně
                # Kontrola délky: base musí být alespoň 3 znaky
                if len(base) >= 3:
                    return base  # Chýle → Chýl

    # ========== SPECIÁLNÍ PŘÍJMENÍ: -ěte, -ěti → -ě ==========
    # Hraběte → Hrabě (genitiv), Hraběti → Hrabě (dativ)
    if lo.endswith('ěte') and len(obs) > 4:
        return obs[:-2]  # Hraběte → Hrabě

    if lo.endswith('ěti') and len(obs) > 4:
        return obs[:-2]  # Hraběti → Hrabě

    # ========== PŘÍJMENÍ KONČÍCÍ NA -LO: nechat beze změny ==========
    # Šídlo je nominativ, neměnit na Šídl!
    # (odstraněno chybné pravidlo které převádělo Šídlo → Šídl)

    # ========== GENITIV: -ky → -ka (před obecným -y) ==========
    # Veverky → Veverka (genitiv příjmení na -ka)
    # Musí být PŘED obecným -y pravidlem!
    if lo.endswith('ky') and len(obs) > 4:
        candidate_ka = obs[:-1] + 'a'  # Veverky → Veverka
        candidate_ka_lo = candidate_ka.lower()

        # Zkontroluj jestli -ka forma je známé příjmení
        # Buď v animal_plant_surnames nebo v common_surnames_a
        if candidate_ka_lo in animal_plant_surnames or candidate_ka_lo in common_surnames_a:
            return candidate_ka  # Veverky → Veverka

        # Nebo pokud kandidát končí na běžný vzor -ička, -ička, -ůčka, -ečka
        if candidate_ka_lo.endswith(('ička', 'ůčka', 'ečka', 'očka')):
            return candidate_ka

    # ========== GENITIV: -y → -a nebo odstranit -y ==========
    # Klímy → Klíma (genitiv mužů), Procházky → Procházka
    # ALE POUZE pokud to NENÍ přídavné jméno (-ský/-cký/-ný)
    if lo.endswith('y') and len(obs) > 3:
        # Skip if it's adjective form
        if not lo.endswith(('ský', 'cký', 'ný')):
            # Seznam příjmení končících na -a v nominativu
            protected_a_surnames = {
                'procházka', 'klíma', 'svoboda', 'skála', 'hora', 'hala',
                'liška', 'vrba', 'ryba', 'kočka', 'sluka', 'janda',
                'blaha', 'kafka', 'smetana',
                # Další běžná příjmení na -a
                'kuba', 'červinka', 'hromádka', 'horčička', 'straka', 'paseka',
                'kuča', 'bárta', 'slabá', 'malá', 'nová'
            }

            # Zkus nejprve -y → -a (pro Klíma, Procházka)
            candidate_a = obs[:-1] + 'a'
            if candidate_a.lower() in protected_a_surnames:
                return candidate_a

            # Heuristika: -y → -a pro příjmení jako Klíma
            if obs[:-1].lower().endswith(('klím', 'dvořák', 'svobod')):
                return candidate_a

            # Běžný případ: jen odstraň -y (Nováky → Novák)
            return obs[:-1]

    # ========== AKUZATIV/DATIV: -u → -a (příjmení končící na -a) ==========
    # Sýkoru → Sýkora, Klíchu → Klícha, Krejču → Krejča
    # Toto platí pro mužská příjmení končící na -a (Sýkora, Klícha, Jura, etc.)
    if lo.endswith('u') and len(obs) > 2:
        base = obs[:-1]
        base_lo = base.lower()

        # PRIORITA 1: Známá příjmení končící na -a
        known_a_surnames = {'sýkor', 'klích', 'krejč', 'hofman', 'jur',
                           'šíd', 'kord', 'hrab', 'pavelk', 'forejt'}
        if base_lo in known_a_surnames:
            return base + 'a'

        # PRIORITA 2: Heuristika - pokud base končí na typický vzor mužského příjmení na -a
        # Typické koncovky: -ka, -ra, -la, -na, -da, -ta, -ša, -ča
        if base_lo.endswith(('k', 'r', 'l', 'n', 'd', 't', 'š', 'č', 'ř', 'c', 'b', 'v', 'j')):
            # Kontrola délky: base musí být alespoň 3 znaky (Jur-u, ne Kr-u)
            if len(base) >= 3:
                return base + 'a'

        # Jinak nechej bez změny (může to být něco jiného)

    # ========== FINÁLNÍ KONTROLA: Kmeny potřebující -a ==========
    # Pokud příjmení je kmen který potřebuje -a na konci (Červink → Červinka)
    surname_stems_needing_a = {
        'kub', 'červink', 'hromádk', 'horčičk', 'strak',
        'kuč', 'bárt', 'procházk', 'klím', 'svobod'
    }
    if obs.lower() in surname_stems_needing_a:
        return obs + 'a'  # Červink → Červinka

    # ========== TYPO/OCR CORRECTION: Opravy běžných chyb ==========
    # Fiael → Fiala, Růžiček → Růžička (pokud je to typo)
    typo_corrections = {
        'fiael': 'fiala',
        'fial': 'fiala',
        'růžiček': 'růžička',
        'ruzicek': 'růžička',
        'novk': 'novák',
        'dvork': 'dvořák',
        'prochzka': 'procházka',
        'cern': 'černý',
        'horak': 'horák',
    }
    if lo in typo_corrections:
        return typo_corrections[lo].capitalize()

    return obs

# =============== Varianty pro nahrazování ===============
def variants_for_first(first: str) -> set:
    """
    Generuje všechny pádové varianty křestního jména včetně:
    - Nominativ, Genitiv, Dativ, Akuzativ, Vokativ, Lokál, Instrumentál
    - Přivlastňovací přídavná jména (Petrův, Janin, atd.)
    """
    f = first.strip()
    if not f: return {''}
    V = {f, f.lower(), f.capitalize()}
    low = f.lower()

    # ========== Ženská jména končící na -a ==========
    if low.endswith('a'):
        stem = f[:-1]
        # Základní pády: Gen/Dat/Akuz/Vok/Lok/Instr
        V |= {stem+'y', stem+'e', stem+'ě', stem+'u', stem+'ou', stem+'o'}

        # Přivlastňovací přídavná jména (Janin dům, Petřina kniha)
        V |= {stem+s for s in ['in','ina','iny','iné','inu','inou','iným','iných','ino']}

        # Speciální případy pro měkčení (Petra → Petře, Veronka → Verunce)
        if stem.endswith('k'):
            V.add(stem[:-1] + 'c' + 'e')
            V.add(stem[:-1] + 'c' + 'i')

        # Speciální měkčení tr → tř (Petra → Petřin)
        if stem.endswith('tr'):
            soft_stem = stem[:-1] + 'ř'
            V |= {soft_stem+s for s in ['in','ina','iny','iné','inu','inou','iným','iných','ino']}

        # Speciální měkčení h → z, ch → š, k → c, r → ř
        if stem.endswith('h'):
            soft_stem = stem[:-1] + 'z'
            V.add(soft_stem + 'e')
            V.add(soft_stem + 'i')
        if stem.endswith('ch'):
            soft_stem = stem[:-2] + 'š'
            V.add(soft_stem + 'e')
            V.add(soft_stem + 'i')
        if stem.endswith(('k', 'g')):
            soft_stem = stem[:-1] + 'c'
            V.add(soft_stem + 'e')
            V.add(soft_stem + 'i')
        if stem.endswith('r') and not stem.endswith('tr'):
            soft_stem = stem[:-1] + 'ř'
            V.add(soft_stem + 'e')
            V.add(soft_stem + 'i')

    # ========== Ženská jména končící na -e/-ie (Alice, Otilie, Julie) ==========
    elif low.endswith(('e', 'ie')):
        # Pro jména jako Alice, Otilie, Julie
        # Dativ/Lokál: Alici, Otilii
        # Akuzativ: Alici, Otilii
        # Genitiv: Alice/Alici, Otilie/Otilii
        # Instrumentál: Alicí, Otilií
        # Vokativ: Alice, Otilie

        # Alice -> Alic + i, Alic + í
        if low.endswith('e') and not low.endswith('ie'):
            stem = f[:-1]  # Alice -> Alic
            V |= {stem+'i', stem+'í', stem+'ii'}
            # Přivlastňovací: Alicin
            V |= {stem+s for s in ['in','ina','iny','iné','inu','inou','iným','iných','ino']}

        # Otilie -> Otili + i, Otili + í
        elif low.endswith('ie'):
            stem = f[:-1]  # Otilie -> Otili
            V |= {stem+'i', stem+'í', stem+'ii'}
            # Přivlastňovací: Otiliin
            base = f[:-2]  # Otilie -> Otil
            V |= {base+s for s in ['in','ina','iny','iné','inu','inou','iným','iných','ino']}

    # ========== Mužská jména ==========
    else:
        # Základní pády
        V |= {f+'a', f+'ovi', f+'e', f+'em', f+'u', f+'om'}

        # Přivlastňovací přídavná jména (Petrův dům, Petrova kniha)
        V |= {f+'ův', f+'ova', f+'ovo', f+'ovu', f+'ovou', f+'ově'}
        V |= {f+'ov'+s for s in ['a','o','y','ě','ým','ých','ou','u','e']}

        # Speciální případy pro zakončení -ek, -el
        if low.endswith('ek'):
            stem_k = f[:-2] + 'k'
            V |= {stem_k+'a', stem_k+'ovi', stem_k+'em', stem_k+'u', stem_k+'e'}
            V.add(f[:-2] + 'ka')

        if low.endswith('el'):
            stem_l = f[:-2] + 'l'
            V |= {stem_l+'a', stem_l+'ovi', stem_l+'em', stem_l+'u', stem_l+'e'}
            V.add(f[:-2] + 'la')

        # Speciální případy pro zakončení -ec
        if low.endswith('ec'):
            stem_c = f[:-2] + 'c'
            V |= {stem_c+'e', stem_c+'i', stem_c+'em', stem_c+'u'}

        # Speciální případ: Jiří → Jiřího, Jiřímu, Jiřím, Jiřího
        if low.endswith('í'):
            stem = f[:-1]
            V |= {stem+'ího', stem+'ímu', stem+'ím', stem+'íh'}

        # Speciální případ: -iš/-aš → měkčení (Lukáš, Tomáš)
        if low.endswith(('áš', 'iš')):
            stem_base = f[:-1]
            V |= {stem_base+'e', stem_base+'i', stem_base+'em', stem_base+'ovi'}

        # Lokál s měkčením
        if not low.endswith(('i', 'í')):
            V |= {f+'ovi', f+'e'}

    return V

def variants_for_surname(surname: str) -> set:
    """
    Generuje všechny pádové varianty příjmení včetně:
    - Všechny pády jednotného i množného čísla
    - Přivlastňovací přídavná jména (Novákův, Novákova)
    - Speciální případy pro -ová, -ský, -ek, -ec, atd.
    """
    s = surname.strip()
    if not s: return {''}
    out = {s, s.lower(), s.capitalize()}
    low = s.lower()

    # ========== Příjmení typu -ová (ženská) ==========
    if low.endswith('ová'):
        base = s[:-1]
        out |= {
            s,
            base+'é',
            base+'ou',
            base+'á',
        }
        base_stem = s[:-3]
        out |= {
            base_stem+'ových',
            base_stem+'ovým',
            base_stem+'ové',
        }
        return out

    # ========== Příjmení typu -ský/-cký (přídavná jména) ==========
    if low.endswith(('ský','cký')):
        stem = s[:-2]
        out |= {
            stem+'ý', stem+'ého', stem+'ému', stem+'ým', stem+'ém',
            stem+'á', stem+'é', stem+'ou',
            stem+'ých', stem+'ými'
        }
        return out

    # ========== Obecná přídavná jména končící na -ý ==========
    if low.endswith('ý'):
        stem = s[:-1]
        out |= {
            stem+'ý', stem+'ého', stem+'ému', stem+'ým', stem+'ém',
            stem+'á', stem+'é', stem+'ou',
            stem+'ých', stem+'ými'
        }
        return out

    # ========== Ženská příjmení na -á (ne -ová) ==========
    if low.endswith('á') and not low.endswith('ová'):
        stem = s[:-1]
        out |= {s, stem+'é', stem+'ou', stem+'á'}
        return out

    # ========== Příjmení typu -ek (Dvořáček, Hájek) ==========
    if low.endswith('ek') and len(s) >= 3:
        stem_k = s[:-2] + 'k'
        out |= {
            s,
            stem_k+'a', stem_k+'ovi', stem_k+'em', stem_k+'u', stem_k+'e',
            stem_k+'y', stem_k+'ou',
        }
        out |= {
            stem_k+'ův', stem_k+'ova', stem_k+'ovo',
            stem_k+'ovu', stem_k+'ovou', stem_k+'ově'
        }
        out |= {
            stem_k+'ů', stem_k+'ům', stem_k+'y',
        }
        return out

    # ========== Příjmení typu -ec (Němec, Konec) ==========
    if low.endswith('ec') and len(s) >= 3:
        stem_c = s[:-2] + 'c'
        out |= {
            s,
            stem_c+'e', stem_c+'i', stem_c+'em', stem_c+'u', stem_c+'y',
        }
        out |= {
            stem_c+'ů', stem_c+'ům', stem_c+'ích', stem_c+'ech', stem_c+'emi',
        }
        out |= {
            stem_c+'ův', stem_c+'ova', stem_c+'ovo',
            stem_c+'ovu', stem_c+'ovou', stem_c+'ově'
        }
        return out

    # ========== Příjmení na -a (mužská i ženská) ==========
    if low.endswith('a') and len(s) >= 2 and not low.endswith('ová'):
        stem = s[:-1]
        out |= {
            s,
            stem+'y', stem+'ovi', stem+'ou', stem+'u', stem+'e', stem+'o',
        }
        out |= {
            stem+'ův', stem+'ova', stem+'ovo',
            stem+'ovu', stem+'ovou', stem+'ově'
        }
        out |= {
            stem+'ů', stem+'ům', stem+'y',
        }
        return out

    # ========== Obecná mužská příjmení (konsonantní kmeny) ==========
    out |= {
        s+'a', s+'ovi', s+'e', s+'em', s+'u',
    }
    out |= {
        s+'ův', s+'ova', s+'ovo',
        s+'ovu', s+'ovou', s+'ově'
    }
    out |= {
        s+'ov'+suf for suf in ['a','o','y','ě','ým','ých','ou','u','e','i']
    }
    out |= {
        s+'ů', s+'ům', s+'y', s+'ích', s+'ech',
    }

    return out

# =============== Regexy ===============

# Vylepšený ADDRESS_RE - zachytává adresy v JAKÉMKOLIV formátu
# Podporuje všechny možné kombinace:
# 1. "bytem ulice číslo, PSČ město" (např. "bytem Lázeňská 145/8, 415 01 Teplice")
# 2. "bytem ulice číslo, město PSČ" (např. "bytem Mánesova 89/12, Chomutov 430 01")
# 3. "ulice číslo, město" (např. "Slezská 67, Opava" - bez PSČ, BEZ prefixu - město POVINNÉ)
# 4. "bytem ulice číslo" (např. "bytem Karlova 12" - s prefixem, město volitelné)
ADDRESS_RE = re.compile(
    r'(?<!\[)'
    r'(?:'
        # SKUPINA 1: S prefixem - město je VOLITELNÉ
        r'(?:'
            r'(?:(?:trvale\s+)?bytem\s+|'
            r'(?:trvalé\s+)?bydlišt[eě]\s*:\s*|'
            r'(?:sídlo(?:\s+podnikání)?|se\s+sídlem)\s*:\s*|'
            r'(?:místo\s+podnikání)\s*:\s*|'
            r'(?:adresa|trvalý\s+pobyt|na\s+adrese)\s*:\s*|'
            r'(?:v\s+ulic[ií]|na\s+ulici|v\s+dom[eě])\s+)'
            r'[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ]'
            r'[a-záčďéěíňóřšťúůýž\s]{2,50}?'
            r'\s+\d{1,4}(?:/\d{1,4})?'  # Číslo domu (povinné)
            r'(?:'  # Nepovinná část s městem/PSČ (jen s prefixem!)
                r',\s*'
                r'(?:'
                    # PSČ město
                    r'\d{3}\s?\d{2}'
                    r'[ \t]+'
                    r'[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž \t]{1,30}'
                    r'(?:[ \t]+\d{1,2})?'
                r'|'
                    # město PSČ
                    r'[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž \t]{1,30}'
                    r'[ \t]+'
                    r'\d{3}\s?\d{2}'
                r'|'
                    # jen město
                    r'[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž \t]{1,30}'
                    r'(?:[ \t]+\d{1,2})?'
                r')'
            r')?'
        r')'
    r'|'
        # SKUPINA 2: BEZ prefixu - město je POVINNÉ (aby se nechytaly náhodná slova)
        r'(?:'
            r'[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ]'
            r'[a-záčďéěíňóřšťúůýž\s]{2,50}?'
            r'\s+\d{1,4}(?:/\d{1,4})?'  # Číslo domu
            r',\s*'  # Čárka (povinná!)
            r'(?:'
                # PSČ město
                r'\d{3}\s?\d{2}'
                r'[ \t]+'
                r'[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž \t]{1,30}'
                r'(?:[ \t]+\d{1,2})?'
            r'|'
                # město PSČ
                r'[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž \t]{1,30}'
                r'[ \t]+'
                r'\d{3}\s?\d{2}'
            r'|'
                # jen město (BEZ PSČ, ale město MUSÍ být!)
                r'[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž \t]{1,30}'
                r'(?:[ \t]+\d{1,2})?'
            r')'
        r')'
    r')'
    r'(?=\s|$|,|\.|;|:|\n|\r|Rodné|IČO|DIČ|Tel|E-mail|Kontakt|OP|Datum|Narozen|\))',
    re.IGNORECASE | re.UNICODE
)

# SPZ/RZ (Státní poznávací značky)
# České SPZ formáty:
# - Formát XYZ NNNN: 4A5 6789, 1P2 3456 (číslice-písmeno-číslice mezera 4 číslice)
# - Formát XYY NNNN: 1AB 2345 (číslice-2písmena mezera 4 číslice)
# - S prefixem: "SPZ: ...", "RZ: ...", "reg. značka: ..."
LICENSE_PLATE_RE = re.compile(
    r'(?:'
    r'(?:SPZ|RZ|reg\.?\s*(?:značka|číslo)?)\s*:?\s*'  # Volitelný prefix
    r')?'
    r'('  # Capture group pro samotnou SPZ
    r'\d[A-Z]\d\s+\d{4}|'  # 4A5 6789 (číslice-písmeno-číslice mezera 4číslice)
    r'\d[A-Z]{2}\s+\d{4}|'  # 1AB 2345 (číslice-2písmena mezera 4číslice)
    r'\d[A-Z]{2}\d{4}|'  # 1AB2345 (bez mezery)
    r'[A-Z]{2}\d{4,5}'  # AB12345
    r')',
    re.IGNORECASE
)

# VIN (Vehicle Identification Number) - 17 znaků
# Formát: TMBCF61Z0L7654321, 1HGBH41JXMN109186
VIN_RE = re.compile(
    r'(?:VIN|Vehicle\s+ID|Identifikační\s+číslo\s+vozidla)\s*[:\-]?\s*([A-HJ-NPR-Z0-9]{17})\b|'
    r'\b([A-HJ-NPR-Z0-9]{17})\b(?=\s*(?:VIN|vozidlo|auto|vehicle))',
    re.IGNORECASE
)

# MAC adresa (Media Access Control)
# Formát: 00:1B:44:11:3A:B7, 00-1B-44-11-3A-B7, 001B.4411.3AB7
MAC_RE = re.compile(
    r'(?:MAC\s+(?:address|adresa)?)\s*[:\-]?\s*([0-9A-F]{2}[:\-][0-9A-F]{2}[:\-][0-9A-F]{2}[:\-][0-9A-F]{2}[:\-][0-9A-F]{2}[:\-][0-9A-F]{2})|'
    r'\b([0-9A-F]{2}[:\-][0-9A-F]{2}[:\-][0-9A-F]{2}[:\-][0-9A-F]{2}[:\-][0-9A-F]{2}[:\-][0-9A-F]{2})\b|'
    r'\b([0-9A-F]{4}\.[0-9A-F]{4}\.[0-9A-F]{4})\b',
    re.IGNORECASE
)

# IMEI (International Mobile Equipment Identity)
# Formát: 15 číslic (např. 123456789012345)
IMEI_RE = re.compile(
    r'(?:IMEI|International\s+Mobile\s+Equipment\s+Identity)\s*[:\-]?\s*(\d{15})\b|'
    r'\b(\d{15})\b(?=\s*(?:IMEI|mobil|telefon|mobile))',
    re.IGNORECASE
)

# IČO (8 číslic)
ICO_RE = re.compile(
    r'(?:IČO?\s*:?\s*)?(?<!\d)(\d{8})(?!\d)',
    re.IGNORECASE
)

# DIČ (CZ + 8-10 číslic)
DIC_RE = re.compile(
    r'\b(CZ\d{8,10})\b',
    re.IGNORECASE
)

# Rodné číslo (6 číslic / 3-4 číslice)
# DŮLEŽITÉ: Musí mít SILNÝ kontext (RČ, Rodné číslo, nar.) - PRIORITA!
# Regex má 2 capture groups - první pro context match, druhý pro standalone
BIRTH_ID_RE = re.compile(
    r'(?:'
    r'(?:RČ|Rodné\s+číslo|r\.?\s?č\.?|nar\.|narozen[aáý]?|Narození)\s*:?\s*(\d{6}/?\d{3,4})|'  # S kontextem (CAPTURE GROUP 1)
    r'(?<!FÚ-)(?<!KS-)(?<!VS-)(?<!čj-)(?<!\d)(\d{6}/\d{3,4})(?!\d)'  # Bez kontextu, ale ne po FÚ-/KS-/VS- (CAPTURE GROUP 2)
    r')',
    re.IGNORECASE
)

# Číslo OP (formát: AB 123456 nebo OP: 123456789)
# DŮLEŽITÉ: Musí být před PHONE_RE!
ID_CARD_RE = re.compile(
    r'(?:'
    r'\b([A-Z]{2}\s?\d{6})\b|'  # Standardní formát: AB 123456
    r'(?:OP|pas|pas\.|pas\.č\.|č\.OP)(?:\s+č\.?)?\s*[:\-]?\s*(\d{6,9})'  # OP: 123456789, OP č. 123456, Pas: 123456
    r')',
    re.IGNORECASE
)

# Email - OPRAVENO: Podpora pro diakritiku v lokální části
# Zachytí: martina.horáková@neoteam.cz, jan.novák@firma.cz, atd.
EMAIL_RE = re.compile(
    r'\b([a-zA-ZáčďéěíňóřšťúůýžÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b'
)

# Telefon (CZ formáty) - NESMÍ zachytit prefix do hodnoty!
# Prefix je mimo capture group, telefon je uvnitř
# DŮLEŽITÉ: Whitelist - NEchytej biometrické/technické prefixy!
# DŮLEŽITÉ: NEchytej částky (čísla následovaná Kč, EUR, USD)
PHONE_RE = re.compile(
    r'(?!'  # Negative lookahead - NEchytej pokud předchází:
    r'(?:IRIS_SCAN|VOICE_RK|HASH_BIO|FINGERPRINT|FACIAL_|RETINA_|PALM_|DNA_)_[A-Z0-9_]*'
    r')'
    r'(?:tel\.?|telefon|mobil|GSM)?\s*:?\s*'  # Volitelný prefix (MIMO capture group!)
    r'('  # START capture group - jen samotné číslo
    r'\+420\s?\d{3}\s?\d{3}\s?\d{3}|'  # +420 xxx xxx xxx
    r'\+420\s?\d{3}\s?\d{2}\s?\d{2}\s?\d{2}|'  # +420 xxx xx xx xx
    r'(?<!\d)\d{3}\s?\d{3}\s?\d{3}(?!\s*(?:Kč|EUR|USD|CZK)\b)(?!\d)|'  # xxx xxx xxx (NE pokud následuje měna!)
    r'(?<!\d)\d{3}\s?\d{2}\s?\d{2}\s?\d{2}(?!\d)'  # xxx xx xx xx
    r')',  # END capture group
    re.IGNORECASE
)

# Bankovní účet (formát: číslo/kód banky NEBO dlouhé číslo s kontextem)
# DŮLEŽITÉ: IBAN je samostatný regex níže!
# NESMÍ zachytit spisové značky (FÚ-xxx/xxxx, KS-xxx/xxxx, VS-xxx)
# DŮLEŽITÉ: Kód banky musí být součástí zachyceného účtu!
BANK_RE = re.compile(
    r'(?:'
    # Standardní formát s předčíslím: 3622-1234567890/0710
    r'(?<!FÚ-)(?<!KS-)(?<!VS-)(?<!čj-)(\d{1,6}-\d{6,16}/\d{4})|'
    # Standardní formát: 123456/2010, 1234567890/3210
    r'(?<!FÚ-)(?<!KS-)(?<!VS-)(?<!čj-)(\d{6,16}/\d{4})|'
    # S kontextem: "číslo účtu: 123456789/0800" - zachytí i s kódem pokud je
    r'(?:číslo\s+účtu|účet|účtu|platba\s+na\s+účet|bankovní\s+účet)\s*:?\s*(\d{6,16}(?:/\d{4})?)'
    r')\b',
    re.IGNORECASE
)

# IBAN (mezinárodní formát bankovního účtu)
# CZ IBAN: CZ + 2 číslice + 20 číslic = 24 znaků celkem
# KRITICKÉ: Musí být před CARD_RE, protože obsahuje dlouhé sekvence číslic!
IBAN_RE = re.compile(
    r'\b(CZ\d{2}(?:\s?\d{4}){5})\b',
    re.IGNORECASE
)

# Datum (DD.MM.YYYY nebo DD. MM. YYYY)
# Obecný pattern pro všechna data
DATE_RE = re.compile(
    r'\b(\d{1,2}\.\s?\d{1,2}\.\s?\d{4})\b'
)

# Datum narození (DOB) - specificky pro "Datum narození:" kontext
DOB_RE = re.compile(
    r'(?:Datum\s+narození|Narozen[aý]?|Nar\.|Narození)\s*:?\s*(\d{1,2}\.\s?\d{1,2}\.\s?\d{4})\b',
    re.IGNORECASE
)

# Datum slovně (např. "15. března 2024")
DATE_WORDS_RE = re.compile(
    r'\b(\d{1,2}\.\s?(?:ledna|února|března|dubna|května|června|července|srpna|září|října|listopadu|prosince)\s?\d{4})\b',
    re.IGNORECASE
)

# Hesla a credentials (KRITICKÉ - hodnotu neukládat!)
PASSWORD_RE = re.compile(
    r'(?:password|heslo|passwd|pwd|pass)\s*[:\-=]\s*([^\s,;\.]{3,50})',
    re.IGNORECASE
)

# Credentials pattern - "Credentials: username / password"
CREDENTIALS_RE = re.compile(
    r'(?i)\b(credentials?|login|přihlašovací\s+údaje)\s*:\s*([A-Za-z0-9._\-@]+)\s*/\s*(\S+)',
    re.IGNORECASE
)

# API klíče, Secrets, Tokens (KRITICKÉ - hodnotu neukládat!)
API_KEY_RE = re.compile(
    r'(?:AWS\s+)?(?:Access\s+)?(?:Key(?:\s+ID)?|Secret(?:\s+Access\s+Key)?|Token|API[_\s]?Key)\s*[:\-=]\s*([A-Za-z0-9+/=_\-]{16,})',
    re.IGNORECASE
)

SECRET_RE = re.compile(
    r'(?:Stripe|SendGrid|GitHub|Secret|Client[_\s]?Secret|Access\s+Token|Personal\s+Access\s+Token)\s*[:\-=]\s*([A-Za-z0-9+/=_\-]{16,})',
    re.IGNORECASE
)

SSH_KEY_RE = re.compile(
    r'(?:ssh-rsa|ssh-ed25519|ecdsa-sha2-nistp256)\s+([A-Za-z0-9+/=]{50,})',
    re.IGNORECASE
)

# Usernames, Account IDs, Hostnames
USERNAME_RE = re.compile(
    r'(?:Login|Username|Uživatel|User|GitHub|Jira|AWS\s+Console)\s*[:\-=]\s*([A-Za-z0-9._\-@]+)',
    re.IGNORECASE
)

ACCOUNT_ID_RE = re.compile(
    r'(?:Account\s+ID|AWS\s+Account)\s*[:\-=]\s*(\d{12})',
    re.IGNORECASE
)

HOSTNAME_RE = re.compile(
    r'(?:hostname|host|server|RDS|endpoint)\s*[:\-=]\s*([a-z0-9\-\.]+\.[a-z]{2,})',
    re.IGNORECASE
)

# IP adresy (IPv4)
IP_RE = re.compile(
    r'\b(?<!\d\.)(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})(?!\.\d)\b'
)

# Platební karty (13-19 číslic)
# Rozšířený pattern: Visa/MC (16), AmEx (15), Diners (14), atd.
# DŮLEŽITÉ: IBAN se zpracovává PŘED tímto regexem!
# DŮLEŽITÉ: Zachytí i "Číslo:" když je to 16-19 číslic (typicky karta)
CARD_RE = re.compile(
    r'(?:'
    # S prefixem: "Číslo karty:", "Karta:", "Card Number:", "Číslo:" (když 16 číslic)
    r'(?:Číslo\s+(?:platební\s+)?karty|Číslo|(?:Platební\s+)?(?:Karta|Card)(?:\s+\d+)?(?:\s+Number)?)\s*[:\-=]?\s*'
    r'('
    r'\d{4}[\s\-]?\d{6}[\s\-]?\d{5}|'  # AmEx: 4-6-5 (15 číslic)
    r'\d{4}[\s\-]?\d{4}[\s\-]?\d{4}[\s\-]?\d{4}(?:[\s\-]?\d{2,3})?|'  # Visa/MC: 16-19
    r'\d{4}[\s\-]?\d{6}[\s\-]?\d{4}'  # Diners: 4-6-4 (14 číslic)
    r')'
    r')',
    re.IGNORECASE
)

# Čísla pojištěnce
INSURANCE_ID_RE = re.compile(
    r'(?:Číslo\s+pojištěnce|Pojišťovna|VZP|ČPZP|ZPŠ|OZP)[,\s:]+(?:číslo\s*:?\s*)?(\d{10})',
    re.IGNORECASE
)

# RFID/Badge čísla
RFID_RE = re.compile(
    r'(?:RFID\s+karta|badge|ID\s+karta)\s*[:#]?\s*([A-Za-z0-9\-_/]+)',
    re.IGNORECASE
)

# ========== SOCIAL MEDIA (KRITICKÉ - PII) ==========

# LinkedIn profily
LINKEDIN_RE = re.compile(
    r'(?:LinkedIn|linkedin)?\s*:?\s*(https?://(?:www\.)?linkedin\.com/in/[A-Za-z0-9\-_]+)',
    re.IGNORECASE
)

# Facebook profily
FACEBOOK_RE = re.compile(
    r'(?:Facebook|facebook)?\s*:?\s*(https?://(?:www\.)?facebook\.com/[A-Za-z0-9\._\-]+)',
    re.IGNORECASE
)

# Instagram handle - POUZE s explicitním kontextem nebo URL
# NEchytej @ z emailů!
INSTAGRAM_RE = re.compile(
    r'(?:Instagram|instagram)\s*:?\s*(@[A-Za-z0-9_]+)|'  # Handle pouze s prefixem "Instagram:"
    r'(https?://(?:www\.)?instagram\.com/[A-Za-z0-9_\.]+)',  # Nebo plné URL
    re.IGNORECASE
)

# Skype ID
SKYPE_RE = re.compile(
    r'(?:Skype|skype)\s*:?\s*([A-Za-z0-9\._\-]+)',
    re.IGNORECASE
)

# ========== BIOMETRIC IDs (KRITICKÉ - GDPR Článek 9) ==========

# Voice ID / Hlasový profil
VOICE_ID_RE = re.compile(
    r'(?:Hlasový\s+profil|Voice\s+ID|VOICE_ID|Voice\s+Profile)\s*[:\-=]?\s*([A-Z0-9_\-]+)',
    re.IGNORECASE
)

# Biometric hash (otisk prstu, sítnice, atd.)
# Zachytí pouze hodnoty po explicitním "hash:", "Hash:", nebo samostatné hash kódy
BIO_HASH_RE = re.compile(
    r'(?:hash|Hash):\s*([A-Z][A-Z0-9_\-]{10,})|'  # hash: HASH_BIO_JP_2024_0156
    r'\b((?:HASH_BIO|IRIS|RETINA|FINGERPRINT|PALM|DNA)_[A-Z0-9_\-]{8,})\b'  # Standalone hash codes
)

# Photo ID / Face ID files
PHOTO_ID_RE = re.compile(
    r'(photo_id_[A-Za-z0-9_\-]+\.(?:jpg|jpeg|png|gif|bmp))|'
    r'(face_id_[A-Za-z0-9_\-]+\.(?:jpg|jpeg|png|gif|bmp))|'
    r'(?:Fotografie|Photo\s+ID|Face\s+ID)\s*[:\-=]?\s*[Uu]loženo.*?\(([A-Za-z0-9_\-]+\.(?:jpg|jpeg|png|gif|bmp))\)',
    re.IGNORECASE
)

# Enhanced API Key - zachytí i složitější formáty
# Minimální délka 8 znaků, aby se předešlo false positives jako "e" nebo "ERP"
API_KEY_ENHANCED_RE = re.compile(
    r'(?:API\s+klíč|API\s+Key|api_key)\s*[:\-=]?\s*([A-Za-z0-9_\-]{8,})',
    re.IGNORECASE
)

# Genetické identifikátory (rs...) - NESMÍ být zachyceny jako ICO!
# Pattern: rs28897696, rs1234567
GENETIC_ID_RE = re.compile(
    r'\b(rs\d{6,})\b',
    re.IGNORECASE
)

# ========== DOPLNĚNÉ GDPR KATEGORIE ==========

# Datum narození (samostatné, ne v rodném čísle)
# Formáty: dd.mm.yyyy, dd/mm/yyyy, dd-mm-yyyy
# S kontextem: "datum narození:", "nar.", "narozen(a)"
BIRTH_DATE_RE = re.compile(
    r'(?:datum\s+narození|nar\.|narozen[aáý]?)\s*[:\-]?\s*'
    r'(\d{1,2}[\./\-]\d{1,2}[\./\-]\d{4})',
    re.IGNORECASE
)

# Číslo pasu
# Formáty: 12345678 (8 číslic), AB123456 (2 písmena + 6 číslic)
PASSPORT_RE = re.compile(
    r'(?:pas|passport|č\.\s*pasu)\s*(?:č\.)?\s*[:\-]?\s*([A-Z]{0,2}\d{6,9})\b',
    re.IGNORECASE
)

# Číslo řidičského průkazu
# Formáty: AB123456, 12345678
DRIVER_LICENSE_RE = re.compile(
    r'(?:ŘP|řidičák|řidičský\s+průkaz|driver\'?s?\s*license)\s*(?:č\.)?\s*[:\-]?\s*([A-Z]{0,2}\d{6,9})\b',
    re.IGNORECASE
)

# Benefitní karty (MultiSport, Sodexo, Edenred, atd.) - DŮLEŽITÉ PII!
# Formáty: 9876543210, MS-123456, SOD/123456, "ID karty: 9876543210"
# Důvod přidání: Unikátní identifikátor osoby, jednoznačně PII
BENEFIT_CARD_RE = re.compile(
    r'(?:'
    r'(?:MultiSport|Sodexo|Edenred|benefitní\s+karta|benefit\s+card)\s*(?:karta|č\.?|ID)?\s*[:\-]?\s*([A-Z]{0,3}[\-/]?\d{6,12})|'
    r'ID\s+karty\s*[:\-]\s*(\d{6,12})'
    r')\b',
    re.IGNORECASE
)

# =============== Třída Anonymizer ===============
class Anonymizer:
    def __init__(self, verbose=False):
        self.verbose = verbose
        self.counter = defaultdict(int)
        self.canonical_persons = []  # list of {first, last, tag}
        self.person_index = {}  # (first_norm, last_norm) -> tag
        self.person_canonical_names = {}  # tag -> canonical full name (as it appears in document)
        self.person_variants = {}  # tag -> set of all variants
        self.entity_map = defaultdict(lambda: defaultdict(set))  # typ -> original -> varianty
        self.entity_index_cache = defaultdict(dict)  # OPTIMIZATION: typ -> original -> idx cache
        self.entity_reverse_map = defaultdict(dict)  # OPTIMIZATION: typ -> variant -> original
        self.source_text = ""  # Store original text for validation
        self._regex_cache = {}  # PERFORMANCE: Cache compiled regex patterns

    def _get_or_create_label(self, typ: str, original: str, store_value: bool = True) -> str:
        """Vrátí existující nebo vytvoří nový štítek pro entitu.

        Args:
            typ: Typ entity (PASSWORD, EMAIL, atd.)
            original: Původní hodnota
            store_value: False pro citlivá data (hesla, API klíče) - uloží jen "***REDACTED***"
        """
        # Normalizace
        orig_norm = original.strip()

        # Speciální cleanup pro ADDRESS - odstraň prefixy "Sídlo:", "Trvalé bydliště:", "Trvalý pobyt:" atd.
        if typ == 'ADDRESS':
            # Odstraň prefix s dvojtečkou (Sídlo:, Adresa:, atd.)
            orig_norm = re.sub(r'^(Sídlo|Trvalé\s+bydliště|Trvalý\s+pobyt|Bydliště|Adresa|Místo\s+podnikání|Se\s+sídlem|Bytem)\s*:\s*', '', orig_norm, flags=re.IGNORECASE)
            # Odstraň prefix bez dvojtečky na začátku (adrese, bytem) - instrumentál/lokál
            orig_norm = re.sub(r'^(adrese|adresa|bytem|bydlišti|sídle)\s+', '', orig_norm, flags=re.IGNORECASE)

        # OPTIMIZATION: Use reverse_map for O(1) lookup instead of O(n) iteration
        # Check if this variant already exists
        if orig_norm in self.entity_reverse_map[typ]:
            existing_orig = self.entity_reverse_map[typ][orig_norm]
            existing_idx = self.entity_index_cache[typ][existing_orig]
            return f"[[{typ}_{existing_idx}]]"

        # Vytvoř nový
        idx = len(self.entity_map[typ]) + 1

        # Pro citlivá data: unikátní placeholder pro každý item
        # Pro běžná data: ukládej skutečnou hodnotu
        if store_value:
            map_key = orig_norm
        else:
            # Každý citlivý item má unikátní klíč, ale zobrazí se jako ***REDACTED***
            map_key = f"***REDACTED_{idx}***"

        self.entity_map[typ][map_key].add(orig_norm)
        self.entity_index_cache[typ][map_key] = idx  # Cache the index
        self.entity_reverse_map[typ][orig_norm] = map_key  # Reverse lookup
        return f"[[{typ}_{idx}]]"

    def _normalize_for_matching(self, text: str) -> str:
        """Normalizuje text pro porovnávání (odstranění diakritiky, lowercase)."""
        if not text: return ""
        n = unicodedata.normalize('NFD', text)
        no_diac = ''.join(c for c in n if not unicodedata.combining(c))
        return re.sub(r'[^A-Za-z]', '', no_diac).lower()

    def _ensure_person_tag(self, first_nom: str, last_nom: str, first_obs: str = None, last_obs: str = None) -> tuple:
        """Zajistí, že pro danou osobu existuje tag a vrátí ho spolu s canonical názvem.

        Args:
            first_nom: Inferred nominative first name (used for matching)
            last_nom: Inferred nominative last name (used for matching)
            first_obs: Observed first name from document (used as canonical if this is a new person)
            last_obs: Observed last name from document (used as canonical if this is a new person)

        Returns:
            tuple: (tag, canonical_full_name)
        """
        # FINÁLNÍ NORMALIZACE - aby canonical_persons obsahovali správné jména
        # Julia → Julie, Maria → Marie, atd.
        # DŮLEŽITÉ: Použij normalize_name_variant (NE infer_first_name_nominative)
        # protože nechceme vytvářet nové formy jmen, které nejsou v dokumentu!
        first_normalized = normalize_name_variant(first_nom) if first_nom else first_nom

        key = (self._normalize_for_matching(first_normalized), self._normalize_for_matching(last_nom))

        if key in self.person_index:
            tag = self.person_index[key]
            # Return existing tag and its canonical name
            canonical_full = self.person_canonical_names[tag]
            return tag, canonical_full

        # Vytvoř nový tag
        self.counter['PERSON'] += 1
        tag = f'[[PERSON_{self.counter["PERSON"]}]]'

        # Ulož do indexu s normalizovaným jménem
        self.person_index[key] = tag

        # CANONICAL: Always use INFERRED nominative initially
        # Post-processing will fix any that aren't in document
        canonical_first = first_normalized
        canonical_last = last_nom
        canonical_full = f'{canonical_first} {canonical_last}'

        self.canonical_persons.append({'first': canonical_first, 'last': canonical_last, 'tag': tag})
        self.person_canonical_names[tag] = canonical_full  # Store canonical name for this tag

        # Vygeneruj všechny pádové varianty (použij normalizované jméno pro generování variant)
        fvars = variants_for_first(first_normalized)
        svars = variants_for_surname(last_nom)
        self.person_variants[tag] = {f'{f} {s}' for f in fvars for s in svars}

        # Store index and reverse map using canonical name (not inferred!)
        self.entity_index_cache['PERSON'][canonical_full] = self.counter['PERSON']
        self.entity_reverse_map['PERSON'][canonical_full] = canonical_full

        return tag, canonical_full

    def _apply_known_people(self, text: str) -> str:
        """Aplikuje známé osoby (již detekované) - nahrazuje všechny pádové varianty stejným tagem."""
        # FÁZE 1: Nahrazení plných jmen (křestní + příjmení)
        for p in self.canonical_persons:
            tag = p['tag']

            # Pro každou variantu této osoby (seřazeno od nejdelší)
            for pat in sorted(self.person_variants[tag], key=len, reverse=True):
                # FILTR: Odmítni zkrácené genitivy
                parts = pat.split()
                if len(parts) == 2:
                    fv, lv = parts
                    fv_lo = fv.lower()

                    # Pokud křestní jméno má 3-5 znaků a končí na 'k' → zkrácený genitiv
                    if 3 <= len(fv) <= 5 and fv_lo[-1] == 'k':
                        continue

                    # Pokud křestní jméno má 3 znaky a nekončí na samohlásku/n/l/r → zkrácený
                    if len(fv) == 3 and not fv_lo[-1] in 'aeiouyáéěíóúůýnlr':
                        continue

                # PERFORMANCE: Cache compiled regex patterns
                if pat not in self._regex_cache:
                    self._regex_cache[pat] = re.compile(r'(?<!\w)'+re.escape(pat)+r'(?!\w)', re.IGNORECASE)
                rx = self._regex_cache[pat]

                def repl(m):
                    surf = m.group(0)
                    # Zaznamenej tuto variantu
                    canonical = f'{p["first"]} {p["last"]}'
                    self.entity_map['PERSON'][canonical].add(surf)
                    return tag

                text = rx.sub(repl, text)

        return text

    def _replace_remaining_people(self, text: str) -> str:
        """Detekuje a nahradí zbývající osoby.

        FÁZE detekce (v pořadí):
        1. FÁZE 3.5: Maiden names - (rozená Novotná), (dříve Svobodová)
        2. FÁZE 3: Samostatná příjmení - "Novák uvedl", "pan Dvořák"
        3. FÁZE 3.7: Samostatná křestní jména - "Jakub pracoval", "Eva řekla"
        4. Jména s titulem - "MUDr. Jan Novák"
        5. Běžná jména - "Jan Novák", "Eva Malá"
        """

        # ========== FÁZE 3.5: MAIDEN NAMES (RODNÁ PŘÍJMENÍ) ==========
        # Pattern: (rozená Novotná), (dříve Svobodová), (roz. Malá)
        # Tyto jsou velmi specifické a mají vysokou prioritu

        def replace_maiden_name(match):
            prefix = match.group(1)  # "rozená", "dříve", "roz."
            surname = match.group(2)  # "Novotná"

            # Odhadni nominativ
            surname_nom = infer_surname_nominative(surname)

            # Vytvoř tag pro příjmení (použijeme prázdné křestní jméno jako marker)
            tag, _ = self._ensure_person_tag("", surname_nom)

            return f"({prefix} {tag})"

        maiden_name_pattern = re.compile(
            r'\((rozená|rozenou|roz\.|dříve|dřív|původně)\s+([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)\)',
            re.UNICODE | re.IGNORECASE
        )

        text = maiden_name_pattern.sub(replace_maiden_name, text)

        # ========== FÁZE 3: SAMOSTATNÁ PŘÍJMENÍ ==========
        # Pattern: "Novák uvedl", "pan Dvořák", "Novákovi bylo", "od Maláové"
        # DŮLEŽITÉ: Nenahrazuj, pokud následuje křestní jméno (to je full name)

        def replace_standalone_surname(match):
            prefix = match.group(1) if match.groups()[0] else ""  # "pan", "paní", ""
            surname = match.group(2) if len(match.groups()) > 1 else match.group(1)

            # Odhadni nominativ
            surname_nom = infer_surname_nominative(surname)

            # Vytvoř tag (prázdné křestní jméno pro standalone příjmení)
            tag, _ = self._ensure_person_tag("", surname_nom)

            if prefix:
                return f"{prefix} {tag}"
            else:
                return tag

        # Pattern pro samostatné příjmení s kontextem
        # Zachytí: "pan Novák", "paní Malá", "Novák uvedl", "Novákovi bylo"
        standalone_surname_pattern = re.compile(
            r'(?:'
            r'(?:pan|paní|pana|paní|panu)\s+([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)|'  # pan Novák
            r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+ov[ia])\s+(?:uvedl|uvedla|řekl|řekla|byl|byla|měl|měla)|'  # Novákovi uvedl
            r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)\s+(?:uvedl|uvedla|řekl|řekla|potvrdil|potvrdila)'  # Novák uvedl
            r')',
            re.UNICODE | re.IGNORECASE
        )

        # POZOR: Tento pattern může zachytit i false positives, takže musíme být opatrní
        # Raději ho zatím zakomentujeme a přidáme později po testování
        # text = standalone_surname_pattern.sub(replace_standalone_surname, text)

        # ========== FÁZE 3.7: SAMOSTATNÁ KŘESTNÍ JMÉNA ==========
        # Pattern: "Jakub pracoval jako...", "Eva řekla...", ale NE "Praha", "Česká", atd.
        def replace_standalone_first_name(match):
            name = match.group(1)
            name_lower = name.lower()

            # Kontrola, zda je to křestní jméno z knihovny
            if name_lower not in CZECH_FIRST_NAMES:
                return match.group(0)

            # Ignore list - slova která vypadají jako jména, ale nejsou
            ignore_words = {
                'praha', 'brno', 'ostrava', 'plzeň', 'česká', 'slovenská',
                'evropa', 'amerika', 'asie', 'afrika', 'čech', 'moravia'
            }
            if name_lower in ignore_words:
                return match.group(0)

            # Vytvoř/najdi tag pro samostatné křestní jméno
            tag, _ = self._ensure_person_tag(name, "")

            return tag

        # Pattern pro samostatné křestní jméno následované slovesem nebo "jako"
        # Rozšířeno o uvozovky a další slovesa
        standalone_first_name_pattern = re.compile(
            r'(?:^|["\s])([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)\s+(?:pracoval|pracovala|řekl|řekla|uvedl|uvedla|jako|byl|byla|je|jsou|měl|měla|dělal|dělala)',
            re.UNICODE | re.MULTILINE
        )

        def replace_standalone_wrapper(match):
            # Zachovej prefix (uvozovky nebo mezeru)
            prefix = match.group(0)[0] if match.group(0)[0] in ('"', ' ', '\n', '\t') else ''
            result = replace_standalone_first_name(match)
            # Pokud bylo nahrazeno, přidej prefix
            if result != match.group(0):
                return prefix + result[len(prefix):] if prefix else result
            return match.group(0)

        text = standalone_first_name_pattern.sub(replace_standalone_wrapper, text)

        # DÁLE: Pattern pro jména s titulem (MUDr. Eva Malá)
        # Tento musí jít PŘED obecným pattern aby titul nebyl ztracen
        titled_pattern = re.compile(
            r'(Ing\.|Mgr\.|Bc\.|MUDr\.|JUDr\.|PhDr\.|RNDr\.|Prof\.|Doc\.|Ph\.D\.|MBA|CSc\.|DrSc\.)\s+'
            r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)\s+'
            r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)',
            re.UNICODE
        )

        def replace_titled(match):
            title = match.group(1)
            first = match.group(2)
            last = match.group(3)

            # ========== VALIDACE - STEJNÁ JAKO V replace_person() ==========

            # 1. Blacklist kritických slov
            critical_blacklist = {
                's.r.o.', 'a.s.', 'spol.', 'k.s.', 'v.o.s.', 'o.p.s.',
                'ltd', 'inc', 'corp', 'gmbh', 'llc',
                'czech', 'republic', 'synlab', 'gymnázium', 'gymnasium',
                'university', 'univerzita', 'fakulta', 'klinika', 'nemocnice',
                'centrum', 'ústav', 'institute', 'academy', 'akademie',
                'kaspersky', 'endpoint', 'latitude', 'archer', 'classic',
                'windows', 'linux', 'android', 'ios', 'office', 'excel',
                'ředitelka', 'ředitel', 'jednatel', 'jednatelka',
                'manager', 'director', 'chief', 'officer',
                'vyšetřující', 'vyšetřovatel', 'lékař', 'doktor', 'sestra'
            }

            combined = f"{first} {last}".lower()
            for word in critical_blacklist:
                if word in combined:
                    return match.group(0)  # Není osoba

            # 2. Role detection - pokud první slovo je role
            role_words = {
                'ředitelka', 'ředitel', 'jednatel', 'jednatelka',
                'manager', 'director', 'chief', 'officer',
                'specialist', 'consultant', 'coordinator',
                'developer', 'architect', 'engineer', 'analyst',
                'vyšetřující', 'vyšetřovatel', 'lékař', 'doktor'
            }
            if first.lower() in role_words:
                return match.group(0)  # Role, ne osoba

            # 3. Validace křestního jména
            first_lo = first.lower()
            common_czech_names = {'jan', 'petr', 'pavel', 'jiří', 'josef', 'tomáš', 'martin', 'jakub', 'david', 'daniel'}

            if first_lo not in CZECH_FIRST_NAMES and first_lo not in common_czech_names:
                # Zkrácené genitivy (Han, Elišk, Radk) - odmítnout
                if len(first) < 3:
                    return match.group(0)
                # Pokud má 3 znaky a nekončí na samohlásku ani n/l/r
                if len(first) == 3 and not first_lo[-1] in 'aeiouyáéěíóúůýnlr':
                    return match.group(0)
                # Zkrácené tvary končící na 'k' (4-5 znaků)
                if 4 <= len(first) <= 5 and first_lo[-1] == 'k':
                    return match.group(0)

            # Vytvoř/najdi tag pro osobu (s inferencí nominativu)
            last_nom = infer_surname_nominative(last)
            first_nom = infer_first_name_nominative(first) or first
            tag, canonical = self._ensure_person_tag(first_nom, last_nom)

            # Ulož původní formu jako variantu (pokud je jiná než kanonická)
            original_form = f"{first} {last}"
            if original_form.lower() != canonical.lower():
                self.entity_map['PERSON'][canonical].add(original_form)

            # Vrať titul + tag
            return f"{title} {tag}"

        # Nahraď titulované osoby NEJPRVE
        text = titled_pattern.sub(replace_titled, text)

        # Pak běžný pattern pro jména bez titulu
        titles = r'(?:Ing\.|Mgr\.|Bc\.|MUDr\.|JUDr\.|PhDr\.|RNDr\.|Ph\.D\.|MBA|CSc\.|DrSc\.)'

        # ========== NOVÝ: Pattern pro "Titul Jméno Příjmení" (3 slova) ==========
        # Tento pattern musí být PŘED běžným 2-slovným patternem!
        # Detekuje např: "Klient Ladislav Konečný", "Žadatel Jan Novák", atd.
        role_person_pattern = re.compile(
            r'\b([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)'  # Potenciální titul/role
            r'\s+'
            r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)'  # Křestní jméno
            r'\s+'
            r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)\b',  # Příjmení
            re.UNICODE
        )

        # Pattern pro jméno příjmení (2 slova)
        person_pattern = re.compile(
            r'\b([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)'
            r'\s+'
            r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)\b',
            re.UNICODE
        )

        # Všechny ignore_words pro rychlé ověření (budeme je potřebovat v obou handlerech)
        # Definujeme ignore_words ZDE, aby byl dostupný v obou funkcích
        ignore_words = {
            # Běžná slova ve smlouvách
            'místo', 'datum', 'částku', 'bytem', 'sídlo', 'adresa',
            'číslo', 'kontakt', 'telefon', 'email', 'rodné', 'narozena',
            'vydán', 'uzavřena', 'podepsána', 'smlouva', 'dohoda',
            # Místa
            'staré', 'město', 'nové', 'město', 'malá', 'strana',
            'václavské', 'náměstí', 'hlavní', 'nádraží',
            'karlovy', 'vary', 'karlova', 'var',  # Karlovy Vary city
            'hradec', 'hradci', 'králové',  # Hradec Králové city
            # Organizace/instituce klíčová slova
            'česká', 'spořitelna', 'komerční', 'banka', 'raiffeisen',
            'credit', 'bank', 'financial', 'global', 'senior',
            'junior', 'lead', 'chief', 'head', 'director',
            # Finance/Investment
            'capital', 'equity', 'value', 'crescendo', 'investment',
            'fund', 'holdings', 'partners', 'assets', 'portfolio',
            # Pozice/role
            'jednatel', 'jednatelka', 'ředitel', 'ředitelka',
            'auditor', 'manager', 'consultant', 'specialist',
            'assistant', 'coordinator', 'analyst',
            # CRITICAL: Tituly a role před jmény ve smlouvách
            'pan', 'paní', 'pán', 'pani',
            'pacient', 'pacientka', 'pacientek',
            'obžalovaný', 'obžalovaná', 'obžalované', 'obžalovaného',
            'zaměstnanec', 'zaměstnankyně', 'zaměstnance',
            'kupující', 'prodávající', 'prodávajícího',
            'stavebník', 'stavebníka',
            'investor', 'investora',
            'dlužník', 'dlužníka', 'věřitel', 'věřitele',
            'odsouzený', 'odsouzená', 'odsouzeného',
            # === KOMPLETNÍ SEZNAM PRÁVNÍCH, ZDRAVOTNICKÝCH A ADMINISTRATIVNÍCH TERMÍNŮ (605 termínů) ===
            # Právní, ekonomické, zdravotnické, pracovněprávní, rodinné, sociální, vzdělávací a ostatní termíny
            # které se mohou vyskytovat před jmény osob a nesmí být detekovány jako křestní jména
            'absolvent', 'absolventa', 'absolventce', 'absolventek', 'absolventem', 'absolventka', 'absolventkou', 'absolventky',
            'absolventovi', 'absolventů', 'azylant', 'azylanta', 'azylantce', 'azylantek', 'azylantem', 'azylantka',
            'azylantkou', 'azylantky', 'azylantovi', 'azylantu', 'azylantů', 'beneficient', 'beneficienta', 'beneficientem',
            'beneficientka', 'beneficientky', 'beneficientovi', 'beneficientů', 'cizince', 'cizincem', 'cizinci', 'cizinců',
            'cizinec', 'cizinek', 'cizinka', 'cizinkou', 'cizinky', 'dlužnic', 'dlužnice', 'dlužnicí',
            'dlužník', 'dlužníka', 'dlužníkem', 'dlužníkovi', 'dlužníku', 'dlužníky', 'dlužníků', 'dotčenou',
            'dotčená', 'dotčené', 'dotčeného', 'dotčenému', 'dotčený', 'dotčených', 'dotčeným', 'držitel',
            'držitelce', 'držitele', 'držitelek', 'držitelem', 'držiteli', 'držitelka', 'držitelkou', 'držitelky',
            'držitelů', 'dárce', 'dárci', 'dárců', 'dárkyni', 'dárkyní', 'dárkyně', 'důchodce',
            'důchodci', 'důchodců', 'důchodkyni', 'důchodkyně', 'důchodkyň', 'evidovanou', 'evidovaná', 'evidované',
            'evidovaného', 'evidovanému', 'evidovaný', 'evidovaných', 'evidovaným', 'hospitalizovanou', 'hospitalizovaná', 'hospitalizované',
            'hospitalizovaného', 'hospitalizovanému', 'hospitalizovaný', 'hospitalizovaných', 'hospitalizovaným', 'invalidní', 'invalidních', 'invalidního',
            'invalidním', 'invalidnímu', 'investor', 'investora', 'investorce', 'investorek', 'investorem', 'investorka',
            'investorkou', 'investorky', 'investorovi', 'investorů', 'investoři', 'jednotlivce', 'jednotlivcem', 'jednotlivci',
            'jednotlivců', 'jednotlivec', 'jmenovanou', 'jmenovaná', 'jmenované', 'jmenovaného', 'jmenovanému', 'jmenovaný',
            'jmenovaných', 'jmenovaným', 'klient', 'klienta', 'klientce', 'klientek', 'klientem', 'klientka',
            'klientkou', 'klientky', 'klientovi', 'klientu', 'klientů', 'kontrolovanou', 'kontrolovaná', 'kontrolované',
            'kontrolovaného', 'kontrolovanému', 'kontrolovaný', 'kontrolovaných', 'kontrolovaným', 'kupující', 'kupujících', 'kupujícího',
            'kupujícím', 'kupujícímu', 'léčenou', 'léčená', 'léčené', 'léčeného', 'léčenému', 'léčený',
            'léčených', 'léčeným', 'léčenými', 'manžel', 'manžela', 'manželce', 'manželek', 'manželem',
            'manželka', 'manželkou', 'manželky', 'manželovi', 'manželé', 'manželů', 'nemocnou', 'nemocná',
            'nemocné', 'nemocného', 'nemocném', 'nemocnému', 'nemocný', 'nemocných', 'nemocným', 'nemocnými',
            'nezletilou', 'nezletilá', 'nezletilé', 'nezletilého', 'nezletilému', 'nezletilý', 'nezletilých', 'nezletilým',
            'obviněnou', 'obviněná', 'obviněné', 'obviněného', 'obviněném', 'obviněnému', 'obviněný', 'obviněných',
            'obviněným', 'obviněnými', 'občan', 'občana', 'občance', 'občanek', 'občanem', 'občanka',
            'občankou', 'občanky', 'občanovi', 'občanu', 'občané', 'občanů', 'obžalovanou', 'obžalovaná',
            'obžalované', 'obžalovaného', 'obžalovaném', 'obžalovanému', 'obžalovaný', 'obžalovaných', 'obžalovaným', 'obžalovanými',
            'odsouzenou', 'odsouzená', 'odsouzené', 'odsouzeného', 'odsouzeném', 'odsouzenému', 'odsouzený', 'odsouzených',
            'odsouzeným', 'odsouzenými', 'opatrovnic', 'opatrovnice', 'opatrovnicí', 'opatrovník', 'opatrovníka', 'opatrovníkem',
            'opatrovníkovi', 'opatrovníku', 'opatrovníků', 'opatrovanec', 'opatrovance', 'opatrovanci', 'opatrovancem', 'opatrovanců',
            'opatrovanka', 'opatrovanky', 'opatrovankou', 'opatrovanek',
            # Náboženské a rodinné tituly
            'matka', 'matky', 'matce', 'matkou', 'matek', 'matkám', 'matkách', 'matkama',
            # Pěstounská péče
            'pěstoun', 'pěstouna', 'pěstounovi', 'pěstounu', 'pěstounem', 'pěstouni', 'pěstounů', 'pěstounům', 'pěstounech', 'pěstouny',
            'pěstounka', 'pěstounky', 'pěstounce', 'pěstounkou', 'pěstounek', 'pěstounkám', 'pěstounkách', 'pěstounkami',
            # Exekutoři / bailiffs
            'exekutor', 'exekutora', 'exekutorovi', 'exekutoru', 'exekutorem', 'exekutoři', 'exekutorů', 'exekutorům', 'exekutorech', 'exekutory',
            'exekutorka', 'exekutorky', 'exekutorce', 'exekutorkou', 'exekutorek', 'exekutorkám', 'exekutorkách', 'exekutorkami',
            'operovanou', 'operovaná', 'operované', 'operovaného', 'operovanému',
            'operovaný', 'operovaných', 'operovaným', 'operovanými', 'oprávněnou', 'oprávněná', 'oprávněné', 'oprávněného',
            'oprávněnému', 'oprávněný', 'oprávněných', 'oprávněným', 'osob', 'osoba', 'osobami', 'osobou',
            'osoby', 'osobám', 'osobě', 'osvojitel', 'osvojitelce', 'osvojitele', 'osvojitelem', 'osvojiteli',
            'osvojitelka', 'osvojitelky', 'osvojitelů', 'ošetřenou', 'ošetřená', 'ošetřené', 'ošetřeného', 'ošetřenému',
            'ošetřený', 'ošetřených', 'ošetřeným', 'ošetřenými', 'pachatel', 'pachatelce', 'pachatele', 'pachatelek',
            'pachateli', 'pachatelka', 'pachatelky', 'pachatelů', 'pacient', 'pacienta', 'pacientce', 'pacientek',
            'pacientem', 'pacientka', 'pacientkou', 'pacientky', 'pacientovi', 'pacientu', 'pacientů', 'pečující',
            'pečujících', 'pečujícího', 'pečujícím', 'pečujícímu', 'plátce', 'plátci', 'plátců', 'plátkyní',
            'plátkyně', 'podezřelou', 'podezřelá', 'podezřelé', 'podezřelého', 'podezřelém', 'podezřelému', 'podezřelý',
            'podezřelých', 'podezřelým', 'podezřelými', 'podnikatel', 'podnikatelce', 'podnikatele', 'podnikatelek', 'podnikateli',
            'podnikatelka', 'podnikatelky', 'podnikatelů', 'pojištěnce', 'pojištěncem', 'pojištěnci', 'pojištěnců', 'pojištěnec',
            'pojištěnek', 'pojištěnka', 'pojištěnkou', 'pojištěnky', 'poručnic', 'poručnice', 'poručnicí', 'poručník',
            'poručníka', 'poručníkem', 'poručníkovi', 'poručníku', 'poručníků', 'poškozenou', 'poškozená', 'poškozené',
            'poškozeného', 'poškozeném', 'poškozenému', 'poškozený', 'poškozených', 'poškozeným', 'poškozenými', 'prodávající',
            'prodávajících', 'prodávajícího', 'prodávajícím', 'prodávajícímu', 'propuštěnou', 'propuštěná', 'propuštěné', 'propuštěného',
            'propuštěném', 'propuštěnému', 'propuštěný', 'propuštěných', 'propuštěným', 'propuštěnými', 'přihlášenou', 'přihlášená',
            'přihlášené', 'přihlášeného', 'přihlášenému', 'přihlášený', 'přihlášených', 'přihlášeným', 'přijatou', 'přijatá',
            'přijatého', 'přijatému', 'přijatévé', 'přijatý', 'přijatých', 'přijatým', 'příjemce', 'příjemci',
            'příjemců', 'příjemkyní', 'příjemkyně', 'registrovanou', 'registrovaná', 'registrované', 'registrovaného', 'registrovanému',
            'registrovaný', 'registrovaných', 'registrovaným', 'rodič', 'rodiče', 'rodičem', 'rodiči', 'rodička',
            'rodičky', 'rodičů', 'rodina', 'rodině', 'rodiny',
            # Rodinné vztahy - VELMI ROZŠÍŘENÉ
            'dítě', 'dítěte', 'dítěti', 'dítětem', 'děti', 'dětí', 'dětem', 'dětmi',
            'sourozenec', 'sourozence', 'sourozenci', 'sourozencem', 'sourozenci', 'sourozencům', 'sourozencích', 'sourozenci',
            'sourozenka', 'sourozence', 'sourozenkou', 'sourozenci', 'sourozenek', 'sourozenkyně',
            # Matka - varianty (základní forma už existuje výše)
            'maminka', 'maminky', 'mamince', 'maminkou', 'maminek', 'maminkám', 'maminkách',
            'máma', 'mámy', 'mámě', 'mámou', 'mám', 'mámám', 'mámách',
            # Otec
            'otec', 'otce', 'otci', 'otcem', 'otcové', 'otců', 'otcům', 'otcích', 'otci',
            'tatínek', 'tatínka', 'tatínkovi', 'tatínku', 'tatínkem', 'tatínkové', 'tatínků', 'tatínkům', 'tatíncích',
            'táta', 'táty', 'tátovi', 'tátu', 'tátou', 'tátové', 'tátů', 'tátům', 'tátech',
            # Rodička - rozšíření (rodič už existuje)
            'rodičku', 'rodičkou', 'rodiček', 'rodičkám', 'rodičkách', 'rodičům', 'rodičích',
            # Bratr a sestra
            'bratr', 'bratra', 'bratru', 'bratrovi', 'bratrem', 'bratři', 'bratrů', 'bratrům', 'bratrech',
            'nevlastní bratr', 'nevlastního bratra', 'nevlastnímu bratrovi',
            'adoptivní bratr', 'adoptivního bratra', 'adoptivnímu bratrovi',
            'sestra', 'sestry', 'sestře', 'sestrou', 'sester', 'sestrám', 'sestrách',
            'nevlastní sestra', 'nevlastní sestry', 'nevlastní sestře', 'nevlastní sestrou',
            'adoptivní sestra', 'adoptivní sestry', 'adoptivní sestře', 'adoptivní sestrou',
            # Syn a dcera
            'syn', 'syna', 'synovi', 'synu', 'synem', 'synové', 'synů', 'synům', 'synech',
            'dcera', 'dcery', 'dceři', 'dcerou', 'dcer', 'dcerám', 'dcerách',
            # Vnuci
            'vnuk', 'vnuka', 'vnukovi', 'vnuku', 'vnukem', 'vnuci', 'vnuků', 'vnukům', 'vnucích',
            'vnučka', 'vnučky', 'vnučce', 'vnučkou', 'vnuček', 'vnučkám', 'vnučkách',
            # Prarodiče
            'dědeček', 'dědečka', 'dědečkovi', 'dědečku', 'dědečkem', 'dědečkové', 'dědečků', 'dědečkům', 'dědečcích',
            'děda', 'dědy', 'dědovi', 'dědu', 'dědou', 'dědové', 'dědů', 'dědům', 'dědech',
            'babička', 'babičky', 'babičce', 'babičkou', 'babiček', 'babičkám', 'babičkách',
            'bába', 'báby', 'bábě', 'bábou', 'báb', 'bábám', 'bábách',
            # Strýc, teta, bratranci
            'strýc', 'strýce', 'strýci', 'strýcem', 'strýcové', 'strýců', 'strýcům', 'strýcích',
            'teta', 'tety', 'tetě', 'tetou', 'tet', 'tetám', 'tetách',
            'bratranec', 'bratrance', 'bratranci', 'bratrancem', 'bratranci', 'bratranců', 'bratrancům', 'bratrancích',
            'sestřenice', 'sestřenici', 'sestřenicí', 'sestřenic', 'sestřenicím', 'sestřenicích',
            # Tchán, tchýně, švagr, švagrová
            'švagr', 'švagra', 'švagrovi', 'švagru', 'švagram', 'švagrové', 'švagrů', 'švagrům', 'švagrech',
            'švagrová', 'švagrové', 'švagrovou', 'švagrových', 'švagrovám', 'švagrovách',
            'tchán', 'tchána', 'tchánovi', 'tchánů', 'tchánům', 'tchánech',
            'tchýně', 'tchýni', 'tchýní', 'tchyň', 'tchyním', 'tchyních',
            'zeť', 'zeti', 'zetě', 'zeťi', 'zeťů', 'zeťům', 'zeťích',
            'snacha', 'snachy', 'snače', 'snachou', 'snach', 'snachám', 'snachách',
            # Partnerské vztahy
            'manžel', 'manžela', 'manželovi', 'manželu', 'manželem', 'manželé', 'manželů', 'manželům', 'manželech',
            'manželka', 'manželky', 'manželce', 'manželkou', 'manželek', 'manželkám', 'manželkách',
            'choť', 'choť', 'choti', 'chotě', 'chotí', 'choť', 'chotí', 'chotím', 'chotích',
            'partner', 'partnera', 'partnerovi', 'partneru', 'partnerem', 'partneři', 'partnerů', 'partnerům', 'partnerech',
            'partnerka', 'partnerky', 'partnerce', 'partnerkou', 'partnerek', 'partnerkám', 'partnerkách',
            'druh', 'druha', 'druhovi', 'druhu', 'druhem', 'druhové', 'druhů', 'druhům', 'druzích',
            'družka', 'družky', 'družce', 'družkou', 'družek', 'družkám', 'družkách',
            'snoubenec', 'snoubence', 'snoubenci', 'snoubencem', 'snoubenci', 'snoubenců', 'snoubencům', 'snoubencích',
            'snoubenka', 'snoubenky', 'snoubence', 'snoubenkou', 'snoubenek', 'snoubenkám', 'snoubenkách',
            'přítele', 'příteli', 'přítelem', 'přátelé', 'přátel', 'přátelům', 'přátelích',
            'přítelkyní', 'přítelkyň', 'přítelkyním', 'přítelkyních',
            # Ex-manželé a rozvedení
            'ex-manžel', 'ex-manžela', 'ex-manželovi', 'ex-manželu', 'ex-manželem',
            'bývalý manžel', 'bývalého manžela', 'bývalému manželovi',
            'ex-manželka', 'ex-manželky', 'ex-manželce', 'ex-manželkou',
            'bývalá manželka', 'bývalé manželky', 'bývalé manželce', 'bývalou manželkou',
            'rozvedenou', 'rozvedená', 'rozvedené', 'rozvedeného',
            'rozvedenému manželovi', 'rozvedeného manžela', 'rozvedený manžel',
            'rozvedená manželka', 'rozvedené manželky', 'rozvedené manželce',
            'vdovec', 'vdovce', 'vdovci', 'vdovcem', 'vdovci', 'vdovců', 'vdovcům', 'vdovcích',
            'vdova', 'vdovy', 'vdově', 'vdovou', 'vdov', 'vdovám', 'vdovách',
            # Děti a věk
            'nezletilý', 'nezletilého', 'nezletilému', 'nezletilém', 'nezletilým', 'nezletilých', 'nezletilými',
            'nezletilá', 'nezletilé', 'nezletilou',
            'potomek', 'potomka', 'potomkovi', 'potomku', 'potomkem', 'potomci', 'potomků', 'potomkům', 'potomcích',
            'narozený', 'narozeného', 'narozenému', 'narozeném', 'narozeným', 'narozených', 'narozenými',
            'narozená', 'narozené', 'narozenou',
            'novorozenec', 'novorozence', 'novorozenci', 'novorozencem', 'novorozenci', 'novorozenců', 'novorozencům', 'novorozencích',
            'mladistvý', 'mladistvého', 'mladistvému', 'mladistvém', 'mladistvým', 'mladistvých', 'mladistvými',
            'mladistvá', 'mladistvé', 'mladistvou',
            'dospělý', 'dospělého', 'dospělému', 'dospělém', 'dospělým', 'dospělých', 'dospělými',
            'dospělá', 'dospělé', 'dospělou',
            'rozvedenou', 'rozvedená', 'rozvedené', 'rozvedeného',
            'rozvedenému', 'rozvedený', 'rozvedených', 'rozvedeným', 'ručitel', 'ručitelce', 'ručitele', 'ručitelek', 'ručitelem',
            'ručiteli', 'ručitelka', 'ručitelkou', 'ručitelky', 'ručitelů', 'samoživitel', 'samoživitele', 'samoživitelka',
            'samoživitelky', 'stavebnice', 'stavebnicí', 'stavebník', 'stavebníka', 'stavebníkem', 'stavebníkovi', 'stavebníku',
            'stavebníků', 'student', 'studenta', 'studentce', 'studentek', 'studentem', 'studenti', 'studentka',
            'studentkou', 'studentky', 'studentovi', 'studentů', 'stěžovatel', 'stěžovatele', 'stěžovatelek', 'stěžovateli',
            'stěžovatelka', 'stěžovatelky', 'stěžovatelů', 'svědek', 'svědka', 'svědkovi', 'svědkyni', 'svědkyní',
            'svědkyně', 'svědků', 'transplantovanou', 'transplantovaná', 'transplantované', 'transplantovaného', 'transplantovanému', 'transplantovaný',
            'transplantovaným', 'uchazeč', 'uchazeče', 'uchazeči', 'uchazečka', 'uchazečky', 'uchazečů', 'uprchlice',
            'uprchlicí', 'uprchlík', 'uprchlíka', 'uprchlíkem', 'uprchlíkovi', 'uprchlíku', 'uprchlíkyně', 'uprchlíků',
            'uvedenou', 'uvedená', 'uvedené', 'uvedeného', 'uvedenému', 'uvedený', 'uvedených', 'uvedeným',
            'vybranou', 'vybraná', 'vybrané', 'vybraného', 'vybranému', 'vybraný', 'vybraných', 'vybraným',
            'vyloučenou', 'vyloučená', 'vyloučené', 'vyloučeného', 'vyloučenému', 'vyloučený', 'vyloučených', 'vyloučeným',
            'vyšetřenou', 'vyšetřená', 'vyšetřené', 'vyšetřeného', 'vyšetřenému', 'vyšetřený', 'vyšetřených', 'vyšetřeným',
            'vyšetřenými', 'vyšetřovanou', 'vyšetřovaná', 'vyšetřované', 'vyšetřovaného', 'vyšetřovanému', 'vyšetřovaný', 'vyšetřovaných',
            'vyšetřovaným', 'vyšetřovanými', 'věřitel', 'věřitelce', 'věřitele', 'věřitelek', 'věřiteli', 'věřitelka',
            'věřitelky', 'věřitelů', 'zadrženou', 'zadržená', 'zadržené', 'zadrženého', 'zadrženém', 'zadrženému',
            'zadržený', 'zadržených', 'zadrženým', 'zadrženými',
            # Zdravotní/fyzický stav - ROZŠÍŘENÉ
            'zraněnou', 'zraněná', 'zraněné', 'zraněného', 'zraněném', 'zraněnému', 'zraněný', 'zraněných', 'zraněným', 'zraněnými',
            # Lehce/těžce zraněný
            'lehce zraněný', 'lehce zraněná', 'lehce zraněné', 'lehce zraněného', 'lehce zraněnému', 'lehce zraněném', 'lehce zraněným',
            'těžce zraněný', 'těžce zraněná', 'těžce zraněné', 'těžce zraněného', 'těžce zraněnému', 'těžce zraněném', 'těžce zraněným',
            # Nemocný
            'nemocný', 'nemocného', 'nemocnému', 'nemocném', 'nemocným',
            'nemocná', 'nemocné', 'nemocnou', 'nemocnými', 'nemocných',
            'těžce nemocný', 'těžce nemocná', 'těžce nemocné', 'těžce nemocného', 'těžce nemocnému', 'těžce nemocném', 'těžce nemocným',
            'chronicky nemocný', 'chronicky nemocná', 'chronicky nemocné', 'chronicky nemocného',
            'duševně nemocný', 'duševně nemocná', 'duševně nemocné', 'duševně nemocného',
            # Hospitalizovaný a další zdravotní stavy
            'hospitalizovaný', 'hospitalizovaného', 'hospitalizovanému', 'hospitalizovaném', 'hospitalizovaným', 'hospitalizovanými', 'hospitalizovaných',
            'hospitalizovaná', 'hospitalizované', 'hospitalizovanou',
            'transportovaný', 'transportovaného', 'transportovanému', 'transportovaném', 'transportovaným', 'transportovanými', 'transportovaných',
            'transportovaná', 'transportované', 'transportovanou',
            'intubovaný', 'intubovaného', 'intubovanému', 'intubovaném', 'intubovaným', 'intubovanými', 'intubovaných',
            'intubovaná', 'intubované', 'intubovanou',
            'resuscitovaný', 'resuscitovaného', 'resuscitovanému', 'resuscitovaném', 'resuscitovaným', 'resuscitovanými', 'resuscitovaných',
            'resuscitovaná', 'resuscitované', 'resuscitovanou',
            'stabilizovaný', 'stabilizovaného', 'stabilizovanému', 'stabilizovaném', 'stabilizovaným', 'stabilizovanými', 'stabilizovaných',
            'stabilizovaná', 'stabilizované', 'stabilizovanou',
            'léčený', 'léčeného', 'léčenému', 'léčeném', 'léčeným', 'léčenými', 'léčených',
            'léčená', 'léčené', 'léčenou',
            'amputovaný', 'amputovaného', 'amputovanému', 'amputovaném', 'amputovaným', 'amputovanými', 'amputovaných',
            'amputovaná', 'amputované', 'amputovanou',
            # Dárce (transplantace)
            'dárce', 'dárci', 'dárců', 'dárcem', 'dárcům', 'dárcích', 'dárkyně', 'dárkyní', 'dárkyň', 'dárkyni',
            # Specifické zdravotní stavy
            'tetraplegie', 'tetraplegik', 'tetraplegika', 'paraplegik', 'paraplegika',
            'imobilní', 'imobilního', 'imobilnímu', 'imobilním', 'imobilních',
            'komatózní', 'komatózního', 'komatóznímu', 'komatózním', 'komatózních',
            'bezvědomý', 'bezvědomého', 'bezvědomému', 'bezvědomém', 'bezvědomým', 'bezvědomých', 'bezvědomými',
            'bezvědomá', 'bezvědomé', 'bezvědomou',
            'umírající', 'umírajícího', 'umírajícímu', 'umírajícím', 'umírajících',
            # Mrtvý/zemřelý/zesnulý
            'mrtvý', 'mrtvého', 'mrtvému', 'mrtvém', 'mrtvým', 'mrtvých', 'mrtvými',
            'mrtvá', 'mrtvé', 'mrtvou',
            'zemřelý', 'zemřelého', 'zemřelému', 'zemřelém', 'zemřelým', 'zemřelých', 'zemřelými',
            'zemřelá', 'zemřelé', 'zemřelou',
            'zesnulý', 'zesnulého', 'zesnulému', 'zesnulém', 'zesnulým', 'zesnulých', 'zesnulými',
            'zesnulá', 'zesnulé', 'zesnulou',
            'deceased', 'exitus',
            'zaměstnance', 'zaměstnancem', 'zaměstnanci', 'zaměstnanců',
            'zaměstnanec', 'zaměstnankyni', 'zaměstnankyní', 'zaměstnankyně', 'zaměstnankyň', 'zaměstnavatel', 'zaměstnavatele', 'zaměstnavatelem',
            'zaměstnavateli', 'zaměstnavatelka', 'zaměstnavatelky', 'zaměstnavatelů', 'zletilou', 'zletilá', 'zletilé', 'zletilého',
            'zletilému', 'zletilý', 'účastnic', 'účastnice', 'účastnicí', 'účastník', 'účastníka', 'účastníkem',
            'účastníkovi', 'účastníku', 'účastníků', 'žadatel', 'žadatelce', 'žadatele', 'žadatelek', 'žadateli',
            'žadatelka', 'žadatelky', 'žadatelů', 'žalobce', 'žalobci', 'žalobců', 'žalobkyni', 'žalobkyní',
            'žalobkyně', 'žalovanou', 'žalovaná', 'žalované', 'žalovaného', 'žalovaném', 'žalovanému', 'žalovaný',
            'žalovaných', 'žalovaným', 'žalovanými', 'žáci', 'žák', 'žáka', 'žákem', 'žákovi',
            'žáku', 'žákyni', 'žákyní', 'žákyně', 'žákyň', 'žáků',
            # Další termíny
            'přítel', 'přítelkyně', 'kolega', 'kolegyně',
            'majitel', 'majitelka',
            'předseda', 'předsedkyně', 'člen', 'členka',
            'věznice', 'věznici', 'vězení',  # prison - IMPORTANT!
            # Ostatní
            'scrum', 'master', 'developer', 'architect', 'engineer',
            'officer', 'professional', 'certified', 'advanced',
            'management', 'legal', 'counsel', 'executive',
            # Pozdravy/oslovení
            'ahoj', 'dobrý', 'den', 'vážený', 'vážená',
            # Značky aut
            'škoda', 'octavia', 'fabia', 'superb', 'kodiaq',
            'volkswagen', 'toyota', 'ford', 'bmw', 'audi',
            # Technologie a software
            'google', 'amazon', 'microsoft', 'apple', 'facebook',
            'cloud', 'web', 'tech', 'solutions', 'data', 'digital',
            'software', 'enterprise', 'premium', 'standard',
            'analytics', 'computer', 'vision', 'protection',
            'security', 'authenticator', 'repository', 'access',
            'personal', 'hub', 'book', 'pro', 'series', 'launch',
            'team', 'development', 'react', 'splunk', 'innovate',
            'ventures', 'credo', 'mayo', 'clinic', 'met', 'london',
            'avenue', 'contractual', 'plánovaná', 'diagno',
            'cisco', 'processing', 'notářská',
            # Zdravotnictví + Léky/Produkty
            'nemocnice', 'poliklinika', 'polikliniek', 'nemocniec',
            'healthcare', 'symbicort', 'turbuhaler', 'spirometr',
            'jaeger', 'medical', 'health', 'pharma', 'pharmaceutical',
            # ===================================================================
            # OPATROVNICTVÍ A PÉČE - ROZŠÍŘENÉ
            # ===================================================================
            # Svěřenec/svěřenka (opatrovník už existuje výše)
            'svěřenec', 'svěřence', 'svěřenci', 'svěřencem', 'svěřenci', 'svěřenců', 'svěřencům', 'svěřencích',
            'svěřenka', 'svěřenky', 'svěřence', 'svěřenkou', 'svěřenek', 'svěřenkám', 'svěřenkách',
            # Osvojenec/osvojenka (osvojitel už existuje výše)
            'osvojenec', 'osvojence', 'osvojenci', 'osvojencem', 'osvojenci', 'osvojenců', 'osvojencům', 'osvojencích',
            'osvojenka', 'osvojenky', 'osvojence', 'osvojenkou', 'osvojenek', 'osvojenkám', 'osvojenkách',
            # Pečovatel/pečovatelka
            'pečovatel', 'pečovatele', 'pečovateli', 'pečovatelem', 'pečovatelé', 'pečovatelů', 'pečovatelům', 'pečovatelích',
            'pečovatelka', 'pečovatelky', 'pečovatelce', 'pečovatelkou', 'pečovatelek', 'pečovatelkám', 'pečovatelkách',
            # Ošetřující/ošetřovatel
            'ošetřující', 'ošetřujícího', 'ošetřujícímu', 'ošetřujícím', 'ošetřujících',
            'ošetřovatel', 'ošetřovatele', 'ošetřovateli', 'ošetřovatelem', 'ošetřovatelé', 'ošetřovatelů', 'ošetřovatelům', 'ošetřovatelích',
            'ošetřovatelka', 'ošetřovatelky', 'ošetřovatelce', 'ošetřovatelkou', 'ošetřovatelek', 'ošetřovatelkám', 'ošetřovatelkách',
            # Opatrovatel/opatrovatelka
            'opatrovatel', 'opatrovatele', 'opatrovateli', 'opatrovatelem', 'opatrovatelé', 'opatrovatelů', 'opatrovatelům', 'opatrovatelích',
            'opatrovatelka', 'opatrovatelky', 'opatrovatelce', 'opatrovatelkou', 'opatrovatelek', 'opatrovatelkám', 'opatrovatelkách',
            # Zákonný zástupce
            'zákonný zástupce', 'zákonného zástupce', 'zákonnému zástupci', 'zákonném zástupci', 'zákonným zástupcem',
            'zákonná zástupkyně', 'zákonné zástupkyně', 'zákonné zástupkyni', 'zákonnou zástupkyní',
            # Kontaktní osoba
            'kontaktní osoba', 'kontaktní osoby', 'kontaktní osobě', 'kontaktní osobou',
            'kontaktní', 'kontaktního', 'kontaktnímu', 'kontaktním', 'kontaktních',
            # ===================================================================
            # ZAMĚSTNÁNÍ A FUNKCE - ROZŠÍŘENÉ
            # ===================================================================
            # Pracovník/pracovnice
            'pracovník', 'pracovníka', 'pracovníkovi', 'pracovníku', 'pracovníkem', 'pracovníci', 'pracovníků', 'pracovníkům', 'pracovnících',
            'pracovnice', 'pracovnici', 'pracovnicí', 'pracovnic', 'pracovnicím', 'pracovnicích',
            # Dělník/dělnice
            'dělník', 'dělníka', 'dělníkovi', 'dělníku', 'dělníkem', 'dělníci', 'dělníků', 'dělníkům', 'dělnících',
            'dělnice', 'dělnici', 'dělnicí', 'dělnic', 'dělnicím', 'dělnicích',
            # Úředník/úřednice
            'úředník', 'úředníka', 'úředníkovi', 'úředníku', 'úředníkem', 'úředníci', 'úředníků', 'úředníkům', 'úřednících',
            'úřednice', 'úřednici', 'úřednicí', 'úřednic', 'úřednicím', 'úřednicích',
            # Ředitel/ředitelka (rozšíření)
            'ředitel', 'ředitele', 'řediteli', 'ředitelem', 'ředitelé', 'ředitelů', 'ředitelům', 'ředitelích',
            'ředitelka', 'ředitelky', 'ředitelce', 'ředitelkou', 'ředitelek', 'ředitelkám', 'ředitelkách',
            'ředitel školy', 'ředitele školy', 'řediteli školy', 'ředitelem školy',
            'ředitelka školy', 'ředitelky školy', 'ředitelce školy', 'ředitelkou školy',
            # Jednatel/jednatelka
            'jednatel', 'jednatele', 'jednateli', 'jednatelem', 'jednatelé', 'jednatelů', 'jednatelům', 'jednatelích',
            'jednatelka', 'jednatelky', 'jednatelce', 'jednatelkou', 'jednatelek', 'jednatelkám', 'jednatelkách',
            'jednatel společnosti',
            # Předseda rozšíření (základní forma už existuje)
            'předsedy', 'předsedovi', 'předsedou', 'předsedů', 'předsedům', 'předsedech',
            'předsedkyni', 'předsedkyní', 'předsedkyň', 'předsedkyním', 'předsedkyních',
            'předseda senátu', 'předsedy senátu', 'předsedovi senátu', 'předsedou senátu',
            # Místopředseda/místopředsedkyně
            'místopředseda', 'místopředsedy', 'místopředsedovi', 'místopředsedou', 'místopředsedové', 'místopředsedů', 'místopředsedům', 'místopředsedech',
            'místopředsedkyně', 'místopředsedkyni', 'místopředsedkyní', 'místopředsedkyň', 'místopředsedkyním', 'místopředsedkyních',
            # Manažer/manažerka
            'manažer', 'manažera', 'manažerovi', 'manažeru', 'manažerem', 'manažeři', 'manažerů', 'manažerům', 'manažerech',
            'manažerka', 'manažerky', 'manažerce', 'manažerkou', 'manažerek', 'manažerkám', 'manažerkách',
            # Vedoucí
            'vedoucí', 'vedoucího', 'vedoucímu', 'vedoucím', 'vedoucích',
            # Vlastník/vlastnice (majitel už existuje)
            'vlastník', 'vlastníka', 'vlastníkovi', 'vlastníku', 'vlastníkem', 'vlastníci', 'vlastníků', 'vlastníkům', 'vlastnících',
            'vlastnice', 'vlastnici', 'vlastnicí', 'vlastnic', 'vlastnicím', 'vlastnicích',
            'vlastník nemovitosti', 'vlastníka nemovitosti', 'vlastníkovi nemovitosti', 'vlastníkem nemovitosti',
            # Spoluvlastník/spoluvlastnice
            'spoluvlastník', 'spoluvlastníka', 'spoluvlastníkovi', 'spoluvlastníku', 'spoluvlastníkem', 'spoluvlastníci', 'spoluvlastníků', 'spoluvlastníkům', 'spoluvlastnících',
            'spoluvlastnice', 'spoluvlastnici', 'spoluvlastnicí', 'spoluvlastnic', 'spoluvlastnicím', 'spoluvlastnicích',
            'spoluvlastník bytu', 'spoluvlastníka bytu', 'spoluvlastníkovi bytu', 'spoluvlastníkem bytu',
            'podílový spoluvlastník', 'podílového spoluvlastníka', 'podílovému spoluvlastníkovi', 'podílovým spoluvlastníkem',
            # Akcionář/akcionářka
            'akcionář', 'akcionáře', 'akcionáři', 'akcionářem', 'akcionáři', 'akcionářů', 'akcionářům', 'akcionářích',
            'akcionářka', 'akcionářky', 'akcionářce', 'akcionářkou', 'akcionářek', 'akcionářkám', 'akcionářkách',
            # Společník/společnice
            'společník', 'společníka', 'společníkovi', 'společníku', 'společníkem', 'společníci', 'společníků', 'společníkům', 'společnících',
            'společnice', 'společnici', 'společnicí', 'společnic', 'společnicím', 'společnicích',
            # Statutární zástupce
            'statutární zástupce', 'statutárního zástupce', 'statutárnímu zástupci', 'statutárním zástupcem',
            # Prokurist/prokuristka
            'prokurist', 'prokurista', 'prokuristovi', 'prokuristu', 'prokuristem', 'prokuristé', 'prokuristů', 'prokuristům', 'prokuristech',
            'prokuristka', 'prokuristky', 'prokuristce', 'prokuristkou', 'prokuristek', 'prokuristkám', 'prokuristkách',
            # Zmocněnec/zmocněnkyně
            'zmocněnec', 'zmocněnce', 'zmocněnci', 'zmocněncem', 'zmocněnci', 'zmocněnců', 'zmocněncům', 'zmocněncích',
            'zmocněnkyně', 'zmocněnkyni', 'zmocněnkyní', 'zmocněnkyň', 'zmocněnkyním', 'zmocněnkyních',
            # ===================================================================
            # PRÁVNÍ OZNAČENÍ - VELMI ROZŠÍŘENÉ
            # ===================================================================
            # Dotčený/zúčastněný
            'dotčený', 'dotčeného', 'dotčenému', 'dotčeném', 'dotčeným', 'dotčených', 'dotčenými',
            'dotčená', 'dotčené', 'dotčenou',
            'zúčastněný', 'zúčastněného', 'zúčastněnému', 'zúčastněném', 'zúčastněným', 'zúčastněných', 'zúčastněnými',
            'zúčastněná', 'zúčastněné', 'zúčastněnou',
            # Zastupující/zástupce rozšíření
            'zastupující', 'zastupujícího', 'zastupujícímu', 'zastupujícím', 'zastupujících',
            'zástupce', 'zástupci', 'zástupcem', 'zástupců', 'zástupcům', 'zástupcích',
            'zástupkyně', 'zástupkyni', 'zástupkyní', 'zástupkyň', 'zástupkyním', 'zástupkyních',
            'zástupce ředitele', 'zástupci ředitele', 'zástupcem ředitele',
            # Advokát/advokátka
            'advokát', 'advokáta', 'advokátovi', 'advokátu', 'advokátem', 'advokáti', 'advokátů', 'advokátům', 'advokátech',
            'advokátka', 'advokátky', 'advokátce', 'advokátkou', 'advokátek', 'advokátkám', 'advokátkách',
            # Obhájce/obhájkyně
            'obhájce', 'obhájci', 'obhájcem', 'obhájců', 'obhájcům', 'obhájcích',
            'obhájkyně', 'obhájkyni', 'obhájkyní', 'obhájkyň', 'obhájkyním', 'obhájkyních',
            # Právní zástupce
            'právní zástupce', 'právního zástupce', 'právnímu zástupci', 'právním zástupcem',
            # Znalec/znalkyně
            'znalec', 'znalce', 'znalci', 'znalcem', 'znalci', 'znalců', 'znalcům', 'znalcích',
            'znalkyně', 'znalkyni', 'znalkyní', 'znalkyň', 'znalkyním', 'znalkyních',
            'soudní znalec', 'soudního znalce', 'soudnímu znalci', 'soudním znalcem',
            # Expert/expertka, odborník/odbornice
            'expert', 'experta', 'expertovi', 'expertu', 'expertem', 'experti', 'expertů', 'expertům', 'expertech',
            'expertka', 'expertky', 'expertce', 'expertkou', 'expertek', 'expertkám', 'expertkách',
            'odborník', 'odborníka', 'odborníkovi', 'odborníku', 'odborníkem', 'odborníci', 'odborníků', 'odborníkům', 'odbornicích',
            'odbornice', 'odbornici', 'odbornicí', 'odbornic', 'odbornicím', 'odbornicích',
            # Soudce/soudkyně
            'soudce', 'soudci', 'soudcem', 'soudců', 'soudcům', 'soudcích',
            'soudkyně', 'soudkyni', 'soudkyní', 'soudkyň', 'soudkyním', 'soudkyních',
            'samosoudce', 'samosoudci', 'samosoudcem', 'samosoudců', 'samosoudcům', 'samosoudcích',
            'samosoudkyně', 'samosoudkyni', 'samosoudkyní', 'samosoudkyň', 'samosoudkyním', 'samosoudkyních',
            # Státní zástupce/prokurátorka
            'státní zástupce', 'státního zástupce', 'státnímu zástupci', 'státním zástupcem',
            'státní zástupkyně', 'státní zástupkyni', 'státní zástupkyní',
            'prokurátor', 'prokurátora', 'prokurátorovi', 'prokurátoři', 'prokurátorů', 'prokurátorům', 'prokurátorech',
            'prokurátorka', 'prokurátorky', 'prokurátorce', 'prokurátorkou', 'prokurátorek', 'prokurátorkám', 'prokurátorkách',
            # Insolvenční správce
            'insolvenční správce', 'insolvenčního správce', 'insolvenčnímu správci', 'insolvenčním správcem',
            'insolvenční správkyně', 'insolvenční správkyni', 'insolvenční správkyní',
            'správce konkursní podstaty', 'správce konkursní podstaty', 'správci konkursní podstaty', 'správcem konkursní podstaty',
            # Probační úředník, kurátor
            'probační úředník', 'probačního úředníka', 'probačnímu úředníkovi', 'probačním úředníkem',
            'probační úřednice', 'probační úřednici', 'probační úřednicí',
            'kurátor', 'kurátora', 'kurátorovi', 'kurátoři', 'kurátorů', 'kurátorům', 'kurátorech',
            'kurátorka', 'kurátorky', 'kurátorce', 'kurátorkou', 'kurátorek', 'kurátorkám', 'kurátorkách',
            'sociální kurátor', 'sociálního kurátora', 'sociálnímu kurátorovi', 'sociálním kurátorem',
            # ===================================================================
            # NEMOVITOSTI
            # ===================================================================
            # Nájemce/nájemkyně
            'nájemce', 'nájemci', 'nájemcem', 'nájemců', 'nájemcům', 'nájemcích',
            'nájemkyně', 'nájemkyni', 'nájemkyní', 'nájemkyň', 'nájemkyním', 'nájemkyních',
            'nájemce bytu', 'nájemci bytu', 'nájemcem bytu',
            # Podnájemce/podnájemkyně
            'podnájemce', 'podnájemci', 'podnájemcem', 'podnájemců', 'podnájemcům', 'podnájemcích',
            'podnájemkyně', 'podnájemkyni', 'podnájemkyní', 'podnájemkyň', 'podnájemkyním', 'podnájemkyních',
            # Pronajímatel/pronajímatelka
            'pronajímatel', 'pronajímatele', 'pronajímateli', 'pronajímatelem', 'pronajímatelé', 'pronajímatelů', 'pronajímatelům', 'pronajímatelích',
            'pronajímatelka', 'pronajímatelky', 'pronajímatelce', 'pronajímatelkou', 'pronajímatelek', 'pronajímatelkám', 'pronajímatelkách',
            # Uživatel/uživatelka
            'uživatel', 'uživatele', 'uživateli', 'uživatelem', 'uživatelé', 'uživatelů', 'uživatelům', 'uživatelích',
            'uživatelka', 'uživatelky', 'uživatelce', 'uživatelkou', 'uživatelek', 'uživatelkám', 'uživatelkách',
            'oprávněný uživatel', 'oprávněného uživatele', 'oprávněnému uživateli', 'oprávněným uživatelem',
            'neoprávněný uživatel', 'neoprávněného uživatele', 'neoprávněnému uživateli', 'neoprávněným uživatelem',
            # Investor stavby
            'investor stavby', 'investora stavby', 'investorovi stavby', 'investorem stavby',
            # ===================================================================
            # ŠKOLSTVÍ - ROZŠÍŘENÉ
            # ===================================================================
            # Absolvent/absolventka
            'absolvent', 'absolventa', 'absolventovi', 'absolventu', 'absolventem', 'absolventi', 'absolventů', 'absolventům', 'absolventech',
            'absolventka', 'absolventky', 'absolventce', 'absolventkou', 'absolventek', 'absolventkám', 'absolventkách',
            # Učitel/učitelka
            'učitel', 'učitele', 'učiteli', 'učitelem', 'učitelé', 'učitelů', 'učitelům', 'učitelích',
            'učitelka', 'učitelky', 'učitelce', 'učitelkou', 'učitelek', 'učitelkám', 'učitelkách',
            'třídní učitel', 'třídního učitele', 'třídnímu učiteli', 'třídním učitelem',
            'třídní učitelka', 'třídní učitelky', 'třídní učitelce', 'třídní učitelkou',
            # Pedagog/pedagožka
            'pedagog', 'pedagoga', 'pedagogovi', 'pedagogu', 'pedagogem', 'pedagogové', 'pedagogů', 'pedagogům', 'pedagogech',
            'pedagožka', 'pedagožky', 'pedagožce', 'pedagožkou', 'pedagožek', 'pedagožkám', 'pedagožkách',
            # Vyučující
            'vyučující', 'vyučujícího', 'vyučujícímu', 'vyučujícím', 'vyučujících',
            # Lektor/lektorka
            'lektor', 'lektora', 'lektorovi', 'lektoru', 'lektorem', 'lektoři', 'lektorů', 'lektorům', 'lektorech',
            'lektorka', 'lektorky', 'lektorce', 'lektorkou', 'lektorek', 'lektorkám', 'lektorkách',
            # ===================================================================
            # SPECIFICKÉ SITUACE
            # ===================================================================
            # Osoba identifikovaná, fyzická osoba
            'osoba identifikovaná jako', 'osoby identifikované jako', 'osobě identifikované jako',
            'osoba jménem', 'osoby jménem', 'osobě jménem', 'osobou jménem',
            'fyzická osoba', 'fyzické osoby', 'fyzické osobě', 'fyzickou osobou',
            # Občan (rozšíření - základní formy už existují)
            'občan', 'občana', 'občanovi', 'občanu', 'občanem', 'občané', 'občanů', 'občanům', 'občanech',
            'občanka', 'občanky', 'občance',
            # Přítomný/nepřítomný
            'přítomný', 'přítomného', 'přítomněm u', 'přítomném', 'přítomným', 'přítomných', 'přítomným i',
            'přítomná', 'přítomné', 'přítomnou',
            'nepřítomný', 'nepřítomného', 'nepřítomněm u', 'nepřítomném', 'nepřítomným', 'nepřítomných', 'nepřítomným i',
            'nepřítomná', 'nepřítomné', 'nepřítomnou',
            'dostavivší se', 'dostavivšího se', 'dostavivšímu se', 'dostavivším se', 'dostavivších se',
            # Jmenovaný/uvedený
            'jmenovaný', 'jmenovaného', 'jmenovanému', 'jmenovaném', 'jmenovaným', 'jmenovaných', 'jmenovanými',
            'jmenovaná', 'jmenované', 'jmenovanou',
            'jmenovaný do funkce', 'jmenovaného do funkce', 'jmenovanému do funkce', 'jmenovaným do funkce',
            'výše uvedený', 'výše uvedeného', 'výše uvedenému', 'výše uvedeném', 'výše uvedeným',
            'výše uvedená', 'výše uvedené', 'výše uvedenou',
            'níže uvedený', 'níže uvedeného', 'níže uvedenému', 'níže uvedeném', 'níže uvedeným',
            'níže uvedená', 'níže uvedené', 'níže uvedenou',
            'shora uvedený', 'shora uvedeného', 'shora uvedenému', 'shora uvedeném', 'shora uvedeným',
            'shora uvedená', 'shora uvedené', 'shora uvedenou',
            # Zvolený/pověřený/delegovaný
            'zvolený', 'zvoleného', 'zvolenému', 'zvoleném', 'zvoleným', 'zvolených', 'zvoleným i',
            'zvolená', 'zvolené', 'zvolenou',
            'pověřený', 'pověřeného', 'pověřenému', 'pověřeném', 'pověřeným', 'pověřených', 'pověřenými',
            'pověřená', 'pověřené', 'pověřenou',
            'delegovaný', 'delegovaného', 'delegovanému', 'delegovaném', 'delegovaným', 'delegovaných', 'delegovanými',
            'delegovaná', 'delegované', 'delegovanou',
            # Zbavený/odvolaný (vyloučený už existuje)
            'zbavený', 'zbaveného', 'zbavenému', 'zbaveném', 'zbaveným', 'zbavených', 'zbavenými',
            'zbavená', 'zbavené', 'zbavenou',
            'odvolaný', 'odvolaného', 'odvolanému', 'odvolaném', 'odvolaným', 'odvolaných', 'odvolanými',
            'odvolaná', 'odvolané', 'odvolanou',
            # Azylant/emigrant/imigrant (uprchlík už existuje)
            'azylant', 'azylanta', 'azylantovi', 'azylantu', 'azylantem', 'azylanti', 'azylantů', 'azylantům', 'azylantech',
            'azylantka', 'azylantky', 'azylantce', 'azylantkou', 'azylantek', 'azylantkám', 'azylantkách',
            'emigrant', 'emigranta', 'emigrantovi', 'emigrantu', 'emigrantem', 'emigranti', 'emigrantů', 'emigrantům', 'emigrantech',
            'emigrantka', 'emigrantky', 'emigrantce', 'emigrantkou', 'emigrantek', 'emigrantkám', 'emigrantkách',
            'imigrant', 'imigranta', 'imigrantovi', 'imigrantu', 'imigrantem', 'imigranti', 'imigrantů', 'imigrantům', 'imigrantech',
            'imigrantka', 'imigrantky', 'imigrantce', 'imigrantkou', 'imigrantek', 'imigrantkám', 'imigrantkách',
            'žadatel o azyl', 'žadatele o azyl', 'žadateli o azyl', 'žadatelem o azyl',
            'cizinec', 'cizince', 'cizinci', 'cizincem', 'cizinci', 'cizinců', 'cizincům', 'cizincích',
            'cizinka', 'cizinky', 'cizince', 'cizinkou', 'cizinek', 'cizinkám', 'cizinkách',
            # Senior/důchodce (rozšíření)
            'senior', 'seniora', 'seniorovi', 'senioru', 'seniorem', 'senioři', 'seniorů', 'seniorům', 'seniorech',
            'seniorka', 'seniorky', 'seniorce', 'seniorkou', 'seniorek', 'seniorkám', 'seniorkách',
            'důchodce', 'důchodci', 'důchodcem', 'důchodců', 'důchodcům', 'důchodcích',
            'důchodkyně', 'důchodkyni', 'důchodkyní', 'důchodkyň', 'důchodkyním', 'důchodkyních',
            # Těhotná/gravidní/rodička
            'těhotná', 'těhotné', 'těhotnou', 'těhotných', 'těhotným', 'těhotnými',
            'gravidní', 'gravidní', 'gravidních', 'gravidním', 'gravidními',
            'porodnice', 'porodnici', 'porodnicí', 'porodnic', 'porodnicím', 'porodnicích',
            # Pozůstalý/dědic
            'pozůstalý', 'pozůstalého', 'pozůstalému', 'pozůstalém', 'pozůstalým', 'pozůstalých', 'pozůstalými',
            'pozůstalá', 'pozůstalé', 'pozůstalou',
            'dědic', 'dědice', 'dědici', 'dědicem', 'dědici', 'dědiců', 'dědicům', 'dědicích',
            'dědička', 'dědičky', 'dědičce', 'dědičkou', 'dědiček', 'dědičkám', 'dědičkách',
            'odkázaný', 'odkázaného', 'odkázanému', 'odkázaném', 'odkázaným', 'odkázaných', 'odkázanými',
            'odkázaná', 'odkázané', 'odkázanou',
            # Bývalý/původní/nový
            'bývalý', 'bývalého', 'bývalému', 'bývalém', 'bývalým', 'bývalých', 'bývalými',
            'bývalá', 'bývalé', 'bývalou',
            'původní', 'původního', 'původnímu', 'původním', 'původních',
            'předchozí', 'předchozího', 'předchozímu', 'předchozím', 'předchozích',
            'nový', 'nového', 'novému', 'novém', 'novým', 'nových', 'novými',
            'nová', 'nové', 'novou',
            'nynější', 'nynějšího', 'nynějšímu', 'nynějším', 'nynějších',
            'současný', 'současného', 'současnému', 'současném', 'současným', 'současných', 'současnými',
            'současná', 'současné', 'současnou',
            # ===================================================================
            # KOMBINOVANÉ TERMÍNY
            # ===================================================================
            # Manželé/rodiče/sourozenci (plurály)
            'manželé', 'manželů', 'manželům', 'manželích', 'manželi',
            # Pan/paní (rozšíření)
            'pan', 'pana', 'panovi', 'panu', 'panem', 'pánové', 'pánů', 'pánům', 'pánech',
            'paní', 'paní', 'paním', 'paních',
            'pán', 'pána', 'pánovi', 'pánu', 'pánom',
            'pani', 'pani', 'paním', 'paních',
            # S tituly
            'pan doktor', 'pana doktora', 'panu doktorovi', 'panem doktorem',
            'paní doktorka', 'paní doktorky', 'paní doktorce', 'paní doktorkou',
            'pan profesor', 'pana profesora', 'panu profesorovi', 'panem profesorem',
            'paní profesorka', 'paní profesorky', 'paní profesorce', 'paní profesorkou',
            'pan inženýr', 'pana inženýra', 'panu inženýrovi', 'panem inženýrem',
            'paní inženýrka', 'paní inženýrky', 'paní inženýrce', 'paní inženýrkou',
            # Další
            'care', 'plus', 'minus', 'service', 'services',
            'group', 'company', 'corp', 'ltd', 'gmbh', 'inc'
        }

        # ========== Handler pro "Titul Jméno Příjmení" (3 slova) ==========
        def replace_role_person(match):
            """
            Zpracuje pattern "Titul Jméno Příjmení" (např. "Klient Ladislav Konečný").
            Pokud první slovo je titul/role v ignore_words, anonymizuje 2. a 3. slovo jako osobu.
            """
            role_word = match.group(1)
            first_obs = match.group(2)
            last_obs = match.group(3)

            # Zkontroluj, jestli první slovo je titul/role
            if role_word.lower() not in ignore_words:
                # První slovo NENÍ titul → vrať original (nechť to zpracuje běžný 2-slovný pattern)
                return match.group(0)

            # První slovo JE titul → anonymizuj jméno a příjmení
            # Použij stejnou logiku jako u běžných osob

            # Infer nominative
            last_nom = infer_surname_nominative(last_obs)
            first_nom = infer_first_name_nominative(first_obs) or first_obs

            # Create/find person tag
            tag, canonical = self._ensure_person_tag(first_nom, last_nom)

            # Save variant if different from canonical
            original_form = f"{first_obs} {last_obs}"
            if original_form.lower() != canonical.lower():
                self.entity_map['PERSON'][canonical].add(original_form)

            # Return: "Titul [[PERSON_X]]"
            return f"{role_word} {tag}"

        def replace_person(match):
            first_obs = match.group(1)
            last_obs = match.group(2)

            # ========== A) BLACKLIST NE-OSOB ==========

            # 1. Blacklist kritických slov (firmy, instituce, produkty, role)
            critical_blacklist = {
                # Firmy a právní formy
                's.r.o.', 'a.s.', 'spol.', 'k.s.', 'v.o.s.', 'o.p.s.',
                'ltd', 'inc', 'corp', 'gmbh', 'llc',
                # Instituce
                'czech', 'republic', 'synlab', 'gymnázium', 'gymnasium',
                'university', 'univerzita', 'fakulta', 'klinika', 'nemocnice',
                'centrum', 'ústav', 'institute', 'academy', 'akademie',
                'motol', 'bulovka', 'thomayer', 'center',
                # Produkty/Software
                'kaspersky', 'endpoint', 'latitude', 'archer', 'classic',
                'windows', 'linux', 'android', 'ios', 'office', 'excel',
                # Role/Pozice (když jsou samostatně)
                'ředitelka', 'ředitel', 'jednatel', 'jednatelka',
                'manager', 'director', 'chief', 'officer',
                'vyšetřující', 'vyšetřovatel', 'lékař', 'doktor', 'sestra'
            }

            # Kontrola, zda hodnota obsahuje blacklist slovo
            combined = f"{first_obs} {last_obs}".lower()
            for word in critical_blacklist:
                if word in combined:
                    return match.group(0)  # Není osoba

            # 2. Rozšířený ignore list (původní)
            ignore_words = {
                # Běžná slova ve smlouvách
                'místo', 'datum', 'částku', 'bytem', 'sídlo', 'adresa',
                'číslo', 'kontakt', 'telefon', 'email', 'rodné', 'narozena',
                'vydán', 'uzavřena', 'podepsána', 'smlouva', 'dohoda',
                # Místa
                'staré', 'město', 'nové', 'město', 'malá', 'strana',
                'václavské', 'náměstí', 'hlavní', 'nádraží',
                'karlovy', 'vary', 'karlova', 'var',  # Karlovy Vary city
                'hradec', 'hradci', 'králové',  # Hradec Králové city
                # Organizace/instituce klíčová slova
                'česká', 'spořitelna', 'komerční', 'banka', 'raiffeisen',
                'credit', 'bank', 'financial', 'global', 'senior',
                'junior', 'lead', 'chief', 'head', 'director',
                # Finance/Investment
                'capital', 'equity', 'value', 'crescendo', 'investment',
                'fund', 'holdings', 'partners', 'assets', 'portfolio',
                # Pozice/role
                'jednatel', 'jednatelka', 'ředitel', 'ředitelka',
                'auditor', 'manager', 'consultant', 'specialist',
                'assistant', 'coordinator', 'analyst',
                # CRITICAL: Tituly a role před jmény ve smlouvách
                'pan', 'paní', 'pán', 'pani',
                'pacient', 'pacientka', 'pacientek',
                'obžalovaný', 'obžalovaná', 'obžalované', 'obžalovaného',
                'zaměstnanec', 'zaměstnankyně', 'zaměstnance',
                'kupující', 'prodávající', 'prodávajícího',
                'stavebník', 'stavebníka',
                'investor', 'investora',
                'dlužník', 'dlužníka', 'věřitel', 'věřitele',
                'odsouzený', 'odsouzená', 'odsouzeného',
                # === KOMPLETNÍ SEZNAM PRÁVNÍCH, ZDRAVOTNICKÝCH A ADMINISTRATIVNÍCH TERMÍNŮ (605 termínů) ===
                # Právní, ekonomické, zdravotnické, pracovněprávní, rodinné, sociální, vzdělávací a ostatní termíny
                # které se mohou vyskytovat před jmény osob a nesmí být detekovány jako křestní jména
                'absolvent', 'absolventa', 'absolventce', 'absolventek', 'absolventem', 'absolventka', 'absolventkou', 'absolventky',
                'absolventovi', 'absolventů', 'azylant', 'azylanta', 'azylantce', 'azylantek', 'azylantem', 'azylantka',
                'azylantkou', 'azylantky', 'azylantovi', 'azylantu', 'azylantů', 'beneficient', 'beneficienta', 'beneficientem',
                'beneficientka', 'beneficientky', 'beneficientovi', 'beneficientů', 'cizince', 'cizincem', 'cizinci', 'cizinců',
                'cizinec', 'cizinek', 'cizinka', 'cizinkou', 'cizinky', 'dlužnic', 'dlužnice', 'dlužnicí',
                'dlužník', 'dlužníka', 'dlužníkem', 'dlužníkovi', 'dlužníku', 'dlužníky', 'dlužníků', 'dotčenou',
                'dotčená', 'dotčené', 'dotčeného', 'dotčenému', 'dotčený', 'dotčených', 'dotčeným', 'držitel',
                'držitelce', 'držitele', 'držitelek', 'držitelem', 'držiteli', 'držitelka', 'držitelkou', 'držitelky',
                'držitelů', 'dárce', 'dárci', 'dárců', 'dárkyni', 'dárkyní', 'dárkyně', 'důchodce',
                'důchodci', 'důchodců', 'důchodkyni', 'důchodkyně', 'důchodkyň', 'evidovanou', 'evidovaná', 'evidované',
                'evidovaného', 'evidovanému', 'evidovaný', 'evidovaných', 'evidovaným', 'hospitalizovanou', 'hospitalizovaná', 'hospitalizované',
                'hospitalizovaného', 'hospitalizovanému', 'hospitalizovaný', 'hospitalizovaných', 'hospitalizovaným', 'invalidní', 'invalidních', 'invalidního',
                'invalidním', 'invalidnímu', 'investor', 'investora', 'investorce', 'investorek', 'investorem', 'investorka',
                'investorkou', 'investorky', 'investorovi', 'investorů', 'investoři', 'jednotlivce', 'jednotlivcem', 'jednotlivci',
                'jednotlivců', 'jednotlivec', 'jmenovanou', 'jmenovaná', 'jmenované', 'jmenovaného', 'jmenovanému', 'jmenovaný',
                'jmenovaných', 'jmenovaným', 'klient', 'klienta', 'klientce', 'klientek', 'klientem', 'klientka',
                'klientkou', 'klientky', 'klientovi', 'klientu', 'klientů', 'kontrolovanou', 'kontrolovaná', 'kontrolované',
                'kontrolovaného', 'kontrolovanému', 'kontrolovaný', 'kontrolovaných', 'kontrolovaným', 'kupující', 'kupujících', 'kupujícího',
                'kupujícím', 'kupujícímu', 'léčenou', 'léčená', 'léčené', 'léčeného', 'léčenému', 'léčený',
                'léčených', 'léčeným', 'léčenými', 'manžel', 'manžela', 'manželce', 'manželek', 'manželem',
                'manželka', 'manželkou', 'manželky', 'manželovi', 'manželé', 'manželů', 'nemocnou', 'nemocná',
                'nemocné', 'nemocného', 'nemocném', 'nemocnému', 'nemocný', 'nemocných', 'nemocným', 'nemocnými',
                'nezletilou', 'nezletilá', 'nezletilé', 'nezletilého', 'nezletilému', 'nezletilý', 'nezletilých', 'nezletilým',
                'obviněnou', 'obviněná', 'obviněné', 'obviněného', 'obviněném', 'obviněnému', 'obviněný', 'obviněných',
                'obviněným', 'obviněnými', 'občan', 'občana', 'občance', 'občanek', 'občanem', 'občanka',
                'občankou', 'občanky', 'občanovi', 'občanu', 'občané', 'občanů', 'obžalovanou', 'obžalovaná',
                'obžalované', 'obžalovaného', 'obžalovaném', 'obžalovanému', 'obžalovaný', 'obžalovaných', 'obžalovaným', 'obžalovanými',
                'odsouzenou', 'odsouzená', 'odsouzené', 'odsouzeného', 'odsouzeném', 'odsouzenému', 'odsouzený', 'odsouzených',
                'odsouzeným', 'odsouzenými', 'opatrovnic', 'opatrovnice', 'opatrovnicí', 'opatrovník', 'opatrovníka', 'opatrovníkem',
                'opatrovníkovi', 'opatrovníku', 'opatrovníků', 'opatrovanec', 'opatrovance', 'opatrovanci', 'opatrovancem', 'opatrovanců',
                'opatrovanka', 'opatrovanky', 'opatrovankou', 'opatrovanek',
            # Náboženské a rodinné tituly
            'matka', 'matky', 'matce', 'matkou', 'matek', 'matkám', 'matkách', 'matkama',
            # Pěstounská péče
            'pěstoun', 'pěstouna', 'pěstounovi', 'pěstounu', 'pěstounem', 'pěstouni', 'pěstounů', 'pěstounům', 'pěstounech', 'pěstouny',
            'pěstounka', 'pěstounky', 'pěstounce', 'pěstounkou', 'pěstounek', 'pěstounkám', 'pěstounkách', 'pěstounkami',
            # Exekutoři / bailiffs
            'exekutor', 'exekutora', 'exekutorovi', 'exekutoru', 'exekutorem', 'exekutoři', 'exekutorů', 'exekutorům', 'exekutorech', 'exekutory',
            'exekutorka', 'exekutorky', 'exekutorce', 'exekutorkou', 'exekutorek', 'exekutorkám', 'exekutorkách', 'exekutorkami',
            'operovanou', 'operovaná', 'operované', 'operovaného', 'operovanému',
                'operovaný', 'operovaných', 'operovaným', 'operovanými', 'oprávněnou', 'oprávněná', 'oprávněné', 'oprávněného',
                'oprávněnému', 'oprávněný', 'oprávněných', 'oprávněným', 'osob', 'osoba', 'osobami', 'osobou',
                'osoby', 'osobám', 'osobě', 'osvojitel', 'osvojitelce', 'osvojitele', 'osvojitelem', 'osvojiteli',
                'osvojitelka', 'osvojitelky', 'osvojitelů', 'ošetřenou', 'ošetřená', 'ošetřené', 'ošetřeného', 'ošetřenému',
                'ošetřený', 'ošetřených', 'ošetřeným', 'ošetřenými', 'pachatel', 'pachatelce', 'pachatele', 'pachatelek',
                'pachateli', 'pachatelka', 'pachatelky', 'pachatelů', 'pacient', 'pacienta', 'pacientce', 'pacientek',
                'pacientem', 'pacientka', 'pacientkou', 'pacientky', 'pacientovi', 'pacientu', 'pacientů', 'pečující',
                'pečujících', 'pečujícího', 'pečujícím', 'pečujícímu', 'plátce', 'plátci', 'plátců', 'plátkyní',
                'plátkyně', 'podezřelou', 'podezřelá', 'podezřelé', 'podezřelého', 'podezřelém', 'podezřelému', 'podezřelý',
                'podezřelých', 'podezřelým', 'podezřelými', 'podnikatel', 'podnikatelce', 'podnikatele', 'podnikatelek', 'podnikateli',
                'podnikatelka', 'podnikatelky', 'podnikatelů', 'pojištěnce', 'pojištěncem', 'pojištěnci', 'pojištěnců', 'pojištěnec',
                'pojištěnek', 'pojištěnka', 'pojištěnkou', 'pojištěnky', 'poručnic', 'poručnice', 'poručnicí', 'poručník',
                'poručníka', 'poručníkem', 'poručníkovi', 'poručníku', 'poručníků', 'poškozenou', 'poškozená', 'poškozené',
                'poškozeného', 'poškozeném', 'poškozenému', 'poškozený', 'poškozených', 'poškozeným', 'poškozenými', 'prodávající',
                'prodávajících', 'prodávajícího', 'prodávajícím', 'prodávajícímu', 'propuštěnou', 'propuštěná', 'propuštěné', 'propuštěného',
                'propuštěném', 'propuštěnému', 'propuštěný', 'propuštěných', 'propuštěným', 'propuštěnými', 'přihlášenou', 'přihlášená',
                'přihlášené', 'přihlášeného', 'přihlášenému', 'přihlášený', 'přihlášených', 'přihlášeným', 'přijatou', 'přijatá',
                'přijatého', 'přijatému', 'přijatévé', 'přijatý', 'přijatých', 'přijatým', 'příjemce', 'příjemci',
                'příjemců', 'příjemkyní', 'příjemkyně', 'registrovanou', 'registrovaná', 'registrované', 'registrovaného', 'registrovanému',
                'registrovaný', 'registrovaných', 'registrovaným', 'rodič', 'rodiče', 'rodičem', 'rodiči', 'rodička',
                'rodičky', 'rodičů', 'rodina', 'rodině', 'rodiny',
                # Rodinné vztahy - VELMI ROZŠÍŘENÉ
                'dítě', 'dítěte', 'dítěti', 'dítětem', 'děti', 'dětí', 'dětem', 'dětmi',
                'sourozenec', 'sourozence', 'sourozenci', 'sourozencem', 'sourozenci', 'sourozencům', 'sourozencích', 'sourozenci',
                'sourozenka', 'sourozence', 'sourozenkou', 'sourozenci', 'sourozenek', 'sourozenkyně',
                # Matka - varianty (základní forma už existuje výše)
                'maminka', 'maminky', 'mamince', 'maminkou', 'maminek', 'maminkám', 'maminkách',
                'máma', 'mámy', 'mámě', 'mámou', 'mám', 'mámám', 'mámách',
                # Otec
                'otec', 'otce', 'otci', 'otcem', 'otcové', 'otců', 'otcům', 'otcích', 'otci',
                'tatínek', 'tatínka', 'tatínkovi', 'tatínku', 'tatínkem', 'tatínkové', 'tatínků', 'tatínkům', 'tatíncích',
                'táta', 'táty', 'tátovi', 'tátu', 'tátou', 'tátové', 'tátů', 'tátům', 'tátech',
                # Rodička - rozšíření (rodič už existuje)
                'rodičku', 'rodičkou', 'rodiček', 'rodičkám', 'rodičkách', 'rodičům', 'rodičích',
                # Bratr a sestra
                'bratr', 'bratra', 'bratru', 'bratrovi', 'bratrem', 'bratři', 'bratrů', 'bratrům', 'bratrech',
                'nevlastní bratr', 'nevlastního bratra', 'nevlastnímu bratrovi',
                'adoptivní bratr', 'adoptivního bratra', 'adoptivnímu bratrovi',
                'sestra', 'sestry', 'sestře', 'sestrou', 'sester', 'sestrám', 'sestrách',
                'nevlastní sestra', 'nevlastní sestry', 'nevlastní sestře', 'nevlastní sestrou',
                'adoptivní sestra', 'adoptivní sestry', 'adoptivní sestře', 'adoptivní sestrou',
                # Syn a dcera
                'syn', 'syna', 'synovi', 'synu', 'synem', 'synové', 'synů', 'synům', 'synech',
                'dcera', 'dcery', 'dceři', 'dcerou', 'dcer', 'dcerám', 'dcerách',
                # Vnuci
                'vnuk', 'vnuka', 'vnukovi', 'vnuku', 'vnukem', 'vnuci', 'vnuků', 'vnukům', 'vnucích',
                'vnučka', 'vnučky', 'vnučce', 'vnučkou', 'vnuček', 'vnučkám', 'vnučkách',
                # Prarodiče
                'dědeček', 'dědečka', 'dědečkovi', 'dědečku', 'dědečkem', 'dědečkové', 'dědečků', 'dědečkům', 'dědečcích',
                'děda', 'dědy', 'dědovi', 'dědu', 'dědou', 'dědové', 'dědů', 'dědům', 'dědech',
                'babička', 'babičky', 'babičce', 'babičkou', 'babiček', 'babičkám', 'babičkách',
                'bába', 'báby', 'bábě', 'bábou', 'báb', 'bábám', 'bábách',
                # Strýc, teta, bratranci
                'strýc', 'strýce', 'strýci', 'strýcem', 'strýcové', 'strýců', 'strýcům', 'strýcích',
                'teta', 'tety', 'tetě', 'tetou', 'tet', 'tetám', 'tetách',
                'bratranec', 'bratrance', 'bratranci', 'bratrancem', 'bratranci', 'bratranců', 'bratrancům', 'bratrancích',
                'sestřenice', 'sestřenici', 'sestřenicí', 'sestřenic', 'sestřenicím', 'sestřenicích',
                # Tchán, tchýně, švagr, švagrová
                'švagr', 'švagra', 'švagrovi', 'švagru', 'švagram', 'švagrové', 'švagrů', 'švagrům', 'švagrech',
                'švagrová', 'švagrové', 'švagrovou', 'švagrových', 'švagrovám', 'švagrovách',
                'tchán', 'tchána', 'tchánovi', 'tchánů', 'tchánům', 'tchánech',
                'tchýně', 'tchýni', 'tchýní', 'tchyň', 'tchyním', 'tchyních',
                'zeť', 'zeti', 'zetě', 'zeťi', 'zeťů', 'zeťům', 'zeťích',
                'snacha', 'snachy', 'snače', 'snachou', 'snach', 'snachám', 'snachách',
                # Partnerské vztahy
                'manžel', 'manžela', 'manželovi', 'manželu', 'manželem', 'manželé', 'manželů', 'manželům', 'manželech',
                'manželka', 'manželky', 'manželce', 'manželkou', 'manželek', 'manželkám', 'manželkách',
                'choť', 'choť', 'choti', 'chotě', 'chotí', 'choť', 'chotí', 'chotím', 'chotích',
                'partner', 'partnera', 'partnerovi', 'partneru', 'partnerem', 'partneři', 'partnerů', 'partnerům', 'partnerech',
                'partnerka', 'partnerky', 'partnerce', 'partnerkou', 'partnerek', 'partnerkám', 'partnerkách',
                'druh', 'druha', 'druhovi', 'druhu', 'druhem', 'druhové', 'druhů', 'druhům', 'druzích',
                'družka', 'družky', 'družce', 'družkou', 'družek', 'družkám', 'družkách',
                'snoubenec', 'snoubence', 'snoubenci', 'snoubencem', 'snoubenci', 'snoubenců', 'snoubencům', 'snoubencích',
                'snoubenka', 'snoubenky', 'snoubence', 'snoubenkou', 'snoubenek', 'snoubenkám', 'snoubenkách',
                'přítele', 'příteli', 'přítelem', 'přátelé', 'přátel', 'přátelům', 'přátelích',
                'přítelkyní', 'přítelkyň', 'přítelkyním', 'přítelkyních',
                # Ex-manželé a rozvedení
                'ex-manžel', 'ex-manžela', 'ex-manželovi', 'ex-manželu', 'ex-manželem',
                'bývalý manžel', 'bývalého manžela', 'bývalému manželovi',
                'ex-manželka', 'ex-manželky', 'ex-manželce', 'ex-manželkou',
                'bývalá manželka', 'bývalé manželky', 'bývalé manželce', 'bývalou manželkou',
                'rozvedenou', 'rozvedená', 'rozvedené', 'rozvedeného',
                'rozvedenému manželovi', 'rozvedeného manžela', 'rozvedený manžel',
                'rozvedená manželka', 'rozvedené manželky', 'rozvedené manželce',
                'vdovec', 'vdovce', 'vdovci', 'vdovcem', 'vdovci', 'vdovců', 'vdovcům', 'vdovcích',
                'vdova', 'vdovy', 'vdově', 'vdovou', 'vdov', 'vdovám', 'vdovách',
                # Děti a věk
                'nezletilý', 'nezletilého', 'nezletilému', 'nezletilém', 'nezletilým', 'nezletilých', 'nezletilými',
                'nezletilá', 'nezletilé', 'nezletilou',
                'potomek', 'potomka', 'potomkovi', 'potomku', 'potomkem', 'potomci', 'potomků', 'potomkům', 'potomcích',
                'narozený', 'narozeného', 'narozenému', 'narozeném', 'narozeným', 'narozených', 'narozenými',
                'narozená', 'narozené', 'narozenou',
                'novorozenec', 'novorozence', 'novorozenci', 'novorozencem', 'novorozenci', 'novorozenců', 'novorozencům', 'novorozencích',
                'mladistvý', 'mladistvého', 'mladistvému', 'mladistvém', 'mladistvým', 'mladistvých', 'mladistvými',
                'mladistvá', 'mladistvé', 'mladistvou',
                'dospělý', 'dospělého', 'dospělému', 'dospělém', 'dospělým', 'dospělých', 'dospělými',
                'dospělá', 'dospělé', 'dospělou',
                'rozvedenou', 'rozvedená', 'rozvedené', 'rozvedeného',
                'rozvedenému', 'rozvedený', 'rozvedených', 'rozvedeným', 'ručitel', 'ručitelce', 'ručitele', 'ručitelek', 'ručitelem',
                'ručiteli', 'ručitelka', 'ručitelkou', 'ručitelky', 'ručitelů', 'samoživitel', 'samoživitele', 'samoživitelka',
                'samoživitelky', 'stavebnice', 'stavebnicí', 'stavebník', 'stavebníka', 'stavebníkem', 'stavebníkovi', 'stavebníku',
                'stavebníků', 'student', 'studenta', 'studentce', 'studentek', 'studentem', 'studenti', 'studentka',
                'studentkou', 'studentky', 'studentovi', 'studentů', 'stěžovatel', 'stěžovatele', 'stěžovatelek', 'stěžovateli',
                'stěžovatelka', 'stěžovatelky', 'stěžovatelů', 'svědek', 'svědka', 'svědkovi', 'svědkyni', 'svědkyní',
                'svědkyně', 'svědků', 'transplantovanou', 'transplantovaná', 'transplantované', 'transplantovaného', 'transplantovanému', 'transplantovaný',
                'transplantovaným', 'uchazeč', 'uchazeče', 'uchazeči', 'uchazečka', 'uchazečky', 'uchazečů', 'uprchlice',
                'uprchlicí', 'uprchlík', 'uprchlíka', 'uprchlíkem', 'uprchlíkovi', 'uprchlíku', 'uprchlíkyně', 'uprchlíků',
                'uvedenou', 'uvedená', 'uvedené', 'uvedeného', 'uvedenému', 'uvedený', 'uvedených', 'uvedeným',
                'vybranou', 'vybraná', 'vybrané', 'vybraného', 'vybranému', 'vybraný', 'vybraných', 'vybraným',
                'vyloučenou', 'vyloučená', 'vyloučené', 'vyloučeného', 'vyloučenému', 'vyloučený', 'vyloučených', 'vyloučeným',
                'vyšetřenou', 'vyšetřená', 'vyšetřené', 'vyšetřeného', 'vyšetřenému', 'vyšetřený', 'vyšetřených', 'vyšetřeným',
                'vyšetřenými', 'vyšetřovanou', 'vyšetřovaná', 'vyšetřované', 'vyšetřovaného', 'vyšetřovanému', 'vyšetřovaný', 'vyšetřovaných',
                'vyšetřovaným', 'vyšetřovanými', 'věřitel', 'věřitelce', 'věřitele', 'věřitelek', 'věřiteli', 'věřitelka',
                'věřitelky', 'věřitelů', 'zadrženou', 'zadržená', 'zadržené', 'zadrženého', 'zadrženém', 'zadrženému',
                'zadržený', 'zadržených', 'zadrženým', 'zadrženými',
                # Zdravotní/fyzický stav - ROZŠÍŘENÉ
                'zraněnou', 'zraněná', 'zraněné', 'zraněného', 'zraněném', 'zraněnému', 'zraněný', 'zraněných', 'zraněným', 'zraněnými',
                # Lehce/těžce zraněný
                'lehce zraněný', 'lehce zraněná', 'lehce zraněné', 'lehce zraněného', 'lehce zraněnému', 'lehce zraněném', 'lehce zraněným',
                'těžce zraněný', 'těžce zraněná', 'těžce zraněné', 'těžce zraněného', 'těžce zraněnému', 'těžce zraněném', 'těžce zraněným',
                # Nemocný
                'nemocný', 'nemocného', 'nemocnému', 'nemocném', 'nemocným',
                'nemocná', 'nemocné', 'nemocnou', 'nemocnými', 'nemocných',
                'těžce nemocný', 'těžce nemocná', 'těžce nemocné', 'těžce nemocného', 'těžce nemocnému', 'těžce nemocném', 'těžce nemocným',
                'chronicky nemocný', 'chronicky nemocná', 'chronicky nemocné', 'chronicky nemocného',
                'duševně nemocný', 'duševně nemocná', 'duševně nemocné', 'duševně nemocného',
                # Hospitalizovaný a další zdravotní stavy
                'hospitalizovaný', 'hospitalizovaného', 'hospitalizovanému', 'hospitalizovaném', 'hospitalizovaným', 'hospitalizovanými', 'hospitalizovaných',
                'hospitalizovaná', 'hospitalizované', 'hospitalizovanou',
                'transportovaný', 'transportovaného', 'transportovanému', 'transportovaném', 'transportovaným', 'transportovanými', 'transportovaných',
                'transportovaná', 'transportované', 'transportovanou',
                'intubovaný', 'intubovaného', 'intubovanému', 'intubovaném', 'intubovaným', 'intubovanými', 'intubovaných',
                'intubovaná', 'intubované', 'intubovanou',
                'resuscitovaný', 'resuscitovaného', 'resuscitovanému', 'resuscitovaném', 'resuscitovaným', 'resuscitovanými', 'resuscitovaných',
                'resuscitovaná', 'resuscitované', 'resuscitovanou',
                'stabilizovaný', 'stabilizovaného', 'stabilizovanému', 'stabilizovaném', 'stabilizovaným', 'stabilizovanými', 'stabilizovaných',
                'stabilizovaná', 'stabilizované', 'stabilizovanou',
                'léčený', 'léčeného', 'léčenému', 'léčeném', 'léčeným', 'léčenými', 'léčených',
                'léčená', 'léčené', 'léčenou',
                'amputovaný', 'amputovaného', 'amputovanému', 'amputovaném', 'amputovaným', 'amputovanými', 'amputovaných',
                'amputovaná', 'amputované', 'amputovanou',
                # Dárce (transplantace)
                'dárce', 'dárci', 'dárců', 'dárcem', 'dárcům', 'dárcích', 'dárkyně', 'dárkyní', 'dárkyň', 'dárkyni',
                # Specifické zdravotní stavy
                'tetraplegie', 'tetraplegik', 'tetraplegika', 'paraplegik', 'paraplegika',
                'imobilní', 'imobilního', 'imobilnímu', 'imobilním', 'imobilních',
                'komatózní', 'komatózního', 'komatóznímu', 'komatózním', 'komatózních',
                'bezvědomý', 'bezvědomého', 'bezvědomému', 'bezvědomém', 'bezvědomým', 'bezvědomých', 'bezvědomými',
                'bezvědomá', 'bezvědomé', 'bezvědomou',
                'umírající', 'umírajícího', 'umírajícímu', 'umírajícím', 'umírajících',
                # Mrtvý/zemřelý/zesnulý
                'mrtvý', 'mrtvého', 'mrtvému', 'mrtvém', 'mrtvým', 'mrtvých', 'mrtvými',
                'mrtvá', 'mrtvé', 'mrtvou',
                'zemřelý', 'zemřelého', 'zemřelému', 'zemřelém', 'zemřelým', 'zemřelých', 'zemřelými',
                'zemřelá', 'zemřelé', 'zemřelou',
                'zesnulý', 'zesnulého', 'zesnulému', 'zesnulém', 'zesnulým', 'zesnulých', 'zesnulými',
                'zesnulá', 'zesnulé', 'zesnulou',
                'deceased', 'exitus',
                'zaměstnance', 'zaměstnancem', 'zaměstnanci', 'zaměstnanců',
                'zaměstnanec', 'zaměstnankyni', 'zaměstnankyní', 'zaměstnankyně', 'zaměstnankyň', 'zaměstnavatel', 'zaměstnavatele', 'zaměstnavatelem',
                'zaměstnavateli', 'zaměstnavatelka', 'zaměstnavatelky', 'zaměstnavatelů', 'zletilou', 'zletilá', 'zletilé', 'zletilého',
                'zletilému', 'zletilý', 'účastnic', 'účastnice', 'účastnicí', 'účastník', 'účastníka', 'účastníkem',
                'účastníkovi', 'účastníku', 'účastníků', 'žadatel', 'žadatelce', 'žadatele', 'žadatelek', 'žadateli',
                'žadatelka', 'žadatelky', 'žadatelů', 'žalobce', 'žalobci', 'žalobců', 'žalobkyni', 'žalobkyní',
                'žalobkyně', 'žalovanou', 'žalovaná', 'žalované', 'žalovaného', 'žalovaném', 'žalovanému', 'žalovaný',
                'žalovaných', 'žalovaným', 'žalovanými', 'žáci', 'žák', 'žáka', 'žákem', 'žákovi',
                'žáku', 'žákyni', 'žákyní', 'žákyně', 'žákyň', 'žáků',
                # Další termíny
                'přítel', 'přítelkyně', 'kolega', 'kolegyně',
                'majitel', 'majitelka',
                'předseda', 'předsedkyně', 'člen', 'členka',
                'věznice', 'věznici', 'vězení',  # prison - IMPORTANT!
                # Ostatní
                'scrum', 'master', 'developer', 'architect', 'engineer',
                'officer', 'professional', 'certified', 'advanced',
                'management', 'legal', 'counsel', 'executive',
                # Pozdravy/oslovení
                'ahoj', 'dobrý', 'den', 'vážený', 'vážená',
                # Značky aut
                'škoda', 'octavia', 'fabia', 'superb', 'kodiaq',
                'volkswagen', 'toyota', 'ford', 'bmw', 'audi',
                # Technologie a software
                'google', 'amazon', 'microsoft', 'apple', 'facebook',
                'cloud', 'web', 'tech', 'solutions', 'data', 'digital',
                'software', 'enterprise', 'premium', 'standard',
                'analytics', 'computer', 'vision', 'protection',
                'security', 'authenticator', 'repository', 'access',
                'personal', 'hub', 'book', 'pro', 'series', 'launch',
                'team', 'development', 'react', 'splunk', 'innovate',
                'ventures', 'credo', 'mayo', 'clinic', 'met', 'london',
                'avenue', 'contractual', 'plánovaná', 'diagno',
                'cisco', 'processing', 'notářská',
                # Zdravotnictví + Léky/Produkty
                'nemocnice', 'poliklinika', 'polikliniek', 'nemocniec',
                'healthcare', 'symbicort', 'turbuhaler', 'spirometr',
                'jaeger', 'medical', 'health', 'pharma', 'pharmaceutical',
                # ===================================================================
                # OPATROVNICTVÍ A PÉČE - ROZŠÍŘENÉ
                # ===================================================================
                # Svěřenec/svěřenka (opatrovník už existuje výše)
                'svěřenec', 'svěřence', 'svěřenci', 'svěřencem', 'svěřenci', 'svěřenců', 'svěřencům', 'svěřencích',
                'svěřenka', 'svěřenky', 'svěřence', 'svěřenkou', 'svěřenek', 'svěřenkám', 'svěřenkách',
                # Osvojenec/osvojenka (osvojitel už existuje výše)
                'osvojenec', 'osvojence', 'osvojenci', 'osvojencem', 'osvojenci', 'osvojenců', 'osvojencům', 'osvojencích',
                'osvojenka', 'osvojenky', 'osvojence', 'osvojenkou', 'osvojenek', 'osvojenkám', 'osvojenkách',
                # Pečovatel/pečovatelka
                'pečovatel', 'pečovatele', 'pečovateli', 'pečovatelem', 'pečovatelé', 'pečovatelů', 'pečovatelům', 'pečovatelích',
                'pečovatelka', 'pečovatelky', 'pečovatelce', 'pečovatelkou', 'pečovatelek', 'pečovatelkám', 'pečovatelkách',
                # Ošetřující/ošetřovatel
                'ošetřující', 'ošetřujícího', 'ošetřujícímu', 'ošetřujícím', 'ošetřujících',
                'ošetřovatel', 'ošetřovatele', 'ošetřovateli', 'ošetřovatelem', 'ošetřovatelé', 'ošetřovatelů', 'ošetřovatelům', 'ošetřovatelích',
                'ošetřovatelka', 'ošetřovatelky', 'ošetřovatelce', 'ošetřovatelkou', 'ošetřovatelek', 'ošetřovatelkám', 'ošetřovatelkách',
                # Opatrovatel/opatrovatelka
                'opatrovatel', 'opatrovatele', 'opatrovateli', 'opatrovatelem', 'opatrovatelé', 'opatrovatelů', 'opatrovatelům', 'opatrovatelích',
                'opatrovatelka', 'opatrovatelky', 'opatrovatelce', 'opatrovatelkou', 'opatrovatelek', 'opatrovatelkám', 'opatrovatelkách',
                # Zákonný zástupce
                'zákonný zástupce', 'zákonného zástupce', 'zákonnému zástupci', 'zákonném zástupci', 'zákonným zástupcem',
                'zákonná zástupkyně', 'zákonné zástupkyně', 'zákonné zástupkyni', 'zákonnou zástupkyní',
                # Kontaktní osoba
                'kontaktní osoba', 'kontaktní osoby', 'kontaktní osobě', 'kontaktní osobou',
                'kontaktní', 'kontaktního', 'kontaktnímu', 'kontaktním', 'kontaktních',
                # ===================================================================
                # ZAMĚSTNÁNÍ A FUNKCE - ROZŠÍŘENÉ
                # ===================================================================
                # Pracovník/pracovnice
                'pracovník', 'pracovníka', 'pracovníkovi', 'pracovníku', 'pracovníkem', 'pracovníci', 'pracovníků', 'pracovníkům', 'pracovnících',
                'pracovnice', 'pracovnici', 'pracovnicí', 'pracovnic', 'pracovnicím', 'pracovnicích',
                # Dělník/dělnice
                'dělník', 'dělníka', 'dělníkovi', 'dělníku', 'dělníkem', 'dělníci', 'dělníků', 'dělníkům', 'dělnících',
                'dělnice', 'dělnici', 'dělnicí', 'dělnic', 'dělnicím', 'dělnicích',
                # Úředník/úřednice
                'úředník', 'úředníka', 'úředníkovi', 'úředníku', 'úředníkem', 'úředníci', 'úředníků', 'úředníkům', 'úřednících',
                'úřednice', 'úřednici', 'úřednicí', 'úřednic', 'úřednicím', 'úřednicích',
                # Ředitel/ředitelka (rozšíření)
                'ředitel', 'ředitele', 'řediteli', 'ředitelem', 'ředitelé', 'ředitelů', 'ředitelům', 'ředitelích',
                'ředitelka', 'ředitelky', 'ředitelce', 'ředitelkou', 'ředitelek', 'ředitelkám', 'ředitelkách',
                'ředitel školy', 'ředitele školy', 'řediteli školy', 'ředitelem školy',
                'ředitelka školy', 'ředitelky školy', 'ředitelce školy', 'ředitelkou školy',
                # Jednatel/jednatelka
                'jednatel', 'jednatele', 'jednateli', 'jednatelem', 'jednatelé', 'jednatelů', 'jednatelům', 'jednatelích',
                'jednatelka', 'jednatelky', 'jednatelce', 'jednatelkou', 'jednatelek', 'jednatelkám', 'jednatelkách',
                'jednatel společnosti',
                # Předseda rozšíření (základní forma už existuje)
                'předsedy', 'předsedovi', 'předsedou', 'předsedů', 'předsedům', 'předsedech',
                'předsedkyni', 'předsedkyní', 'předsedkyň', 'předsedkyním', 'předsedkyních',
                'předseda senátu', 'předsedy senátu', 'předsedovi senátu', 'předsedou senátu',
                # Místopředseda/místopředsedkyně
                'místopředseda', 'místopředsedy', 'místopředsedovi', 'místopředsedou', 'místopředsedové', 'místopředsedů', 'místopředsedům', 'místopředsedech',
                'místopředsedkyně', 'místopředsedkyni', 'místopředsedkyní', 'místopředsedkyň', 'místopředsedkyním', 'místopředsedkyních',
                # Manažer/manažerka
                'manažer', 'manažera', 'manažerovi', 'manažeru', 'manažerem', 'manažeři', 'manažerů', 'manažerům', 'manažerech',
                'manažerka', 'manažerky', 'manažerce', 'manažerkou', 'manažerek', 'manažerkám', 'manažerkách',
                # Vedoucí
                'vedoucí', 'vedoucího', 'vedoucímu', 'vedoucím', 'vedoucích',
                # Vlastník/vlastnice (majitel už existuje)
                'vlastník', 'vlastníka', 'vlastníkovi', 'vlastníku', 'vlastníkem', 'vlastníci', 'vlastníků', 'vlastníkům', 'vlastnících',
                'vlastnice', 'vlastnici', 'vlastnicí', 'vlastnic', 'vlastnicím', 'vlastnicích',
                'vlastník nemovitosti', 'vlastníka nemovitosti', 'vlastníkovi nemovitosti', 'vlastníkem nemovitosti',
                # Spoluvlastník/spoluvlastnice
                'spoluvlastník', 'spoluvlastníka', 'spoluvlastníkovi', 'spoluvlastníku', 'spoluvlastníkem', 'spoluvlastníci', 'spoluvlastníků', 'spoluvlastníkům', 'spoluvlastnících',
                'spoluvlastnice', 'spoluvlastnici', 'spoluvlastnicí', 'spoluvlastnic', 'spoluvlastnicím', 'spoluvlastnicích',
                'spoluvlastník bytu', 'spoluvlastníka bytu', 'spoluvlastníkovi bytu', 'spoluvlastníkem bytu',
                'podílový spoluvlastník', 'podílového spoluvlastníka', 'podílovému spoluvlastníkovi', 'podílovým spoluvlastníkem',
                # Akcionář/akcionářka
                'akcionář', 'akcionáře', 'akcionáři', 'akcionářem', 'akcionáři', 'akcionářů', 'akcionářům', 'akcionářích',
                'akcionářka', 'akcionářky', 'akcionářce', 'akcionářkou', 'akcionářek', 'akcionářkám', 'akcionářkách',
                # Společník/společnice
                'společník', 'společníka', 'společníkovi', 'společníku', 'společníkem', 'společníci', 'společníků', 'společníkům', 'společnících',
                'společnice', 'společnici', 'společnicí', 'společnic', 'společnicím', 'společnicích',
                # Statutární zástupce
                'statutární zástupce', 'statutárního zástupce', 'statutárnímu zástupci', 'statutárním zástupcem',
                # Prokurist/prokuristka
                'prokurist', 'prokurista', 'prokuristovi', 'prokuristu', 'prokuristem', 'prokuristé', 'prokuristů', 'prokuristům', 'prokuristech',
                'prokuristka', 'prokuristky', 'prokuristce', 'prokuristkou', 'prokuristek', 'prokuristkám', 'prokuristkách',
                # Zmocněnec/zmocněnkyně
                'zmocněnec', 'zmocněnce', 'zmocněnci', 'zmocněncem', 'zmocněnci', 'zmocněnců', 'zmocněncům', 'zmocněncích',
                'zmocněnkyně', 'zmocněnkyni', 'zmocněnkyní', 'zmocněnkyň', 'zmocněnkyním', 'zmocněnkyních',
                # ===================================================================
                # PRÁVNÍ OZNAČENÍ - VELMI ROZŠÍŘENÉ
                # ===================================================================
                # Dotčený/zúčastněný
                'dotčený', 'dotčeného', 'dotčenému', 'dotčeném', 'dotčeným', 'dotčených', 'dotčenými',
                'dotčená', 'dotčené', 'dotčenou',
                'zúčastněný', 'zúčastněného', 'zúčastněnému', 'zúčastněném', 'zúčastněným', 'zúčastněných', 'zúčastněnými',
                'zúčastněná', 'zúčastněné', 'zúčastněnou',
                # Zastupující/zástupce rozšíření
                'zastupující', 'zastupujícího', 'zastupujícímu', 'zastupujícím', 'zastupujících',
                'zástupce', 'zástupci', 'zástupcem', 'zástupců', 'zástupcům', 'zástupcích',
                'zástupkyně', 'zástupkyni', 'zástupkyní', 'zástupkyň', 'zástupkyním', 'zástupkyních',
                'zástupce ředitele', 'zástupci ředitele', 'zástupcem ředitele',
                # Advokát/advokátka
                'advokát', 'advokáta', 'advokátovi', 'advokátu', 'advokátem', 'advokáti', 'advokátů', 'advokátům', 'advokátech',
                'advokátka', 'advokátky', 'advokátce', 'advokátkou', 'advokátek', 'advokátkám', 'advokátkách',
                # Obhájce/obhájkyně
                'obhájce', 'obhájci', 'obhájcem', 'obhájců', 'obhájcům', 'obhájcích',
                'obhájkyně', 'obhájkyni', 'obhájkyní', 'obhájkyň', 'obhájkyním', 'obhájkyních',
                # Právní zástupce
                'právní zástupce', 'právního zástupce', 'právnímu zástupci', 'právním zástupcem',
                # Znalec/znalkyně
                'znalec', 'znalce', 'znalci', 'znalcem', 'znalci', 'znalců', 'znalcům', 'znalcích',
                'znalkyně', 'znalkyni', 'znalkyní', 'znalkyň', 'znalkyním', 'znalkyních',
                'soudní znalec', 'soudního znalce', 'soudnímu znalci', 'soudním znalcem',
                # Expert/expertka, odborník/odbornice
                'expert', 'experta', 'expertovi', 'expertu', 'expertem', 'experti', 'expertů', 'expertům', 'expertech',
                'expertka', 'expertky', 'expertce', 'expertkou', 'expertek', 'expertkám', 'expertkách',
                'odborník', 'odborníka', 'odborníkovi', 'odborníku', 'odborníkem', 'odborníci', 'odborníků', 'odborníkům', 'odbornicích',
                'odbornice', 'odbornici', 'odbornicí', 'odbornic', 'odbornicím', 'odbornicích',
                # Soudce/soudkyně
                'soudce', 'soudci', 'soudcem', 'soudců', 'soudcům', 'soudcích',
                'soudkyně', 'soudkyni', 'soudkyní', 'soudkyň', 'soudkyním', 'soudkyních',
                'samosoudce', 'samosoudci', 'samosoudcem', 'samosoudců', 'samosoudcům', 'samosoudcích',
                'samosoudkyně', 'samosoudkyni', 'samosoudkyní', 'samosoudkyň', 'samosoudkyním', 'samosoudkyních',
                # Státní zástupce/prokurátorka
                'státní zástupce', 'státního zástupce', 'státnímu zástupci', 'státním zástupcem',
                'státní zástupkyně', 'státní zástupkyni', 'státní zástupkyní',
                'prokurátor', 'prokurátora', 'prokurátorovi', 'prokurátoři', 'prokurátorů', 'prokurátorům', 'prokurátorech',
                'prokurátorka', 'prokurátorky', 'prokurátorce', 'prokurátorkou', 'prokurátorek', 'prokurátorkám', 'prokurátorkách',
                # Insolvenční správce
                'insolvenční správce', 'insolvenčního správce', 'insolvenčnímu správci', 'insolvenčním správcem',
                'insolvenční správkyně', 'insolvenční správkyni', 'insolvenční správkyní',
                'správce konkursní podstaty', 'správce konkursní podstaty', 'správci konkursní podstaty', 'správcem konkursní podstaty',
                # Probační úředník, kurátor
                'probační úředník', 'probačního úředníka', 'probačnímu úředníkovi', 'probačním úředníkem',
                'probační úřednice', 'probační úřednici', 'probační úřednicí',
                'kurátor', 'kurátora', 'kurátorovi', 'kurátoři', 'kurátorů', 'kurátorům', 'kurátorech',
                'kurátorka', 'kurátorky', 'kurátorce', 'kurátorkou', 'kurátorek', 'kurátorkám', 'kurátorkách',
                'sociální kurátor', 'sociálního kurátora', 'sociálnímu kurátorovi', 'sociálním kurátorem',
                # ===================================================================
                # NEMOVITOSTI
                # ===================================================================
                # Nájemce/nájemkyně
                'nájemce', 'nájemci', 'nájemcem', 'nájemců', 'nájemcům', 'nájemcích',
                'nájemkyně', 'nájemkyni', 'nájemkyní', 'nájemkyň', 'nájemkyním', 'nájemkyních',
                'nájemce bytu', 'nájemci bytu', 'nájemcem bytu',
                # Podnájemce/podnájemkyně
                'podnájemce', 'podnájemci', 'podnájemcem', 'podnájemců', 'podnájemcům', 'podnájemcích',
                'podnájemkyně', 'podnájemkyni', 'podnájemkyní', 'podnájemkyň', 'podnájemkyním', 'podnájemkyních',
                # Pronajímatel/pronajímatelka
                'pronajímatel', 'pronajímatele', 'pronajímateli', 'pronajímatelem', 'pronajímatelé', 'pronajímatelů', 'pronajímatelům', 'pronajímatelích',
                'pronajímatelka', 'pronajímatelky', 'pronajímatelce', 'pronajímatelkou', 'pronajímatelek', 'pronajímatelkám', 'pronajímatelkách',
                # Uživatel/uživatelka
                'uživatel', 'uživatele', 'uživateli', 'uživatelem', 'uživatelé', 'uživatelů', 'uživatelům', 'uživatelích',
                'uživatelka', 'uživatelky', 'uživatelce', 'uživatelkou', 'uživatelek', 'uživatelkám', 'uživatelkách',
                'oprávněný uživatel', 'oprávněného uživatele', 'oprávněnému uživateli', 'oprávněným uživatelem',
                'neoprávněný uživatel', 'neoprávněného uživatele', 'neoprávněnému uživateli', 'neoprávněným uživatelem',
                # Investor stavby
                'investor stavby', 'investora stavby', 'investorovi stavby', 'investorem stavby',
                # ===================================================================
                # ŠKOLSTVÍ - ROZŠÍŘENÉ
                # ===================================================================
                # Absolvent/absolventka
                'absolvent', 'absolventa', 'absolventovi', 'absolventu', 'absolventem', 'absolventi', 'absolventů', 'absolventům', 'absolventech',
                'absolventka', 'absolventky', 'absolventce', 'absolventkou', 'absolventek', 'absolventkám', 'absolventkách',
                # Učitel/učitelka
                'učitel', 'učitele', 'učiteli', 'učitelem', 'učitelé', 'učitelů', 'učitelům', 'učitelích',
                'učitelka', 'učitelky', 'učitelce', 'učitelkou', 'učitelek', 'učitelkám', 'učitelkách',
                'třídní učitel', 'třídního učitele', 'třídnímu učiteli', 'třídním učitelem',
                'třídní učitelka', 'třídní učitelky', 'třídní učitelce', 'třídní učitelkou',
                # Pedagog/pedagožka
                'pedagog', 'pedagoga', 'pedagogovi', 'pedagogu', 'pedagogem', 'pedagogové', 'pedagogů', 'pedagogům', 'pedagogech',
                'pedagožka', 'pedagožky', 'pedagožce', 'pedagožkou', 'pedagožek', 'pedagožkám', 'pedagožkách',
                # Vyučující
                'vyučující', 'vyučujícího', 'vyučujícímu', 'vyučujícím', 'vyučujících',
                # Lektor/lektorka
                'lektor', 'lektora', 'lektorovi', 'lektoru', 'lektorem', 'lektoři', 'lektorů', 'lektorům', 'lektorech',
                'lektorka', 'lektorky', 'lektorce', 'lektorkou', 'lektorek', 'lektorkám', 'lektorkách',
                # ===================================================================
                # SPECIFICKÉ SITUACE
                # ===================================================================
                # Osoba identifikovaná, fyzická osoba
                'osoba identifikovaná jako', 'osoby identifikované jako', 'osobě identifikované jako',
                'osoba jménem', 'osoby jménem', 'osobě jménem', 'osobou jménem',
                'fyzická osoba', 'fyzické osoby', 'fyzické osobě', 'fyzickou osobou',
                # Občan (rozšíření - základní formy už existují)
                'občan', 'občana', 'občanovi', 'občanu', 'občanem', 'občané', 'občanů', 'občanům', 'občanech',
                'občanka', 'občanky', 'občance',
                # Přítomný/nepřítomný
                'přítomný', 'přítomného', 'přítomněm u', 'přítomném', 'přítomným', 'přítomných', 'přítomným i',
                'přítomná', 'přítomné', 'přítomnou',
                'nepřítomný', 'nepřítomného', 'nepřítomněm u', 'nepřítomném', 'nepřítomným', 'nepřítomných', 'nepřítomným i',
                'nepřítomná', 'nepřítomné', 'nepřítomnou',
                'dostavivší se', 'dostavivšího se', 'dostavivšímu se', 'dostavivším se', 'dostavivších se',
                # Jmenovaný/uvedený
                'jmenovaný', 'jmenovaného', 'jmenovanému', 'jmenovaném', 'jmenovaným', 'jmenovaných', 'jmenovanými',
                'jmenovaná', 'jmenované', 'jmenovanou',
                'jmenovaný do funkce', 'jmenovaného do funkce', 'jmenovanému do funkce', 'jmenovaným do funkce',
                'výše uvedený', 'výše uvedeného', 'výše uvedenému', 'výše uvedeném', 'výše uvedeným',
                'výše uvedená', 'výše uvedené', 'výše uvedenou',
                'níže uvedený', 'níže uvedeného', 'níže uvedenému', 'níže uvedeném', 'níže uvedeným',
                'níže uvedená', 'níže uvedené', 'níže uvedenou',
                'shora uvedený', 'shora uvedeného', 'shora uvedenému', 'shora uvedeném', 'shora uvedeným',
                'shora uvedená', 'shora uvedené', 'shora uvedenou',
                # Zvolený/pověřený/delegovaný
                'zvolený', 'zvoleného', 'zvolenému', 'zvoleném', 'zvoleným', 'zvolených', 'zvoleným i',
                'zvolená', 'zvolené', 'zvolenou',
                'pověřený', 'pověřeného', 'pověřenému', 'pověřeném', 'pověřeným', 'pověřených', 'pověřenými',
                'pověřená', 'pověřené', 'pověřenou',
                'delegovaný', 'delegovaného', 'delegovanému', 'delegovaném', 'delegovaným', 'delegovaných', 'delegovanými',
                'delegovaná', 'delegované', 'delegovanou',
                # Zbavený/odvolaný (vyloučený už existuje)
                'zbavený', 'zbaveného', 'zbavenému', 'zbaveném', 'zbaveným', 'zbavených', 'zbavenými',
                'zbavená', 'zbavené', 'zbavenou',
                'odvolaný', 'odvolaného', 'odvolanému', 'odvolaném', 'odvolaným', 'odvolaných', 'odvolanými',
                'odvolaná', 'odvolané', 'odvolanou',
                # Azylant/emigrant/imigrant (uprchlík už existuje)
                'azylant', 'azylanta', 'azylantovi', 'azylantu', 'azylantem', 'azylanti', 'azylantů', 'azylantům', 'azylantech',
                'azylantka', 'azylantky', 'azylantce', 'azylantkou', 'azylantek', 'azylantkám', 'azylantkách',
                'emigrant', 'emigranta', 'emigrantovi', 'emigrantu', 'emigrantem', 'emigranti', 'emigrantů', 'emigrantům', 'emigrantech',
                'emigrantka', 'emigrantky', 'emigrantce', 'emigrantkou', 'emigrantek', 'emigrantkám', 'emigrantkách',
                'imigrant', 'imigranta', 'imigrantovi', 'imigrantu', 'imigrantem', 'imigranti', 'imigrantů', 'imigrantům', 'imigrantech',
                'imigrantka', 'imigrantky', 'imigrantce', 'imigrantkou', 'imigrantek', 'imigrantkám', 'imigrantkách',
                'žadatel o azyl', 'žadatele o azyl', 'žadateli o azyl', 'žadatelem o azyl',
                'cizinec', 'cizince', 'cizinci', 'cizincem', 'cizinci', 'cizinců', 'cizincům', 'cizincích',
                'cizinka', 'cizinky', 'cizince', 'cizinkou', 'cizinek', 'cizinkám', 'cizinkách',
                # Senior/důchodce (rozšíření)
                'senior', 'seniora', 'seniorovi', 'senioru', 'seniorem', 'senioři', 'seniorů', 'seniorům', 'seniorech',
                'seniorka', 'seniorky', 'seniorce', 'seniorkou', 'seniorek', 'seniorkám', 'seniorkách',
                'důchodce', 'důchodci', 'důchodcem', 'důchodců', 'důchodcům', 'důchodcích',
                'důchodkyně', 'důchodkyni', 'důchodkyní', 'důchodkyň', 'důchodkyním', 'důchodkyních',
                # Těhotná/gravidní/rodička
                'těhotná', 'těhotné', 'těhotnou', 'těhotných', 'těhotným', 'těhotnými',
                'gravidní', 'gravidní', 'gravidních', 'gravidním', 'gravidními',
                'porodnice', 'porodnici', 'porodnicí', 'porodnic', 'porodnicím', 'porodnicích',
                # Pozůstalý/dědic
                'pozůstalý', 'pozůstalého', 'pozůstalému', 'pozůstalém', 'pozůstalým', 'pozůstalých', 'pozůstalými',
                'pozůstalá', 'pozůstalé', 'pozůstalou',
                'dědic', 'dědice', 'dědici', 'dědicem', 'dědici', 'dědiců', 'dědicům', 'dědicích',
                'dědička', 'dědičky', 'dědičce', 'dědičkou', 'dědiček', 'dědičkám', 'dědičkách',
                'odkázaný', 'odkázaného', 'odkázanému', 'odkázaném', 'odkázaným', 'odkázaných', 'odkázanými',
                'odkázaná', 'odkázané', 'odkázanou',
                # Bývalý/původní/nový
                'bývalý', 'bývalého', 'bývalému', 'bývalém', 'bývalým', 'bývalých', 'bývalými',
                'bývalá', 'bývalé', 'bývalou',
                'původní', 'původního', 'původnímu', 'původním', 'původních',
                'předchozí', 'předchozího', 'předchozímu', 'předchozím', 'předchozích',
                'nový', 'nového', 'novému', 'novém', 'novým', 'nových', 'novými',
                'nová', 'nové', 'novou',
                'nynější', 'nynějšího', 'nynějšímu', 'nynějším', 'nynějších',
                'současný', 'současného', 'současnému', 'současném', 'současným', 'současných', 'současnými',
                'současná', 'současné', 'současnou',
                # ===================================================================
                # KOMBINOVANÉ TERMÍNY
                # ===================================================================
                # Manželé/rodiče/sourozenci (plurály)
                'manželé', 'manželů', 'manželům', 'manželích', 'manželi',
                # Pan/paní (rozšíření)
                'pan', 'pana', 'panovi', 'panu', 'panem', 'pánové', 'pánů', 'pánům', 'pánech',
                'paní', 'paní', 'paním', 'paních',
                'pán', 'pána', 'pánovi', 'pánu', 'pánom',
                'pani', 'pani', 'paním', 'paních',
                # S tituly
                'pan doktor', 'pana doktora', 'panu doktorovi', 'panem doktorem',
                'paní doktorka', 'paní doktorky', 'paní doktorce', 'paní doktorkou',
                'pan profesor', 'pana profesora', 'panu profesorovi', 'panem profesorem',
                'paní profesorka', 'paní profesorky', 'paní profesorce', 'paní profesorkou',
                'pan inženýr', 'pana inženýra', 'panu inženýrovi', 'panem inženýrem',
                'paní inženýrka', 'paní inženýrky', 'paní inženýrce', 'paní inženýrkou',
                # Další
                'care', 'plus', 'minus', 'service', 'services',
                'group', 'company', 'corp', 'ltd', 'gmbh', 'inc'
            }

            # Kontrola proti ignore listu
            # NOVÁ LOGIKA: Pokud první slovo je titul a druhé je příjmení → anonymizuj příjmení
            if first_obs.lower() in ignore_words:
                # První slovo je titul → může to být "Titul Příjmení" (bez křestního jména)
                # Zkontroluj, zda druhé slovo není také v ignore_words
                if last_obs.lower() in ignore_words:
                    return match.group(0)  # Obě slova jsou ignorovaná → není osoba

                # První slovo = titul, druhé = příjmení → anonymizuj příjmení
                # Pokud existuje osoba s tímto příjmením, anonymizuj ji
                # Infer nominative of surname
                last_nom = infer_surname_nominative(last_obs)

                # Try to find existing person with this surname
                tag = None
                canonical_full = None
                for existing_tag, existing_canonical in self.person_canonical_names.items():
                    # Check if surname matches
                    canonical_parts = existing_canonical.split()
                    if len(canonical_parts) >= 2:  # Has both first and last name
                        existing_surname = canonical_parts[-1]
                        if existing_surname.lower() == last_nom.lower():
                            tag = existing_tag
                            canonical_full = existing_canonical
                            # Save variant if different from canonical surname
                            if last_obs.lower() != existing_surname.lower():
                                self.entity_map['PERSON'][canonical_full].add(last_obs)
                            break

                # If existing person found, anonymize it
                if tag:
                    return f"{first_obs} {tag}"
                else:
                    # No existing person with this surname → don't anonymize
                    # (We can't create a person with just a surname)
                    return match.group(0)

            # Pokud druhé slovo je v ignore_words (ale první ne) → není osoba
            if last_obs.lower() in ignore_words:
                return match.group(0)

            # 3. Detekce firem, produktů, institucí (neměly by být PERSON)
            non_person_patterns = [
                # Tech/Software
                r'\b(tech|cloud|web|solutions?|data|digital|software|analytics)\b',
                r'\b(team|hub|enterprise|premium|standard|professional)\b',
                r'\b(google|amazon|microsoft|apple|facebook|splunk|cisco)\b',
                r'\b(repository|authenticator|vision|protection|security)\b',
                # Tech brands/phones
                r'\b(samsung|huawei|xiaomi|nokia|sony|lg|motorola|oppo|vivo)\b',
                r'\b(galaxy|iphone|pixel|nexus|xperia)\b',
                # Finance/Investment
                r'\b(capital|equity|value|investment|fund|holdings|assets)\b',
                r'\b(crescendo|ventures|partners|portfolio)\b',
                # Management/Business
                r'\b(management|processing|executive|legal|counsel)\b',
                r'\b(clinic|series|launch|innovate|healthcare)\b',
                # Products/Medical
                r'\b(symbicort|turbuhaler|spirometr|jaeger)\b',
                r'\b(pharma|pharmaceutical|medical)\b',
                # Company/organization words
                r'\b(společnost[ií]?)\b',  # Společností, společnosti, společnost
                r'\b(někter[ýáéěí])\b',  # Některý, některá, některé, některí
                # Religious/Place names
                r'\b(svat[éého]|svatá)\b',  # Svaté, Svatého, Svatá (Saint)
                r'\b(kostel|chrám|kaple|církev)\b',  # Church, temple, chapel
                # Czech cities and places
                r'\b(nový\s+jičín|nové\s+město|staré\s+město)\b',
                r'\b(mladá\s+boleslav|české\s+budějovice|hradec\s+králové|hradci\s+králové)\b',
                # Company suffixes když jsou uprostřed
                r'\b(group|company|corp|ltd|gmbh|inc|services?)\b'
            ]
            for pattern in non_person_patterns:
                if re.search(pattern, combined):
                    return match.group(0)

            # 4. Detekce názvů firem (končí na s.r.o., a.s., spol., Ltd. atd.)
            context_after = text[match.end():match.end()+20]
            if re.search(r'^\s*(s\.r\.o\.|a\.s\.|spol\.|k\.s\.|v\.o\.s\.|ltd\.?|inc\.?)', context_after, re.IGNORECASE):
                return match.group(0)

            # ========== B) VALIDACE ČESKÉ OSOBY ==========

            # 5. Požadované patterny pro skutečnou osobu
            # Max 2-3 tokeny (již splněno regex patternem)
            # Každý token začíná velkým písmenem (již splněno)

            # 6. Validace českého příjmení (poslední token)
            last_lo = last_obs.lower()

            # Typické české příjmení koncovky
            valid_surname_suffixes = (
                'ová', 'á',  # ženské
                'ek', 'ák', 'ík', 'ský', 'cký', 'čák', 'ec', 'el',  # mužské
                'a',  # Svoboda, Skála, Liška
                'ý', 'í',  # přídavná jména
                # Další běžné koncovky
                'an', 'en', 'in', 'on', 'un',  # Urban, Marin, Kubín, atd.
                'eš', 'iš', 'uš', 'áš', 'íš',  # Beneš, Kříž, Lukáš, atd.
                'or', 'ar', 'ir', 'ur',  # Gregor, Kohár, atd.
                'ov', 'ev', 'av', 'iv',  # Petrov, Medveděv, atd.
                'áč', 'ič', 'oč', 'ůč',  # Horváč, Novič, atd.
                'át', 'ůt', 'ut', 'et'   # Sovát, Kůt, atd.
            )

            # Pokud příjmení nekončí na typickou koncovku → pravděpodobně není osoba
            # ALE: pokud je to jednoslabičné anglické slovo (např. "Met", "Hub"), může to být produkt/firma
            if not last_lo.endswith(valid_surname_suffixes):
                # Zkontroluj, jestli je to jednoslabičné anglické slovo (firma/produkt)
                # Např: "Met London", "Hub Team", "Pro Series"
                if len(last_obs) <= 3 or last_obs.lower() in {'hub', 'pro', 'met', 'net', 'web', 'app', 'lab', 'dev'}:
                    return match.group(0)  # Pravděpodobně firma/produkt
                # Jinak je to OK (může to být méně běžné české příjmení)

            # 7. Validace křestního jména (musí být v knihovně nebo mít typickou českou strukturu)
            first_lo = first_obs.lower()

            # Whitelist běžných českých jmen (nejsou v knihovně, ale jsou validní)
            common_czech_names = {'jan', 'petr', 'pavel', 'jiří', 'josef', 'tomáš', 'martin', 'jakub', 'david', 'daniel'}

            # Pokud křestní jméno JE v knihovně nebo v whitelistu → OK
            if first_lo in CZECH_FIRST_NAMES or first_lo in common_czech_names:
                pass  # OK
            else:
                # Není v knihovně ani v whitelistu → kontroluj strukturu

                # Pokud má méně než 3 znaky → není validní (např. "Me", "Jo")
                if len(first_obs) < 3:
                    return match.group(0)

                # Pokud má 3 znaky:
                if len(first_obs) == 3:
                    # Pokud nekončí na samohlásku ani na 'n', 'l', 'r' → zkrácený genitiv
                    # "Jan", "Dan", "Ivo" = OK
                    # "Han" (z "Hana"), "Jev" (z "Eva") = NENÍ OK
                    if not first_lo[-1] in 'aeiouyáéěíóúůýnlr':
                        return match.group(0)

                # Pokud má 4-5 znaků:
                if 4 <= len(first_obs) <= 5:
                    # Pokud končí na 'k' → skoro vždy zkrácený genitiv (Elišk, Radk)
                    if first_lo[-1] == 'k':
                        return match.group(0)
                    # Pokud nekončí na samohlásku ani na typickou mužskou koncovku
                    if not first_lo[-1] in 'aeiouyáéěíóúůýnlršm':
                        return match.group(0)

            # 8. Detekce rolí ("Ředitelka Centrum")
            # Pokud první slovo je role → není to osoba
            role_words = {
                'ředitelka', 'ředitel', 'jednatel', 'jednatelka',
                'manager', 'director', 'chief', 'officer',
                'specialist', 'consultant', 'coordinator',
                'developer', 'architect', 'engineer', 'analyst'
            }
            if first_obs.lower() in role_words:
                return match.group(0)  # Role, ne osoba

            # ========== C) INFERENCE KANONICKÉHO JMÉNA ==========

            # NEJDŘÍV zjisti pohlaví podle křestního jména (před inferencí příjmení!)
            first_lo = first_obs.lower()

            # Detekce pohlaví podle křestního jména
            # DŮLEŽITÉ: Musíme nejdřív udělat inferenci, protože "Pavla" může být genitiv od "Pavel" (mužské)!
            # Inference už obsahuje normalizaci variant (Julia→Julie, atd.)
            first_nom_temp = infer_first_name_nominative(first_obs)
            first_nom_temp_lo = first_nom_temp.lower() if first_nom_temp else first_lo

            is_female_firstname = False
            if first_nom_temp_lo in CZECH_FIRST_NAMES:
                # Jméno je v knihovně - zkontroluj, zda je ženské
                is_female_firstname = first_nom_temp_lo.endswith(('a', 'e', 'ie', 'ia'))
            else:
                # Není v knihovně - heuristika podle koncovky
                is_female_firstname = first_lo.endswith(('a', 'e', 'ie', 'ia'))

            # ========== DŮLEŽITÁ KONTROLA: JMÉNO NA -A MŮŽE BÝT GENITIV MUŽSKÉHO JMÉNA! ==========
            # Před přidáním -ová k příjmení zkontroluj, zda jméno končící na -a není genitiv mužského jména
            # Příklady: Marka → Marek, Oldřicha → Oldřich, Stanislava → Stanislav, atd.
            could_be_male_genitive = False
            if first_lo.endswith('a') and len(first_lo) > 2:
                # Zkus odstranit -a a podívat se, zda vznikne známé mužské jméno
                candidate_male = first_obs[:-1]
                candidate_male_lo = candidate_male.lower()

                # Seznam běžných mužských jmen (nominativ bez -a)
                common_male_names = {
                    'marek', 'oldřich', 'bedřich', 'stanislav', 'radoslav', 'václav',
                    'leon', 'albert', 'erik', 'teodor', 'viktor', 'igor', 'artur',
                    'oleksandr', 'sergej', 'oleg', 'mihail', 'denis', 'ivan',
                    'lubomír', 'přemysl', 'tadeáš', 'rostislav', 'ctibor',
                    # Jména která už jsou v male_genitiv_a (níže)
                    'josef', 'emil', 'odon', 'štěpán', 'maxim', 'adam',
                    'matěj', 'jakub', 'lukáš', 'jan', 'petr', 'pavel',
                    'karel', 'michal', 'tomáš', 'aleš', 'miloš', 'leoš'
                }

                # Známé genitivní formy mužských jmen (kde odstranění -a nedá správný tvar)
                # Marka → Marek, Karla → Karel, Michala → Michal
                male_genitive_forms = {
                    'marka': 'marek',
                    'karla': 'karel',
                    'michala': 'michal',
                    'pavla': 'pavel',
                    'vita': 'vito',  # Cizí jména
                    'bruna': 'bruno',
                    'lea': 'leo'
                }

                # Pokud je to známá genitivní forma, je to určitě genitiv mužského jména
                if first_lo in male_genitive_forms:
                    could_be_male_genitive = True
                    is_female_firstname = False  # Override!
                # Nebo pokud po odstranění -a vznikne známé mužské jméno
                elif candidate_male_lo in common_male_names or candidate_male_lo in CZECH_FIRST_NAMES:
                    # Ověř, že kandidát nevypadá jako ženské jméno
                    if not candidate_male_lo.endswith(('a', 'e', 'ie', 'ia', 'y')):
                        could_be_male_genitive = True
                        is_female_firstname = False  # Override! Není to ženské jméno, je to genitiv mužského!

            # Nejdřív inference příjmení
            last_nom = infer_surname_nominative(last_obs)

            # DŮLEŽITÉ: Pokud je křestní jméno ženské a příjmení je v mužském tvaru,
            # převeď příjmení na ženský tvar (-ová)
            # ALE POUZE pokud jméno NENÍ pravděpodobně genitiv mužského jména!
            last_nom_lo = last_nom.lower()
            if is_female_firstname and not could_be_male_genitive and not last_nom_lo.endswith(('ová', 'á')):
                # Příjmení je mužské, ale jméno je ženské → přidej -ová
                # Novák → Nováková, Šustr → Šustrová
                if last_nom_lo.endswith(('k', 't', 'r', 's', 'n', 'l', 'c', 'č', 'š', 'ž', 'd', 'ď', 'ť', 'ň')):
                    last_nom = last_nom + 'ová'

            # DŮLEŽITÉ: Oprava kanonického příjmení
            # Pokud už máme v canonical_persons nějaký tvar tohoto příjmení (např. "Procházka"),
            # použij ten existující kmen místo nového inference
            # Např: "Jakub Procházka" → canonical = "Procházka"
            #       "Petra Procházková" → canonical by měl být "Procházková" (ne "Procházek")

            # Hledej existující kmen příjmení v canonical_persons
            # ALE POUZE pokud mají STEJNÝ ROD (mužské vs ženské)
            existing_surname_stem = None

            # Nejdřív zjisti rod aktuálního příjmení
            current_is_female = last_nom.lower().endswith(('ová', 'á'))

            for p in self.canonical_persons:
                existing_last = p['last']
                if not existing_last:
                    continue

                existing_last_lo = existing_last.lower()

                # Zjisti rod existujícího příjmení
                existing_is_female = existing_last_lo.endswith(('ová', 'á'))

                # MUSÍ být stejný rod!
                if existing_is_female != current_is_female:
                    continue  # Různý rod → skip

                # Porovnej kmeny příjmení
                # Procházka vs Procházka → kmen = "Procházk" (oba mužské) ✓
                # Procházková vs Procházková → kmen = "Procházk" (obě ženské) ✓
                # Procházka vs Procházková → skip (různý rod) ✗

                # Jednoduché pravidlo: odstraň koncovky -ová, -a, -ek, -el, -ec
                def get_stem(surname):
                    s = surname.lower()
                    if s.endswith('ová'):
                        return s[:-3]  # Procházková → Procházk
                    elif s.endswith('ek'):
                        return s[:-2] + 'k'  # Hájek → Hájk
                    elif s.endswith('el'):
                        return s[:-2] + 'l'  # Havel → Havl
                    elif s.endswith('ec'):
                        return s[:-2] + 'c'  # Němec → Němc
                    elif s.endswith('a'):
                        return s[:-1]  # Procházka → Procházk, Paseka → Pasek
                    elif s.endswith('á'):
                        return s[:-1]  # Malá → Mal
                    else:
                        return s  # Novák → Novák

                existing_stem = get_stem(existing_last)
                current_stem = get_stem(last_nom)

                # Pokud kmeny se shodují A mají stejný rod → použij existující tvar
                if existing_stem == current_stem:
                    existing_surname_stem = existing_last
                    break

            # Pokud jsme našli existující kmen (STEJNÝ ROD), použij ho místo inference
            if existing_surname_stem:
                last_nom = existing_surname_stem

            # Určení rodu podle příjmení
            is_female_surname = last_nom.lower().endswith(('ová', 'á'))

            # Inference křestního jména podle rodu příjmení
            # first_lo už je definováno výše (řádek 1670)

            # Gender-aware inference založená na příjmení
            is_female_surname = last_nom.lower().endswith(('ová', 'á'))

            if is_female_surname:
                # Příjmení je ženské → jméno musí být ženské
                first_lo = first_obs.lower() if first_obs else ''

                # Pro pádové koncovky použij inference
                if first_lo.endswith(('u', 'i', 'ě', 'ou', 'ií', 'n', 'y', 'e')):
                    # Pavlu→Pavla, Pavli→Pavla, Pavlou→Pavla, Elen→Elena, Boženy→Božena, Denise→Denis
                    inferred = infer_first_name_nominative(first_obs)
                    # Ale pokud inference vrátila mužské jméno, přidej 'a'
                    if not inferred.lower().endswith(('a', 'e', 'ie', 'ia', 'y')):
                        first_nom = (inferred + 'a').capitalize()
                    else:
                        first_nom = inferred
                elif first_lo.endswith(('a', 'ie', 'ia')):
                    # Nominativ (končí na typickou ženskou koncovku) → pouze normalizuj varianty
                    first_nom = normalize_name_variant(first_obs) if first_obs else first_obs
                else:
                    # Souhláska → přidej 'a'
                    first_nom = (first_obs + 'a').capitalize() if first_obs else first_obs
            else:
                # Příjmení je mužské → použij standardní inference
                # DŮLEŽITÉ: Kontrola, jestli jméno končící na -a není genitiv od mužského jména
                # Příklad: "Josefa Malého" → Josef (genitiv), ne Josefa (ženské jméno)
                first_lo = first_obs.lower() if first_obs else ''

                if first_lo.endswith('a') and len(first_lo) > 2:
                    # Zkus odstranit -a a podívat se, jestli vznikne mužské jméno
                    candidate_male = first_obs[:-1]
                    candidate_male_lo = candidate_male.lower()

                    # Seznam známých mužských jmen, která mají genitiv na -a
                    male_genitiv_a = {
                        'josef', 'emil', 'odon', 'štěpán', 'maxim', 'adam',
                        'matěj', 'jakub', 'lukáš', 'jan', 'petr', 'pavel',
                        'marek', 'oldřich', 'bedřich', 'stanislav', 'radoslav', 'václav',
                        'leon', 'albert', 'erik', 'teodor', 'viktor', 'igor', 'artur',
                        'oleksandr', 'sergej', 'oleg', 'mihail', 'denis', 'ivan',
                        'lubomír', 'přemysl', 'tadeáš', 'rostislav', 'ctibor',
                        'karel', 'michal', 'tomáš', 'aleš', 'miloš', 'leoš', 'radim'
                    }

                    # Známé genitivní formy (kde odstranění -a nedá správný tvar)
                    male_genitive_forms = {
                        'marka': 'marek',
                        'karla': 'karel',
                        'michala': 'michal',
                        'pavla': 'pavel',
                        'vita': 'vito',
                        'bruna': 'bruno',
                        'lea': 'leo'
                    }

                    # Pokud je to známá genitivní forma s mapováním, použij správný tvar
                    if first_lo in male_genitive_forms:
                        first_nom = male_genitive_forms[first_lo].capitalize()
                    # Nebo pokud po odstranění -a vznikne známé mužské jméno, použij ho
                    elif candidate_male_lo in male_genitiv_a or candidate_male_lo in CZECH_FIRST_NAMES:
                        # Ověř, že to není ženské jméno jako Josefa/Emila (která jsou samostatná)
                        # Pokud existuje jak mužská tak ženská forma, preferuj mužskou u mužského příjmení
                        if not candidate_male_lo.endswith(('a', 'e', 'ie', 'ia', 'y')):
                            first_nom = candidate_male
                        else:
                            first_nom = infer_first_name_nominative(first_obs) if first_obs else first_obs
                    else:
                        # Jinak použij standardní inference
                        first_nom = infer_first_name_nominative(first_obs) if first_obs else first_obs
                elif first_lo.endswith('e') and len(first_lo) > 2:
                    # Kontrola pro jména končící na -e (Denise → Denis)
                    # U mužského příjmení preferuj mužskou formu bez -e
                    candidate_male = first_obs[:-1]
                    candidate_male_lo = candidate_male.lower()

                    # Pokud existuje mužská forma bez -e, použij ji
                    if candidate_male_lo in CZECH_FIRST_NAMES:
                        first_nom = candidate_male
                    else:
                        # Jinak použij standardní inference
                        first_nom = infer_first_name_nominative(first_obs) if first_obs else first_obs
                else:
                    # Standardní inference pro ostatní případy
                    first_nom = infer_first_name_nominative(first_obs) if first_obs else first_obs

            # Vytvoř nebo najdi tag pro tuto osobu
            # PASS OBSERVED FORMS to ensure canonical matches what's in the document!
            tag, canonical = self._ensure_person_tag(first_nom, last_nom, first_obs, last_obs)

            # Ulož původní formu jako variantu (pokud je jiná než kanonická)
            original_form = f"{first_obs} {last_obs}"
            if original_form.lower() != canonical.lower():
                self.entity_map['PERSON'][canonical].add(original_form)

            return tag

        # ========== ELEGANTNÍ ŘEŠENÍ: finditer() + filtrace ==========
        # Místo sekvenčního sub() najdeme VŠECHNY matche a vyfiltrujeme nevalidní

        # Definice role slov (která nejsou křestní jména)
        role_words = {
            'ředitelka', 'ředitel', 'jednatel', 'jednatelka',
            'manager', 'director', 'chief', 'officer',
            'specialist', 'consultant', 'coordinator',
            'developer', 'architect', 'engineer', 'analyst',
            'řidič', 'řidička', 'řidiče', 'řidiči', 'řidičem',  # všechny pády
            'klient', 'klienta', 'klientka', 'klientky', 'klientovi',
            'pacient', 'pacienta', 'pacientka', 'pacientky',
            'žadatel', 'žadatele', 'žadatelka', 'žadatelky'
        }

        # 1. Najdi všechny 3-slovné matche (včetně překrývajících se!)
        # DŮLEŽITÉ: finditer() nenachází překryvy, musíme hledat manuálně
        matches_3word = []
        pos = 0
        while pos < len(text):
            match = role_person_pattern.search(text, pos)
            if not match:
                break
            role_word = match.group(1)
            # Platný match pouze pokud první slovo JE v ignore_words (je to titul/role)
            if role_word.lower() in ignore_words:
                matches_3word.append(match)
            # Posun o 1 znak pro nalezení překrývajících se matchů
            pos = match.start() + 1

        # 2. Najdi všechny 2-slovné matche (včetně překrývajících se!)
        matches_2word = []
        pos = 0
        while pos < len(text):
            match = person_pattern.search(text, pos)
            if not match:
                break
            first_obs = match.group(1)
            # Platný match pouze pokud první slovo NENÍ role word
            # (tj. vypadá jako křestní jméno, ne jako "Řidič", "Klient", atd.)
            if first_obs.lower() not in role_words:
                matches_2word.append(match)
            # Posun o 1 znak pro nalezení překrývajících se matchů
            pos = match.start() + 1

        # 3. Kombinuj matche a odstraň překryvy (preferuj delší = 3-slovné)
        all_matches = []

        # Přidej 3-slovné (mají prioritu)
        for match in matches_3word:
            all_matches.append(('3word', match))

        # Přidej 2-slovné, ale pouze pokud se nepřekrývají s 3-slovnými
        for match in matches_2word:
            overlaps = False
            for _, m3 in [m for m in all_matches if m[0] == '3word']:
                # Překryv = matche sdílejí nějaký znak
                if not (match.end() <= m3.start() or match.start() >= m3.end()):
                    overlaps = True
                    break
            if not overlaps:
                all_matches.append(('2word', match))

        # 4. Seřaď podle pozice (od konce, aby se neposunuly indexy při nahrazování)
        all_matches.sort(key=lambda x: x[1].start(), reverse=True)

        # 5. Aplikuj replacementy od konce
        for match_type, match in all_matches:
            if match_type == '3word':
                replacement = replace_role_person(match)
            else:  # '2word'
                replacement = replace_person(match)

            # Nahraď v textu
            text = text[:match.start()] + replacement + text[match.end():]

        # ========== NOVÝ: Detekce samostatných příjmení ==========
        # Po zpracování celých jmen, hledej samostatně se vyskytující příjmení
        # Např: "Podnikatel Benedikt Mencl (...) Menclovi bylo doporučeno..."
        # "Menclovi" (dativ od "Mencl") by mělo být [[PERSON_X]]

        # Vytvoř slovník: příjmení_varianta → (tag, canonical_full)
        surname_to_person = {}
        for tag, canonical_full in self.person_canonical_names.items():
            canonical_parts = canonical_full.split()
            if len(canonical_parts) >= 2:  # Has both first and last name
                surname = canonical_parts[-1]
                # Generuj všechny pádové varianty příjmení
                surname_variants = variants_for_surname(surname)
                for variant in surname_variants:
                    # Pokud varianta ještě není v mapě, přidej ji
                    if variant.lower() not in surname_to_person:
                        surname_to_person[variant.lower()] = (tag, canonical_full)

        # Pattern pro samostatné velké slovo (pravděpodobně příjmení)
        # POZOR: Musíme vyloučit slova, která jsou už anonymizovaná nebo jsou běžná slova
        standalone_word_pattern = re.compile(
            r'(?<!\w)([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]{2,})(?!\w)',
            re.UNICODE
        )

        def replace_standalone_surname(match):
            word = match.group(1)
            word_lo = word.lower()

            # Zkontroluj, jestli je to známé příjmení
            if word_lo in surname_to_person:
                tag, canonical_full = surname_to_person[word_lo]
                # Zaznamenej tuto variantu
                if word != canonical_full:
                    self.entity_map['PERSON'][canonical_full].add(word)
                return tag

            # Není známé příjmení → vrať originál
            return match.group(0)

        # Aplikuj pattern na text
        text = standalone_word_pattern.sub(replace_standalone_surname, text)

        return text

    def anonymize_entities(self, text: str) -> str:
        """Anonymizuje všechny entity (adresy, kontakty, IČO, atd.)."""

        # DŮLEŽITÉ: Pořadí je klíčové! Od nejvíce specifických po nejméně specifické
        # KRITICKÉ: Credentials (hesla, API klíče) PRVNÍ s store_value=False!

        # 1. CREDENTIALS (username / password) - NEJPRVE!
        def replace_credentials(match):
            username = match.group(2)
            password = match.group(3)
            username_tag = self._get_or_create_label('USERNAME', username)
            # TEST MODE: store_value=True (ukládá plnou hodnotu)
            password_tag = self._get_or_create_label('PASSWORD', password, store_value=True)
            return f"{match.group(1)}: {username_tag} / {password_tag}"
        text = CREDENTIALS_RE.sub(replace_credentials, text)

        # 2. HESLA (TEST MODE: store_value=True)
        def replace_password(match):
            return self._get_or_create_label('PASSWORD', match.group(1), store_value=True)
        text = PASSWORD_RE.sub(replace_password, text)

        # 2. API KLÍČE, SECRETS (TEST MODE: store_value=True)
        def replace_api_key(match):
            return self._get_or_create_label('API_KEY', match.group(1), store_value=True)
        text = API_KEY_RE.sub(replace_api_key, text)

        def replace_secret(match):
            return self._get_or_create_label('SECRET', match.group(1), store_value=True)
        text = SECRET_RE.sub(replace_secret, text)

        # 3. SSH KLÍČE (TEST MODE: store_value=True)
        def replace_ssh_key(match):
            return self._get_or_create_label('SSH_KEY', match.group(1), store_value=True)
        text = SSH_KEY_RE.sub(replace_ssh_key, text)

        # 3.5. IBAN (PŘED kartami! IBAN má dlouhé číselné sekvence)
        # TEST MODE: store_value=True (ukládá plnou hodnotu)
        def replace_iban(match):
            return self._get_or_create_label('IBAN', match.group(1), store_value=True)
        text = IBAN_RE.sub(replace_iban, text)

        # 4. PLATEBNÍ KARTY (TEST MODE: store_value=True)
        def replace_card(match):
            # CARD_RE má 1 capture group
            card = match.group(1)
            if card:
                return self._get_or_create_label('CARD', card, store_value=True)
            return match.group(0)
        text = CARD_RE.sub(replace_card, text)

        # 5. USERNAMES, ACCOUNTS, HOSTNAMES
        def replace_username(match):
            return self._get_or_create_label('USERNAME', match.group(1))
        text = USERNAME_RE.sub(replace_username, text)

        def replace_account_id(match):
            return self._get_or_create_label('ACCOUNT_ID', match.group(1))
        text = ACCOUNT_ID_RE.sub(replace_account_id, text)

        def replace_hostname(match):
            return self._get_or_create_label('HOST', match.group(1))
        text = HOSTNAME_RE.sub(replace_hostname, text)

        # 6. IP ADRESY
        def replace_ip(match):
            return self._get_or_create_label('IP', match.group(1))
        text = IP_RE.sub(replace_ip, text)

        # 7. ČÍSLA POJIŠTĚNCE
        def replace_insurance_id(match):
            return self._get_or_create_label('INSURANCE_ID', match.group(1))
        text = INSURANCE_ID_RE.sub(replace_insurance_id, text)

        # 8. RFID/BADGE
        def replace_rfid(match):
            return self._get_or_create_label('RFID', match.group(1))
        text = RFID_RE.sub(replace_rfid, text)

        # 8.1. SOCIÁLNÍ SÍTĚ (KRITICKÉ - PII)
        def replace_linkedin(match):
            return self._get_or_create_label('LINKEDIN', match.group(1))
        text = LINKEDIN_RE.sub(replace_linkedin, text)

        def replace_facebook(match):
            return self._get_or_create_label('FACEBOOK', match.group(1))
        text = FACEBOOK_RE.sub(replace_facebook, text)

        def replace_instagram(match):
            # Instagram má 2 capture groups - handle nebo URL
            handle = match.group(1) if match.group(1) else match.group(2)
            return self._get_or_create_label('INSTAGRAM', handle)
        text = INSTAGRAM_RE.sub(replace_instagram, text)

        def replace_skype(match):
            return self._get_or_create_label('SKYPE', match.group(1))
        text = SKYPE_RE.sub(replace_skype, text)

        # 8.2. BIOMETRICKÉ IDENTIFIKÁTORY (KRITICKÉ - GDPR Článek 9)
        def replace_voice_id(match):
            return self._get_or_create_label('VOICE_ID', match.group(1))
        text = VOICE_ID_RE.sub(replace_voice_id, text)

        def replace_bio_hash(match):
            # BIO_HASH_RE má 2 capture groups
            bio_hash = match.group(1) if match.group(1) else match.group(2)
            return self._get_or_create_label('BIO_HASH', bio_hash)
        text = BIO_HASH_RE.sub(replace_bio_hash, text)

        def replace_photo_id(match):
            # Photo ID má 3 capture groups
            photo_id = match.group(1) if match.group(1) else (match.group(2) if match.group(2) else match.group(3))
            return self._get_or_create_label('PHOTO_ID', photo_id)
        text = PHOTO_ID_RE.sub(replace_photo_id, text)

        def replace_api_key_enhanced(match):
            return self._get_or_create_label('API_KEY', match.group(1))
        text = API_KEY_ENHANCED_RE.sub(replace_api_key_enhanced, text)

        # 9. ADRESY (před jmény, aby "Novákova 45" nebylo osobou)
        def replace_address(match):
            matched_text = match.group(0)
            # Filter out medical/technical terms that are not addresses
            medical_terms = ['hla', 'kompatibilní', 'donor', 'recipient', 'transfuze']
            if any(term in matched_text.lower() for term in medical_terms):
                return matched_text  # Not an address, return unchanged
            return self._get_or_create_label('ADDRESS', matched_text)
        text = ADDRESS_RE.sub(replace_address, text)

        # 10. EMAILY (před ostatními, protože obsahují speciální znaky)
        def replace_email(match):
            return self._get_or_create_label('EMAIL', match.group(1))
        text = EMAIL_RE.sub(replace_email, text)

        # 11. DATUM NAROZENÍ (před BIRTH_ID a všemi daty)
        def replace_dob(match):
            return self._get_or_create_label('DATE', match.group(1))
        text = DOB_RE.sub(replace_dob, text)

        # 12. RODNÁ ČÍSLA (PŘED BANK! Jinak "Rodné číslo: 850123/1234" by se rozpadlo)
        # KRITICKÁ PRIORITA: Silný kontext ("Rodné číslo:") má přednost před bank účty
        def replace_birth_id(match):
            birth_id = match.group(1) if match.group(1) else match.group(2)
            if birth_id:
                # Odstranit lomítko pokud není přítomné
                birth_id_clean = birth_id.replace(' ', '')
                return self._get_or_create_label('BIRTH_ID', birth_id_clean)
            return match.group(0)
        text = BIRTH_ID_RE.sub(replace_birth_id, text)

        # 13. BANKOVNÍ ÚČTY (po BIRTH_ID, aby RČ nebylo zachyceno jako účet)
        # DŮLEŽITÉ: Kód banky musí být součástí tagu (BANK_RE ho zachytí pokud je přítomen)
        def replace_bank(match):
            # BANK_RE má 3 capture groups - získej první non-None
            account = match.group(1) or match.group(2) or match.group(3)
            if account:
                # TEST MODE: store_value=True (ukládá plnou hodnotu)
                return self._get_or_create_label('BANK', account, store_value=True)
            return match.group(0)
        text = BANK_RE.sub(replace_bank, text)

        # 14. ČÍSLA OP (KRITICKÉ: MUSÍ být PŘED telefony!)
        def replace_id_card(match):
            id_card = match.group(1) if match.group(1) else match.group(2)
            if id_card:
                return self._get_or_create_label('ID_CARD', id_card)
            return match.group(0)
        text = ID_CARD_RE.sub(replace_id_card, text)

        # 14.5. VARIABILNÍ SYMBOL (PŘED telefony! VS čísla nejsou telefony)

        # 14.6. KONSTANTNÍ SYMBOL

        # 14.7. SPECIFICKÝ SYMBOL

        # 14.8. ČÍSLO LÉKAŘE / LICENSE ID

        # 14.9. ČÍSLO JEDNACÍ / CASE ID

        # 14.10. SPISOVÁ ZNAČKA SOUDU

        # 14.11. ČÍSLO POJISTNÉ SMLOUVY

        # 14.12. ČÍSLO SMLOUVY (generické)

        # 15. TELEFONY (po ID_CARD a VS, ale PŘED částkami!)
        def replace_phone(match):
            # PHONE_RE má capture group (1) pro samotné číslo (bez prefixu!)
            return self._get_or_create_label('PHONE', match.group(1))
        text = PHONE_RE.sub(replace_phone, text)

        # 16. DIČ (před IČO)
        def replace_dic(match):
            return self._get_or_create_label('DIC', match.group(1))
        text = DIC_RE.sub(replace_dic, text)

        # 16.5. GENETICKÉ IDENTIFIKÁTORY (PŘED IČO! rs... nejsou IČO)
        def replace_genetic_id(match):
            return self._get_or_create_label('GENETIC_ID', match.group(1))
        text = GENETIC_ID_RE.sub(replace_genetic_id, text)

        # 16.6. DATUM NAROZENÍ
        def replace_birth_date(match):
            return self._get_or_create_label('BIRTH_DATE', match.group(1))
        text = BIRTH_DATE_RE.sub(replace_birth_date, text)

        # 16.7. ČÍSLO PASU
        def replace_passport(match):
            return self._get_or_create_label('PASSPORT', match.group(1))
        text = PASSPORT_RE.sub(replace_passport, text)

        # 16.8. ŘIDIČSKÝ PRŮKAZ
        def replace_driver_license(match):
            return self._get_or_create_label('DRIVER_LICENSE', match.group(1))
        text = DRIVER_LICENSE_RE.sub(replace_driver_license, text)

        # 16.9. BENEFITNÍ KARTY (MultiSport, Sodexo) - PII!
        def replace_benefit_card(match):
            card_id = match.group(1) if match.group(1) else match.group(2)
            if card_id:
                return self._get_or_create_label('BENEFIT_CARD', card_id)
            return match.group(0)
        text = BENEFIT_CARD_RE.sub(replace_benefit_card, text)

        # 17. IČO
        def replace_ico(match):
            full = match.group(0)
            # Ale ne pokud je to DIČ (CZ prefix)
            if 'CZ' in full.upper():
                return full
            # A ne pokud je to rodné číslo
            if '/' in full:
                return full
            # A ne pokud je to číslo OP
            if match.group(1):
                return self._get_or_create_label('ICO', match.group(1))
            return full
        text = ICO_RE.sub(replace_ico, text)

        # 18. SPZ / License Plates
        def replace_license_plate(match):
            # LICENSE_PLATE_RE má capture group (1) pro celou SPZ
            plate = match.group(1) if match.lastindex and match.lastindex >= 1 else match.group(0)

            # Filtruj false positives
            # Nesmí být jen číslice (789456, 456789)
            if plate.replace(' ', '').isdigit():
                return match.group(0)

            # Nesmí začínat "EU " (EU 2016)
            if plate.upper().startswith('EU '):
                return match.group(0)

            # Kontrola kontextu - nesmí být po "od", "z", "do", "roku"
            start_pos = match.start()
            context_before = text[max(0, start_pos-10):start_pos].lower()
            if any(word in context_before for word in ['od ', 'z ', 'do ', 'roku ']):
                return match.group(0)

            return self._get_or_create_label('LICENSE_PLATE', plate)
        text = LICENSE_PLATE_RE.sub(replace_license_plate, text)

        # 18.1. VIN (Vehicle Identification Number)
        def replace_vin(match):
            vin = match.group(1) if match.group(1) else match.group(2)
            if vin:
                return self._get_or_create_label('VIN', vin)
            return match.group(0)
        text = VIN_RE.sub(replace_vin, text)

        # 18.2. MAC ADRESA
        def replace_mac(match):
            mac = match.group(1) or match.group(2) or match.group(3)
            if mac:
                return self._get_or_create_label('MAC', mac)
            return match.group(0)
        text = MAC_RE.sub(replace_mac, text)

        # 18.3. IMEI (International Mobile Equipment Identity)
        def replace_imei(match):
            imei = match.group(1) if match.group(1) else match.group(2)
            if imei:
                return self._get_or_create_label('IMEI', imei)
            return match.group(0)
        text = IMEI_RE.sub(replace_imei, text)

        # 19. ČÁSTKY (AŽ NAKONEC! Po telefonech a všech číselných identifikátorech)

        # 20. MASKOVÁNÍ CVV/EXPIRACE u karet
        # Nahradí "CVV: 123" → "CVV: ***" a "exp: 12/26" → "exp: **/**"
        text = re.sub(r'(CVV|CVC)\s*:\s*\d{3,4}', r'\1: ***', text, flags=re.IGNORECASE)
        text = re.sub(r'exp(?:\.|\s+|iration)?\s*:?\s*\d{2}/\d{2,4}', r'exp: **/**', text, flags=re.IGNORECASE)

        # 20.5. CLEANUP biometrických identifikátorů - odstraň [[PHONE_*]] z bio prefixů
        # "IRIS_SCAN_PD_[[PHONE_10]]" → "IRIS_SCAN_PD_10"
        # "VOICE_RK_[[PHONE_11]]" → "VOICE_RK_11"
        text = re.sub(
            r'(IRIS_SCAN|VOICE_RK|HASH_BIO|FINGERPRINT|FACIAL|RETINA|PALM|DNA)_([A-Z0-9_]*)\[\[PHONE_(\d+)\]\]',
            r'\1_\2\3',
            text
        )

        # 20.6. BANK FRAGMENTS - zachyť fragmenty účtů s vloženým [[BIRTH_ID_*]]
        # "1928[[BIRTH_ID_6]]" → "[[BANK_x]]"
        # "Číslo účtu: 6677[[BIRTH_ID_21]]" → "Číslo účtu: [[BANK_x]]"
        def replace_bank_fragment(match):
            # Celý fragment (čísla před + [[BIRTH_ID_*]] + čísla po) → [[BANK_*]]
            # TEST MODE: store_value=True
            return self._get_or_create_label('BANK', match.group(1), store_value=True)

        # Pattern: číslice (volitelně) + [[BIRTH_ID_*]] + číslice (volitelně) v kontextu "účt"
        bank_fragment_pattern = re.compile(
            r'(?:číslo\s+účtu|účet|účtu|platba\s+na\s+účet|bankovní\s+účet)\s*:?\s*(\d{0,10}\[\[BIRTH_ID_\d+\]\]\d{0,10})',
            re.IGNORECASE
        )
        text = bank_fragment_pattern.sub(replace_bank_fragment, text)

        # 21. END-SCAN - finální kontrola citlivých dat (chytá zbytky nalepené na ]])
        text = self._end_scan(text)

        return text

    def _luhn_check(self, card_number: str) -> bool:
        """Validace Luhn algoritmem pro platební karty."""
        digits = [int(d) for d in card_number if d.isdigit()]
        if len(digits) < 13 or len(digits) > 19:
            return False

        checksum = 0
        for i, digit in enumerate(reversed(digits)):
            if i % 2 == 1:
                digit *= 2
                if digit > 9:
                    digit -= 9
            checksum += digit

        return checksum % 10 == 0

    def _end_scan(self, text: str) -> str:
        """Finální sken po všech náhradách - chytá případné zbytky citlivých dat."""

        # LUHN END-SCAN - zachytí všechny Luhn-validní karty (13-19 číslic)
        luhn_pattern = re.compile(r'\b(\d[\s\-]?){12,18}\d\b')
        def final_luhn_card(match):
            candidate = match.group(0)
            # Přeskoč pokud už je v [[CARD_*]]
            if '[[CARD_' in text[max(0, match.start()-10):match.end()+10]:
                return match.group(0)
            # Validuj Luhn
            if self._luhn_check(candidate):
                # TEST MODE: store_value=True
                return self._get_or_create_label('CARD', candidate, store_value=True)
            return match.group(0)
        text = luhn_pattern.sub(final_luhn_card, text)

        # IBAN (pokud unikl) - TEST MODE
        def final_iban(match):
            if not '[[IBAN_' in text[max(0, match.start()-10):match.start()+30]:
                return self._get_or_create_label('IBAN', match.group(1), store_value=True)
            return match.group(0)
        text = IBAN_RE.sub(final_iban, text)

        # Platební karty (pokud unikly nebo mají CVV/exp. datum) - TEST MODE
        def final_card(match):
            card = match.group(1) if match.group(1) else match.group(2)
            if card and not '[[CARD_' in text[max(0, match.start()-10):match.start()+len(card)+10]:
                return self._get_or_create_label('CARD', card, store_value=True)
            return match.group(0)
        text = CARD_RE.sub(final_card, text)

        # Hesla (pokud unikla nebo jsou nalepená na jiných entitách) - TEST MODE
        def final_password(match):
            return self._get_or_create_label('PASSWORD', match.group(1), store_value=True)
        text = PASSWORD_RE.sub(final_password, text)

        # IP adresy (pokud unikly)
        def final_ip(match):
            # Přeskoč pokud už je součástí URL/hostname
            if not re.search(r'[a-z]', text[max(0, match.start()-10):match.start()], re.IGNORECASE):
                return self._get_or_create_label('IP', match.group(1))
            return match.group(0)
        text = IP_RE.sub(final_ip, text)

        # Usernames (pokud unikly)
        def final_username(match):
            return self._get_or_create_label('USERNAME', match.group(1))
        text = USERNAME_RE.sub(final_username, text)

        # API klíče, secrets (pokud unikly) - TEST MODE
        def final_api(match):
            return self._get_or_create_label('API_KEY', match.group(1), store_value=True)
        text = API_KEY_RE.sub(final_api, text)

        def final_secret(match):
            return self._get_or_create_label('SECRET', match.group(1), store_value=True)
        text = SECRET_RE.sub(final_secret, text)

        # Pojištěnce (pokud unikla)
        def final_insurance(match):
            return self._get_or_create_label('INSURANCE_ID', match.group(1))
        text = INSURANCE_ID_RE.sub(final_insurance, text)

        # RFID (pokud uniklo)
        def final_rfid(match):
            return self._get_or_create_label('RFID', match.group(1))
        text = RFID_RE.sub(final_rfid, text)

        # POST-PASS: Karty v kontextu "Číslo:" které unikly hlavnímu CARD_RE
        # Hledáme specificky "Číslo: 4532..." v blízkosti slov karta/card/platební
        card_context_pattern = re.compile(
            r'(?:platební\s+karta|karta|card)[\s\S]{0,100}?'  # Lookforward pro kontext
            r'Číslo\s*:\s*'
            r'(\d{4}[\s\-]?\d{4}[\s\-]?\d{4}[\s\-]?\d{4}(?:[\s\-]?\d{1,3})?)',  # 16-19 číslic
            re.IGNORECASE
        )

        def replace_card_context(match):
            card = match.group(1)
            # Check if already tagged
            if f"[[CARD_" not in text[max(0, match.start()-20):min(len(text), match.end()+20)]:
                return match.group(0).replace(card, self._get_or_create_label('CARD', card, store_value=True))
            return match.group(0)

        text = card_context_pattern.sub(replace_card_context, text)

        # POST-PASS: ALL_CAPS jména bez diakritiky (ALENA DVORAKOVA)
        # Najdi známé osoby a nahraď jejich ALL_CAPS varianty bez diakritiky
        for p in self.canonical_persons:
            # Vytvoř ALL_CAPS verzi bez diakritiky
            import unicodedata
            def remove_dia(s):
                nfd = unicodedata.normalize('NFD', s)
                return ''.join(c for c in nfd if unicodedata.category(c) != 'Mn')

            if p['first'] and p['last']:
                first_caps = remove_dia(p['first']).upper()
                last_caps = remove_dia(p['last']).upper()
                # Nahraď pokud nalezeno
                pattern = rf'\b{re.escape(first_caps)}\s+{re.escape(last_caps)}\b'
                text = re.sub(pattern, p['tag'], text)

        # POST-PASS: Místo narození
        birth_place_pattern = re.compile(
            r'(Místo\s+narození\s*:)\s*([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)',
            re.IGNORECASE
        )
        def replace_birth_place(match):
            prefix = match.group(1)
            place = match.group(2)
            # Simple tag - don't reuse ADDRESS tags
            tag = self._get_or_create_label('BIRTH_PLACE', place)
            return f"{prefix} {tag}"
        text = birth_place_pattern.sub(replace_birth_place, text)

        # POST-PASS: Jednoduché adresy (ulice číslo, město) které unikly ADDRESS_RE
        # Pattern: Ulice 123/45, Praha 3  (bez PSČ nebo kontextu "bytem", "sídlo:" apod.)
        # ROZŠÍŘENO: Teď podporuje JAKÉKOLIV město, ne jen Praha/Brno/Ostrava/Plzeň
        simple_addr_pattern = re.compile(
            r'\b([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+(?:\s+[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)?)\s+'  # Ulice (1-2 slova)
            r'(\d+(?:/\d+)?)\s*,\s*'  # Číslo
            r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+(?:\s+\d+)?)'  # Město (jakékoliv + volitelné číslo obvodu)
            r'(?=\s|$|,|\.|;|:|\n|\r)',  # Lookahead - konec nebo interpunkce
            re.IGNORECASE
        )
        def replace_simple_addr(match):
            addr = match.group(0)
            # Přeskoč pokud už je tagovaná
            if '[[ADDRESS_' not in text[max(0, match.start()-10):min(len(text), match.end()+10)]:
                return self._get_or_create_label('ADDRESS', addr)
            return addr
        text = simple_addr_pattern.sub(replace_simple_addr, text)

        # POST-PASS: Adresy - nahraď všechny výskyty známých adres (i bez PSČ)
        # Projdi všechny adresy v entity_map a nahraď všechny jejich částečné výskyty
        if 'ADDRESS' in self.entity_map:
            for addr_key in self.entity_map['ADDRESS'].keys():
                # Získej tag
                if addr_key in self.entity_index_cache.get('ADDRESS', {}):
                    idx = self.entity_index_cache['ADDRESS'][addr_key]
                    tag = f"[[ADDRESS_{idx}]]"

                    # Generuj varianty: bez PSČ, bez čárek, atd.
                    # Např. "Karlovo náměstí 12/34, 120 00 Praha 2" → "Karlovo náměstí 12/34, Praha 2"
                    # Odstran PSČ pattern: \d{3}\s?\d{2}
                    addr_no_psc = re.sub(r',?\s*\d{3}\s?\d{2}\s*', ', ', addr_key).strip(', ')

                    # Nahraď varianty (ale pouze pokud nejsou už tagované)
                    for variant in [addr_key, addr_no_psc]:
                        if variant and len(variant) > 15:  # Min. délka
                            # Najdi a nahraď všechny výskyty kromě již tagovaných
                            if variant in text:
                                # Split by variant
                                parts = text.split(variant)
                                if len(parts) > 1:
                                    # Nahraď pouze pokud kolem není tag
                                    new_parts = []
                                    for i, part in enumerate(parts[:-1]):
                                        new_parts.append(part)
                                        # Check if not already tagged (look at end of previous part and start of next)
                                        if not (part.endswith('[[') or parts[i+1].startswith(']]')):
                                            new_parts.append(tag)
                                        else:
                                            new_parts.append(variant)  # Keep original if tagged
                                    new_parts.append(parts[-1])
                                    text = ''.join(new_parts)

        return text

    def _fix_canonical_names_not_in_document(self):
        """Fix canonical names that don't appear in the source document.

        If the inferred canonical name doesn't actually exist in the document,
        replace it with the most common variant that DOES exist.
        This prevents cases like "Danya" (not in doc) instead of "Dana" (in doc).
        """
        if not self.source_text:
            return

        fixed_count = 0
        for person in self.canonical_persons:
            canonical_full = f"{person['first']} {person['last']}"

            # Check if canonical first name appears in document
            first_in_doc = person['first'] in self.source_text

            if not first_in_doc:
                # Find the most common variant that IS in the document
                variants = self.entity_map['PERSON'].get(canonical_full, set())

                # Count occurrences of each variant's first name
                first_name_counts = {}
                for variant in variants:
                    parts = variant.split(' ', 1)
                    if len(parts) == 2:
                        first_variant = parts[0]
                        # Count occurrences in source_text
                        count = self.source_text.count(first_variant)
                        if count > 0:
                            first_name_counts[first_variant] = count

                # Pick the most common variant and INFER its nominative
                if first_name_counts:
                    best_first_variant = max(first_name_counts, key=first_name_counts.get)
                    # Infer nominative from the variant (e.g., "Dany" → "Dana")
                    best_first = infer_first_name_nominative(best_first_variant)
                    print(f"  [FIX] Canonical '{person['first']}' not in document, using '{best_first}' (inferred from '{best_first_variant}' which appears {first_name_counts[best_first_variant]}x)")
                    person['first'] = best_first

                    # Update canonical name in person_canonical_names
                    new_canonical_full = f"{best_first} {person['last']}"
                    if person['tag'] in self.person_canonical_names:
                        self.person_canonical_names[person['tag']] = new_canonical_full

                    # Update entity_map key (only if different!)
                    if canonical_full != new_canonical_full:
                        if canonical_full in self.entity_map['PERSON']:
                            variants = self.entity_map['PERSON'][canonical_full]
                            if new_canonical_full not in self.entity_map['PERSON']:
                                self.entity_map['PERSON'][new_canonical_full] = set()
                            self.entity_map['PERSON'][new_canonical_full] |= variants
                            del self.entity_map['PERSON'][canonical_full]

                    fixed_count += 1

        if fixed_count > 0:
            print(f"  [FIX] Fixed {fixed_count} canonical names not in document")

    def _fix_gender_mismatches(self):
        """Fix gender mismatches between first and last names.

        Example: "Stanislav Horáková" (male first + female last) → "Stanislava Horáková"
        """
        if not self.source_text:
            return

        fixed_count = 0
        for person in self.canonical_persons:
            first = person['first']
            last = person['last']
            canonical_full = f"{first} {last}"

            # Check gender consistency
            first_is_male = not first.lower().endswith('a')  # Rough heuristic: -a is usually female
            last_is_female = last.lower().endswith(('ová', 'á'))  # Female surname endings

            # Mismatch: male first name + female surname
            if first_is_male and last_is_female:
                # Try to find female version of first name in variants
                variants = self.entity_map['PERSON'].get(canonical_full, set())

                for variant in variants:
                    parts = variant.split(' ', 1)
                    if len(parts) == 2:
                        variant_first = parts[0]
                        # Check if variant has female first name (ends with -a)
                        if variant_first.lower().endswith('a') and variant_first.lower() != first.lower():
                            # Check if female version is in document
                            if variant_first in self.source_text or variant in self.source_text:
                                print(f"  [GENDER-FIX] '{canonical_full}' má gender mismatch!")
                                print(f"              → OPRAVUJI na '{variant_first} {last}' (z varianty '{variant}')")

                                person['first'] = variant_first

                                # Update person_canonical_names
                                new_canonical = f"{variant_first} {last}"
                                if person['tag'] in self.person_canonical_names:
                                    self.person_canonical_names[person['tag']] = new_canonical

                                # Update entity_map
                                if canonical_full != new_canonical and canonical_full in self.entity_map['PERSON']:
                                    old_variants = self.entity_map['PERSON'][canonical_full]
                                    if new_canonical not in self.entity_map['PERSON']:
                                        self.entity_map['PERSON'][new_canonical] = set()
                                    self.entity_map['PERSON'][new_canonical] |= old_variants
                                    del self.entity_map['PERSON'][canonical_full]

                                fixed_count += 1
                                break

        if fixed_count > 0:
            print(f"  [GENDER-FIX] Opraveno {fixed_count} gender mismatchů\n")

    def _deduplicate_persons(self):
        """Sloučí duplicitní osoby se stejným inferred nominativem nebo sdílenými variantami.

        Tento krok je důležitý protože různé pádové formy (Karel/Karla/Karlu)
        mohou vytvořit separátní osoby pokud nejsou správně detekované jako varianty.
        """
        from collections import defaultdict

        print(f"  [DEDUP] Starting deduplication, {len(self.canonical_persons)} persons to check")

        # PHASE 1: Group by shared variants in entity_map
        # If person A has "Radka Veverky" as variant and person B is "Radka Veverky", merge them
        to_merge = []  # List of (primary_idx, duplicate_idx) pairs

        for i, person_a in enumerate(self.canonical_persons):
            canonical_a = f"{person_a['first']} {person_a['last']}"
            variants_a = self.entity_map['PERSON'].get(canonical_a, set())

            for j, person_b in enumerate(self.canonical_persons):
                if i >= j:  # Only check each pair once
                    continue

                canonical_b = f"{person_b['first']} {person_b['last']}"

                # Check if canonical_b is in variants_a (person B's canonical is a variant of A)
                if canonical_b in variants_a:
                    print(f"  [DEDUP] Found variant overlap: '{canonical_a}' has variant '{canonical_b}'")
                    to_merge.append((i, j))
                    continue

                # Check if canonical_a is in variants_b (person A's canonical is a variant of B)
                variants_b = self.entity_map['PERSON'].get(canonical_b, set())
                if canonical_a in variants_b:
                    print(f"  [DEDUP] Found variant overlap: '{canonical_b}' has variant '{canonical_a}'")
                    to_merge.append((i, j))
                    continue

                # NEW: Check if inferred forms match (catches cases where inference doesn't work during initial detection)
                # Example: "Adéla Jarošová" (infers to itself) vs "Adéle Jarošové" (should infer to "Adéla Jarošová")
                inferred_a_first = infer_first_name_nominative(person_a['first'])
                inferred_a_last = infer_surname_nominative(person_a['last'])
                inferred_b_first = infer_first_name_nominative(person_b['first'])
                inferred_b_last = infer_surname_nominative(person_b['last'])

                # DEBUG for Daniel
                if ('daniel' in canonical_a.lower() and 'mlynář' in canonical_a.lower()) or \
                   ('daniel' in canonical_b.lower() and 'mlynář' in canonical_b.lower()):
                    print(f"  [DEBUG-DAN] Comparing:")
                    print(f"    A: '{canonical_a}' (first='{person_a['first']}', last='{person_a['last']}') -> inferred: ({inferred_a_first}, {inferred_a_last})")
                    print(f"    B: '{canonical_b}' (first='{person_b['first']}', last='{person_b['last']}') -> inferred: ({inferred_b_first}, {inferred_b_last})")

                # Normalize inferred forms
                inferred_a_first_norm = normalize_name_variant(inferred_a_first) if inferred_a_first else inferred_a_first
                inferred_b_first_norm = normalize_name_variant(inferred_b_first) if inferred_b_first else inferred_b_first

                # Compare normalized inferred forms
                key_a_first = self._normalize_for_matching(inferred_a_first_norm)
                key_a_last = self._normalize_for_matching(inferred_a_last)
                key_b_first = self._normalize_for_matching(inferred_b_first_norm)
                key_b_last = self._normalize_for_matching(inferred_b_last)

                if (key_a_first == key_b_first and key_a_last == key_b_last):
                    print(f"  [DEDUP] Found inferred match: '{canonical_a}' ({inferred_a_first_norm} {inferred_a_last}) == '{canonical_b}' ({inferred_b_first_norm} {inferred_b_last})")
                    print(f"          Keys: ({key_a_first},{key_a_last}) == ({key_b_first},{key_b_last})")
                    to_merge.append((i, j))
                    continue

        # Merge persons from Phase 1 (variant overlap)
        merged_count_phase1 = 0
        merged_indices = set()  # Track which indices have been merged

        for primary_idx, dup_idx in to_merge:
            # Skip if either has already been merged
            if primary_idx in merged_indices or dup_idx in merged_indices:
                continue

            primary = self.canonical_persons[primary_idx]
            duplicate = self.canonical_persons[dup_idx]

            primary_canonical = f"{primary['first']} {primary['last']}"
            dup_canonical = f"{duplicate['first']} {duplicate['last']}"

            # Merge entity_map variants
            if dup_canonical in self.entity_map['PERSON']:
                dup_variants = self.entity_map['PERSON'][dup_canonical]
                if primary_canonical not in self.entity_map['PERSON']:
                    self.entity_map['PERSON'][primary_canonical] = set()
                self.entity_map['PERSON'][primary_canonical] |= dup_variants
                del self.entity_map['PERSON'][dup_canonical]

            # Mark duplicate for removal (can't remove during iteration)
            merged_indices.add(dup_idx)
            merged_count_phase1 += 1

        # Remove merged persons (do this after collecting all indices)
        if merged_indices:
            self.canonical_persons = [p for i, p in enumerate(self.canonical_persons) if i not in merged_indices]
            print(f"  [DEDUP] Phase 1: Merged {merged_count_phase1} persons based on variant overlap")

        # PHASE 2: Group persons by their inferred nominative
        by_nominative = defaultdict(list)
        person_keys = {}  # person -> list of keys

        for person in self.canonical_persons:
            canonical = f"{person['first']} {person['last']}"

            # Infer nominative from canonical form
            first_nom = infer_first_name_nominative(person['first'])
            last_nom = infer_surname_nominative(person['last'])

            # Normalize for matching
            first_norm = normalize_name_variant(first_nom) if first_nom else first_nom
            key1 = (self._normalize_for_matching(first_norm), self._normalize_for_matching(last_nom))

            # For ambiguous names ending in -a (could be female OR male genitive), try BOTH
            # Example: "Radka" could be female name OR genitive of "Radko"
            keys = [key1]

            if person['first'].lower().endswith('a') and len(person['first']) > 2:
                # Try removing -a to get potential male name
                potential_male = person['first'][:-1]
                # Try with -o ending (foreign names: Radka -> Radko, Iva -> Ivo)
                if potential_male[-1].lower() not in 'aeiouyáéěíóúůý':
                    potential_male_o = potential_male + 'o'
                    key2 = (self._normalize_for_matching(potential_male_o), self._normalize_for_matching(last_nom))
                    keys.append(key2)

            # Handle dative forms ending in -ovi (Radkovi -> try both Radko and Radek)
            if person['first'].lower().endswith('ovi') and len(person['first']) > 4:
                stem = person['first'][:-3]  # Remove -ovi
                # Try with -o ending (Radkovi -> Radko)
                if stem[-1].lower() not in 'aeiouyáéěíóúůý':
                    stem_o = stem + 'o'
                    key3 = (self._normalize_for_matching(stem_o), self._normalize_for_matching(last_nom))
                    if key3 not in keys:
                        keys.append(key3)

            # ALSO generate keys from ALL variants of this person
            # This catches cases where inference doesn't work properly
            variants = self.entity_map['PERSON'].get(canonical, set())
            for variant in variants:
                # Parse variant into first and last
                variant_parts = variant.split()
                if len(variant_parts) >= 2:
                    var_first = variant_parts[0]
                    var_last = variant_parts[-1]

                    # Infer nominative from variant
                    var_first_nom = infer_first_name_nominative(var_first)
                    var_last_nom = infer_surname_nominative(var_last)

                    # Create key
                    var_first_norm = normalize_name_variant(var_first_nom) if var_first_nom else var_first_nom
                    var_key = (self._normalize_for_matching(var_first_norm), self._normalize_for_matching(var_last_nom))

                    if var_key not in keys:
                        keys.append(var_key)

            # Store all keys for this person
            person_keys[canonical] = keys
            for key in keys:
                by_nominative[key].append(person)

            # Debug: check specific duplicates
            debug_names = ['Karel Řehoř', 'Karla Řehoř', 'Radko Veverka', 'Radka Veverky', 'Radkovi Veverkovi',
                          'Lívia Králová', 'Lívií Králové', 'Marek Holý', 'Markovi Holému',
                          'Adéla Jarošová', 'Adéle Jarošové', 'Petr Dohnal', 'Petru Dohnal']
            if any(name in canonical for name in debug_names):
                print(f"  [DEDUP] {canonical} -> inferred: {first_nom} {last_nom} -> keys: {keys}")

        # Debug: print groups with duplicates
        dup_groups = [(k, g) for k, g in by_nominative.items() if len(g) > 1]
        print(f"  [DEDUP] Total unique keys: {len(by_nominative)}, Groups with duplicates: {len(dup_groups)}")
        if dup_groups:
            print(f"  [DEDUP] Found {len(dup_groups)} groups with duplicates:")
            for key, group in dup_groups[:10]:  # Show first 10
                print(f"    {key}: {[(p['first'], p['last']) for p in group]}")

        # Find and merge duplicates
        merged_count = 0
        for key, group in by_nominative.items():
            if len(group) <= 1:
                continue  # No duplicates

            # Keep first person, merge others into it
            primary = group[0]
            primary_tag = primary['tag']
            primary_canonical = f"{primary['first']} {primary['last']}"

            for duplicate in group[1:]:
                dup_tag = duplicate['tag']
                dup_canonical = f"{duplicate['first']} {duplicate['last']}"

                # Merge entity_map variants
                if dup_canonical in self.entity_map['PERSON']:
                    dup_variants = self.entity_map['PERSON'][dup_canonical]
                    if primary_canonical not in self.entity_map['PERSON']:
                        self.entity_map['PERSON'][primary_canonical] = set()
                    self.entity_map['PERSON'][primary_canonical] |= dup_variants
                    del self.entity_map['PERSON'][dup_canonical]

                # Remove duplicate from canonical_persons
                self.canonical_persons.remove(duplicate)

                # Update person_canonical_names if needed
                if dup_tag in self.person_canonical_names:
                    del self.person_canonical_names[dup_tag]

                merged_count += 1

        total_merged = merged_count_phase1 + merged_count
        if merged_count > 0:
            print(f"  [DEDUP] Phase 2: Merged {merged_count} persons based on inferred nominative")

        # PHASE 3: Fix ambiguous female/male names with male surnames
        # Example: "Radka Hofman" (female name + male surname) should be "Radek Hofman" (male)
        # This happens when "Radka" is genitiv of "Radek" but is detected as female name
        merged_count_phase3 = 0
        ambiguous_names = {'radka': 'radek', 'janka': 'janek', 'mirka': 'mirek', 'petra': 'petr'}

        persons_to_remove = []
        for person in self.canonical_persons:
            first_lo = person['first'].lower()
            last_lo = person['last'].lower()

            # Check if this is an ambiguous female name with male surname
            if first_lo in ambiguous_names and not last_lo.endswith(('ová', 'á')):
                # This might be a genitiv - check if male version exists
                male_first = ambiguous_names[first_lo].capitalize()
                male_canonical = f"{male_first} {person['last']}"

                # Find if male version exists in canonical_persons
                male_person = None
                for p in self.canonical_persons:
                    if f"{p['first']} {p['last']}" == male_canonical:
                        male_person = p
                        break

                if male_person:
                    # Merge female version into male version
                    female_canonical = f"{person['first']} {person['last']}"
                    print(f"  [DEDUP] Phase 3: Merging ambiguous '{female_canonical}' → '{male_canonical}'")

                    # Merge entity_map variants
                    if female_canonical in self.entity_map['PERSON']:
                        female_variants = self.entity_map['PERSON'][female_canonical]
                        if male_canonical not in self.entity_map['PERSON']:
                            self.entity_map['PERSON'][male_canonical] = set()
                        self.entity_map['PERSON'][male_canonical] |= female_variants
                        del self.entity_map['PERSON'][female_canonical]

                    # Mark for removal
                    persons_to_remove.append(person)

                    # Update person_canonical_names
                    if person['tag'] in self.person_canonical_names:
                        del self.person_canonical_names[person['tag']]

                    merged_count_phase3 += 1

        # Remove merged persons
        for person in persons_to_remove:
            self.canonical_persons.remove(person)

        if merged_count_phase3 > 0:
            print(f"  [DEDUP] Phase 3: Merged {merged_count_phase3} persons based on ambiguous male/female names")

        # PHASE 4: Typo/OCR error correction - rename typo canonicals to correct forms
        # Example: "Fiael" → "Fiala", "Růžiček" → "Růžička"
        corrected_count_phase4 = 0
        typo_corrections = {
            'fiael': 'fiala', 'fial': 'fiala',
            'růžiček': 'růžička', 'ruzicek': 'růžička',
            'novk': 'novák', 'dvork': 'dvořák',
            'prochzka': 'procházka', 'cern': 'černý', 'horak': 'horák',
        }

        for person in self.canonical_persons:
            last_lo = person['last'].lower()

            # Check if surname is a known typo
            if last_lo in typo_corrections:
                correct_surname = typo_corrections[last_lo]
                typo_canonical = f"{person['first']} {person['last']}"
                correct_canonical = f"{person['first']} {correct_surname.capitalize()}"

                print(f"  [DEDUP] Phase 4: Correcting typo '{typo_canonical}' → '{correct_canonical}'")

                # Update person's last name
                person['last'] = correct_surname.capitalize()

                # Update entity_map key
                if typo_canonical in self.entity_map['PERSON']:
                    variants = self.entity_map['PERSON'][typo_canonical]
                    # Move variants to correct key
                    if correct_canonical not in self.entity_map['PERSON']:
                        self.entity_map['PERSON'][correct_canonical] = set()
                    self.entity_map['PERSON'][correct_canonical] |= variants
                    del self.entity_map['PERSON'][typo_canonical]

                # Update person_canonical_names mapping
                if person['tag'] in self.person_canonical_names:
                    self.person_canonical_names[person['tag']] = correct_canonical

                corrected_count_phase4 += 1

        if corrected_count_phase4 > 0:
            print(f"  [DEDUP] Phase 4: Corrected {corrected_count_phase4} typo surnames")

        total_merged = merged_count_phase1 + merged_count + merged_count_phase3
        if total_merged > 0:
            print(f"  [DEBUG] Total merged: {total_merged} duplicate persons")

    def anonymize_docx(self, input_path: str, output_path: str, json_map: str, txt_map: str):
        """Hlavní metoda pro anonymizaci DOCX dokumentu."""
        print(f"\n🔍 Zpracovávám: {Path(input_path).name}")

        # Načti dokument
        import time
        start_time = time.time()
        doc = Document(input_path)
        print(f"  [DEBUG] Document loaded in {time.time() - start_time:.1f}s")

        # Store source text for validation (check if inferred names are in document)
        self.source_text = ' '.join([p.text for p in doc.paragraphs])

        # Zpracuj všechny odstavce
        start_time = time.time()
        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            original = para.text

            # POŘADÍ JE KLÍČOVÉ!
            # 1. Nejprve anonymizuj entity (adresy, IČO, telefony...)
            text = self.anonymize_entities(original)

            # 2. Potom aplikuj známé osoby
            text = self._apply_known_people(text)

            # 3. Nakonec detekuj nové osoby
            text = self._replace_remaining_people(text)

            if text != original:
                para.text = text

        # POST-PROCESSING: Fix canonical names that are not in source document
        self._fix_canonical_names_not_in_document()

        # POST-PROCESSING: Fix gender mismatches (e.g., Stanislav Horáková → Stanislava Horáková)
        self._fix_gender_mismatches()

        print(f"  [DEBUG] Paragraphs processed in {time.time() - start_time:.1f}s")

        # Zpracuj tabulky
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.text.strip():
                            continue

                        original = para.text
                        text = self.anonymize_entities(original)
                        text = self._apply_known_people(text)
                        text = self._replace_remaining_people(text)

                        if text != original:
                            para.text = text

        # POST-PROCESSING: Deduplicate persons AFTER all extraction (including tables)
        self._deduplicate_persons()

        # Ulož dokument
        start_time = time.time()
        doc.save(output_path)
        print(f"  [DEBUG] Document saved in {time.time() - start_time:.1f}s")

        # Vytvoř mapy (předej doc a path aby se nečetl znovu)
        start_time = time.time()
        self._create_maps(json_map, txt_map, input_path, doc)
        print(f"  [DEBUG] Maps created in {time.time() - start_time:.1f}s")

        print(f"✅ Hotovo! Nalezeno {len(self.canonical_persons)} osob")

    def _create_maps(self, json_path: str, txt_path: str, source_file: str, doc=None):
        """Vytvoří JSON a TXT mapy náhrad."""

        # Cleanup nepoužitých tagů před vytvořením map
        # Použij předaný dokument nebo načti z disku
        if doc is None:
            from docx import Document
            doc = Document(txt_path.replace('_map.txt', '_anon.docx'))

        anon_text = "\n".join([p.text for p in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    anon_text += "\n".join([p.text for p in cell.paragraphs])

        # SKIP cleanup - not needed in TEST MODE and causes issues with index mapping
        # Všechny entity zůstanou v mapě i když nejsou použity v textu

        # JSON mapa
        json_data = {
            "version": "1.0",
            "generated_at": datetime.now().isoformat(),
            "source_file": Path(source_file).name,
            "entities": []
        }

        # VALIDACE: Načti zdrojový dokument pro kontrolu existence entit
        from docx import Document as DocxDocument
        source_doc = DocxDocument(source_file)
        source_text = '\n'.join([p.text for p in source_doc.paragraphs])
        # Přidej text z tabulek
        for table in source_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    source_text += '\n' + '\n'.join([p.text for p in cell.paragraphs])


        # ========== AUTOMATICKÁ OPRAVA: Kanonická jména musí mít varianty ve smlouvě! ==========
        print("\n🔍 AUTO-OPRAVA: Kontroluji a opravuji kanonická jména...")

        persons_to_delete = []  # Persons that are completely invalid
        fixed_count = 0

        for i, person in enumerate(self.canonical_persons):
            canonical_first = person['first']
            canonical_last = person['last']
            canonical_full = f"{canonical_first} {canonical_last}"

            # Získej všechny varianty této osoby
            variants = self.entity_map['PERSON'].get(canonical_full, set())

            # Kontrola: Je kanonické jméno NEBO alespoň jedna varianta ve smlouvě?
            found_in_doc = canonical_full in source_text or canonical_first in source_text

            # Pokud kanonické není, zkontroluj varianty (celé jméno i křestní jméno)
            if not found_in_doc and variants:
                # Find which variant is actually in the document
                found_variant = None
                for variant in variants:
                    if variant in source_text:
                        found_variant = variant
                        break
                    # Also check just the first name
                    variant_first = variant.split(' ', 1)[0] if ' ' in variant else variant
                    if variant_first in source_text:
                        found_variant = variant
                        break

                if found_variant:
                    found_in_doc = True

            # Decision: Fix or Delete?
            if not found_in_doc:
                if variants:
                    # Has variants but none in doc → Try to fix from most common variant
                    print(f"  ⚠️  '{canonical_full}' není ve smlouvě, ale má varianty: {variants}")
                    print(f"      → Pokouším se opravit z nejčastější varianty...")

                    # Count occurrences of each variant
                    variant_counts = {}
                    for variant in variants:
                        count = source_text.count(variant)
                        if count > 0:
                            variant_counts[variant] = count

                    if variant_counts:
                        # Pick most common variant and infer nominative
                        best_variant = max(variant_counts, key=variant_counts.get)
                        parts = best_variant.split(' ', 1)
                        if len(parts) == 2:
                            corrected_first = infer_first_name_nominative(parts[0])
                            corrected_last = infer_surname_nominative(parts[1])
                            print(f"      ✅ OPRAVENO: '{canonical_full}' → '{corrected_first} {corrected_last}' (z varianty '{best_variant}')")

                            person['first'] = corrected_first
                            person['last'] = corrected_last

                            # Update person_canonical_names
                            new_canonical = f"{corrected_first} {corrected_last}"
                            if person['tag'] in self.person_canonical_names:
                                self.person_canonical_names[person['tag']] = new_canonical

                            # Update entity_map
                            if canonical_full != new_canonical and canonical_full in self.entity_map['PERSON']:
                                old_variants = self.entity_map['PERSON'][canonical_full]
                                if new_canonical not in self.entity_map['PERSON']:
                                    self.entity_map['PERSON'][new_canonical] = set()
                                self.entity_map['PERSON'][new_canonical] |= old_variants
                                del self.entity_map['PERSON'][canonical_full]

                            fixed_count += 1
                        else:
                            print(f"      ❌ Nelze opravit - chybný formát varianty")
                            persons_to_delete.append(i)
                    else:
                        # No variant is in document → completely invalid
                        print(f"  ❌ '{canonical_full}' - žádná varianta ve smlouvě → MAŽU!")
                        persons_to_delete.append(i)
                else:
                    # No variants at all → completely made up
                    print(f"  ❌ '{canonical_full}' - žádné varianty, vymyšlená osoba → MAŽU!")
                    persons_to_delete.append(i)

        # Delete invalid persons (reverse order to preserve indices)
        for i in reversed(persons_to_delete):
            person = self.canonical_persons[i]
            canonical_full = f"{person['first']} {person['last']}"
            tag = person['tag']

            print(f"  🗑️  Mažu neplatnou osobu: {canonical_full} ({tag})")

            # Remove from canonical_persons
            del self.canonical_persons[i]

            # Remove from entity_map
            if canonical_full in self.entity_map['PERSON']:
                del self.entity_map['PERSON'][canonical_full]

            # Remove from person_canonical_names
            if tag in self.person_canonical_names:
                del self.person_canonical_names[tag]

            # Remove from person_index
            for key, value in list(self.person_index.items()):
                if value == tag:
                    del self.person_index[key]

        if fixed_count > 0 or persons_to_delete:
            print(f"\n  ✅ AUTO-OPRAVA dokončena:")
            if fixed_count > 0:
                print(f"     - Opraveno: {fixed_count} osob")
            if persons_to_delete:
                print(f"     - Smazáno: {len(persons_to_delete)} neplatných osob")
            print()
        else:
            print("  ✅ Všechna kanonická jména jsou v pořádku!\n")

        # Osoby - ukládáme VŠECHNY původní formy z dokumentu
        for p in self.canonical_persons:
            canonical_full = f'{p["first"]} {p["last"]}'

            # Získej všechny původní formy z entity_map
            original_forms = self.entity_map['PERSON'].get(canonical_full, {canonical_full})

            # Pro každou původní formu vytvoř samostatný záznam
            # ALE POUZE pokud existuje ve zdrojovém dokumentu!
            for original_form in original_forms:
                if original_form in source_text:
                    json_data["entities"].append({
                        "type": "PERSON",
                        "label": p['tag'],
                        "original": original_form,
                        "occurrences": 1
                    })

        # Ostatní entity (kromě PERSON, který už je v canonical_persons)
        for typ, entities in self.entity_map.items():
            if typ == 'PERSON':
                continue  # Skip PERSON - already handled in canonical_persons
            for idx, (original, variants) in enumerate(entities.items(), 1):
                # VALIDACE: Přidej jen entity které existují ve zdrojovém dokumentu
                if original in source_text:
                    json_data["entities"].append({
                        "type": typ,
                        "label": f"[[{typ}_{idx}]]",
                        "original": original,
                        "occurrences": len(variants)
                    })

        # Ulož JSON
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)

        # TXT mapa
        with open(txt_path, 'w', encoding='utf-8') as f:
            # Osoby
            if self.canonical_persons:
                f.write("OSOBY\n")
                f.write("-----\n")
                for p in self.canonical_persons:
                    canonical_full = f'{p["first"]} {p["last"]}'
                    f.write(f"{p['tag']}: {canonical_full}\n")

                    # Vypsat všechny nalezené varianty této osoby
                    if canonical_full in self.entity_map['PERSON']:
                        variants = self.entity_map['PERSON'][canonical_full]
                        # Vypsat pouze varianty odlišné od kanonického tvaru
                        for variant in sorted(variants):
                            if variant.lower() != canonical_full.lower():
                                f.write(f"  - {variant}\n")
                f.write("\n")

            # Ostatní entity (kromě PERSON, který už je v OSOBY)
            for typ, entities in sorted(self.entity_map.items()):
                if typ == 'PERSON':
                    continue  # Skip PERSON - already handled in OSOBY section
                if entities:
                    f.write(f"{typ}\n")
                    for idx, (original, variants) in enumerate(entities.items(), 1):
                        label = f"[[{typ}_{idx}]]"
                        # Pro citlivá data zobraz jen ***REDACTED*** bez čísla
                        display_value = "***REDACTED***" if original.startswith("***REDACTED_") else original
                        f.write(f"{label}: {display_value}\n")
                    f.write("\n")

# =============== Batch processing ===============
def batch_anonymize(folder_path, names_json="cz_names.v1.json"):
    """Zpracuje všechny DOCX soubory v adresáři."""
    folder = Path(folder_path)
    docx_files = sorted([f for f in folder.glob("*.docx") if not f.name.startswith('~') and '_anon' not in f.name])

    if not docx_files:
        print(f"Nebyly nalezeny žádné .docx soubory v adresáři {folder_path}")
        return

    print(f"\n📁 Zpracovávám {len(docx_files)} souborů v adresáři {folder_path}\n")

    global CZECH_FIRST_NAMES
    CZECH_FIRST_NAMES = load_names_library(names_json)

    for path in docx_files:
        print(f"\n{'='*60}")
        base = path.stem
        out_docx = path.parent / f"{base}_anon.docx"
        out_json = path.parent / f"{base}_map.json"
        out_txt = path.parent / f"{base}_map.txt"

        try:
            a = Anonymizer(verbose=False)
            a.anonymize_docx(str(path), str(out_docx), str(out_json), str(out_txt))
            print(f"✅ Výstupy: {out_docx.name}, {out_json.name}, {out_txt.name}")
        except Exception as e:
            print(f"❌ CHYBA při zpracování {path.name}: {e}")
            import traceback
            traceback.print_exc()

# =============== Main ===============
def main():
    import argparse
    ap = argparse.ArgumentParser(description="Anonymizace českých DOCX s JSON knihovnou jmen")
    ap.add_argument("docx_path", nargs='?', help="Cesta k .docx souboru nebo adresáři")
    ap.add_argument("--names-json", default="cz_names.v1.json", help="Cesta k JSON knihovně jmen")
    ap.add_argument("--batch", action="store_true", help="Zpracovat všechny .docx v adresáři")
    args = ap.parse_args()

    try:
        global CZECH_FIRST_NAMES
        CZECH_FIRST_NAMES = load_names_library(args.names_json)

        if args.batch and args.docx_path:
            batch_anonymize(args.docx_path, args.names_json)
            return 0

        if args.batch and not args.docx_path:
            # Batch mode v aktuálním adresáři
            batch_anonymize(".", args.names_json)
            return 0

        # Single file mode
        if not args.docx_path:
            print("❌ Chybí cesta k souboru. Použij: python script.py <soubor.docx>")
            print("   Nebo: python script.py --batch <adresář>")
            return 2

        path = Path(args.docx_path)
        if not path.exists():
            print(f"❌ Soubor nenalezen: {path}")
            return 2

        base = path.stem
        out_docx = path.parent / f"{base}_anon.docx"
        out_json = path.parent / f"{base}_map.json"
        out_txt = path.parent / f"{base}_map.txt"

        # Kontrola zamčených souborů
        files_locked = False
        for out_file in [out_docx, out_json, out_txt]:
            if out_file.exists():
                try:
                    with open(out_file, 'a'):
                        pass
                except PermissionError:
                    files_locked = True
                    break

        if files_locked:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            out_docx = path.parent / f"{base}_anon_{timestamp}.docx"
            out_json = path.parent / f"{base}_map_{timestamp}.json"
            out_txt = path.parent / f"{base}_map_{timestamp}.txt"
            print(f"\n⚠️  Výstupní soubory jsou otevřené v jiné aplikaci!")
            print(f"   Vytvářím nové soubory s časovým razítkem: {timestamp}\n")

        a = Anonymizer(verbose=False)
        a.anonymize_docx(str(path), str(out_docx), str(out_json), str(out_txt))

        print(f"\n✅ Výstupy:")
        print(f" - {out_docx}")
        print(f" - {out_json}")
        print(f" - {out_txt}")
        print(f"\n📊 Statistiky:")
        print(f" - Nalezeno osob: {len(a.canonical_persons)}")
        print(f" - Celkem entit: {sum(len(e) for e in a.entity_map.values())}")

        return 0

    except Exception as e:
        print(f"\n❌ CHYBA: {e}")
        import traceback
        traceback.print_exc()
        return 1

if __name__ == "__main__":
    sys.exit(main())