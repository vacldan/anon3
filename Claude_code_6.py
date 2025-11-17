# -*- coding: utf-8 -*-
"""
Czech DOCX Anonymizer – Complete v7.0
- Načítá jména z JSON knihovny (cz_names.v1.json)
- Kompletní anonymizace podle GDPR
- Vylepšená detekce adres, osob, kontaktů
Výstupy: <basename>_anon.docx / _map.json / _map.txt
"""

import sys, re, json, unicodedata
from typing import Optional, Set
from pathlib import Path
from collections import defaultdict, OrderedDict
from docx import Document
from datetime import datetime

# =============== Globální proměnné ===============
CZECH_FIRST_NAMES = set()

# =============== Načítání knihovny jmen ===============
def load_names_library(json_path: str = "cz_names.v1.json") -> Set[str]:
    """Načte česká jména z JSON souboru."""
    try:
        script_dir = Path(__file__).parent if '__file__' in globals() else Path.cwd()
        json_file = script_dir / json_path

        if not json_file.exists():
            json_file = Path.cwd() / json_path

        if not json_file.exists():
            print(f"⚠️  Varování: {json_path} nenalezen, používám prázdnou knihovnu!")
            return set()

        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
            names = set()
            if isinstance(data, dict):
                # Nová struktura: {"firstnames": {"M": [...], "F": [...], "U": [...]}}
                if 'firstnames' in data:
                    firstnames = data['firstnames']
                    if isinstance(firstnames, dict):
                        for gender_key in ['M', 'F', 'U']:
                            if gender_key in firstnames:
                                names.update(firstnames[gender_key])
                # Stará struktura: {"male": [...], "female": [...]}
                else:
                    names.update(data.get('male', []))
                    names.update(data.get('female', []))
            elif isinstance(data, list):
                names.update(data)

            # Převod na lowercase pro jednodušší porovnávání
            names = {name.lower() for name in names}
            print(f"✓ Načteno {len(names)} jmen z knihovny")
            return names
    except Exception as e:
        print(f"⚠️  Chyba při načítání {json_path}: {e}")
        return set()

# =============== Varianty pro nahrazování ===============
def variants_for_first(first: str) -> set:
    """Generuje všechny pádové varianty křestního jména."""
    f = first.strip()
    if not f: return {''}

    V = {f, f.lower(), f.capitalize()}
    low = f.lower()

    # Ženská jména na -a
    if low.endswith('a'):
        stem = f[:-1]
        # 7 pádů: nominativ, genitiv, dativ, akuzativ, vokativ, lokál, instrumentál
        V |= {stem+'y', stem+'e', stem+'ě', stem+'u', stem+'ou', stem+'o'}
        # Přivlastňovací: Janin, Petřina
        V |= {stem+s for s in ['in','ina','iny','iné','inu','inou','iným','iných']}
        # Speciální případy
        if stem.endswith('tr'):
            V |= {stem[:-1]+'ř'+s for s in ['in','ina','iny','iné','inu','inou']}
    else:
        # Mužská jména
        V |= {f+'a', f+'ovi', f+'e', f+'em', f+'u', f+'om'}
        # Přivlastňovací: Petrův, Pavlův
        V |= {f+'ův'} | {f+'ov'+s for s in ['a','o','y','ě','ým','ých']}
        # Speciální případy
        if low.endswith('ek'): V.add(f[:-2] + 'ka')
        if low.endswith('el'): V.add(f[:-2] + 'la')
        if low.endswith('ec'): V.add(f[:-2] + 'ce')

    # Bez diakritiky
    V |= {unicodedata.normalize('NFKD', v).encode('ascii','ignore').decode('ascii') for v in list(V)}
    return V

def variants_for_surname(surname: str) -> set:
    """Generuje všechny pádové varianty příjmení."""
    s = surname.strip()
    if not s: return {''}

    out = {s, s.lower(), s.capitalize()}
    low = s.lower()

    # Ženská příjmení na -ová
    if low.endswith('ová'):
        base = s[:-1]
        out |= {s, base+'é', base+'ou'}
        return out

    # Přídavná jména -ský, -cký, -ý
    if low.endswith(('ský','cký','ý')):
        if low.endswith(('ský','cký')):
            stem = s[:-2]
        else:
            stem = s[:-1]
        out |= {stem+'ý', stem+'ého', stem+'ému', stem+'ým', stem+'ém'}
        out |= {stem+'á', stem+'é', stem+'ou'}
        return out

    # Ženská na -á
    if low.endswith('á'):
        stem = s[:-1]
        out |= {s, stem+'é', stem+'ou'}
        return out

    # Speciální případy
    if low.endswith('ek') and len(s) >= 3:
        stem_k = s[:-2] + 'k'
        out |= {s, stem_k+'a', stem_k+'ovi', stem_k+'em', stem_k+'u', stem_k+'e'}
        return out

    if low.endswith('el') and len(s) >= 3:
        stem_l = s[:-2] + 'l'
        out |= {s, stem_l+'a', stem_l+'ovi', stem_l+'em', stem_l+'u'}
        return out

    if low.endswith('ec') and len(s) >= 3:
        stem_c = s[:-2] + 'c'
        out |= {s, stem_c+'e', stem_c+'i', stem_c+'em', stem_c+'u'}
        return out

    # Standardní mužská příjmení
    out |= {s+'a', s+'ovi', s+'e', s+'em', s+'u', s+'ům', s+'em'}
    # Množné číslo: u Nováků
    out |= {s+'ů', s+'ům'}

    return out

# =============== Inference funkcí ===============
def _male_genitive_to_nominative(obs: str) -> Optional[str]:
    """Převede pozorovaný tvar (např. genitiv) na nominativ pro mužská jména."""
    lo = obs.lower()
    cands = []

    # Speciální případy: -ka → -ek, -la → -el
    if lo.endswith('ka') and len(obs) > 2:
        cands.append(obs[:-2] + 'ek')
    if lo.endswith('la') and len(obs) > 2:
        cands.append(obs[:-2] + 'el')
    if lo.endswith('ce') and len(obs) > 2:
        cands.append(obs[:-2] + 'ec')

    # Genitiv/Dativ: -a → remove
    if lo.endswith('a') and len(obs) > 1:
        cands.append(obs[:-1])

    # Dativ: -ovi → remove
    if lo.endswith('ovi') and len(obs) > 3:
        cands.append(obs[:-3])

    # Instrumentál: -em → remove
    if lo.endswith('em') and len(obs) > 2:
        cands.append(obs[:-2])

    for c in cands:
        if c.lower() in CZECH_FIRST_NAMES:
            return c.capitalize()

    return cands[0].capitalize() if cands else None

def infer_first_name_nominative(obs: str) -> str:
    """Odhadne nominativ křestního jména z pozorovaného tvaru."""
    lo = obs.lower()

    # DŮLEŽITÉ: Kontrola, zda už je v nominativu (v knihovně jmen)
    if lo in CZECH_FIRST_NAMES:
        return obs.capitalize()

    # Speciální případy - zkrácená jména (Han → Hana, Mart → Marta, Martin → Martina)
    # Priorita: nejdřív zkus +ina (pro Martin → Martina), pak +a
    if lo + 'ina' in CZECH_FIRST_NAMES:
        return (obs + 'ina').capitalize()
    if lo + 'a' in CZECH_FIRST_NAMES:
        return (obs + 'a').capitalize()

    # Ženská jména - pádové varianty
    if lo.endswith(('y', 'ě', 'e', 'u', 'ou')):
        # Zkus -a variantu
        stem = obs[:-1] if not lo.endswith('ou') else obs[:-2]
        if (stem + 'a').lower() in CZECH_FIRST_NAMES:
            return (stem + 'a').capitalize()

    # Mužská jména - genitiv/dativ/instrumentál
    male_nom = _male_genitive_to_nominative(obs)
    if male_nom:
        return male_nom

    # Pokud nic nepomohlo, vrať původní tvar s velkým písmenem
    return obs.capitalize()

def infer_surname_nominative(obs: str) -> str:
    """Odhadne nominativ příjmení z pozorovaného tvaru."""
    lo = obs.lower()

    # Ženská příjmení -ové, -ou → -ová
    if lo.endswith('é') and len(obs) > 3:
        return obs[:-1] + 'á'
    if lo.endswith('ou') and len(obs) > 3:
        return obs[:-2] + 'á'

    # Přídavná jména
    if lo.endswith(('ého', 'ému', 'ým', 'ém')):
        if lo.endswith('ého'):
            return obs[:-3] + 'ý'
        elif lo.endswith('ému'):
            return obs[:-3] + 'ý'
        elif lo.endswith('ým'):
            return obs[:-2] + 'ý'
        elif lo.endswith('ém'):
            return obs[:-2] + 'ý'

    # Speciální -ka, -la, -ce → -ek, -el, -ec (ale ne běžná příjmení!)
    common_surnames_a = {'svoboda', 'skála', 'hora', 'kula', 'hala'}
    if lo.endswith('ka') and len(obs) > 3 and lo not in common_surnames_a:
        return obs[:-2] + 'ek'
    if lo.endswith('la') and len(obs) > 3 and lo not in common_surnames_a:
        return obs[:-2] + 'el'
    if lo.endswith('ce') and len(obs) > 3:
        return obs[:-2] + 'ec'

    # Dativ -ovi → remove (ale jen pokud je to opravdu dativ, ne součást jména)
    if lo.endswith('ovi') and len(obs) > 5:
        return obs[:-3]

    # Instrumentál -em → remove (ale jen pokud je to opravdu instrumentál)
    if lo.endswith('em') and len(obs) > 4 and not lo.endswith(('em', 'lem', 'rem')):
        return obs[:-2]

    # Genitiv -a → NEODSTRAŇUJ! Mnoho příjmení končí na -a v nominativu (Svoboda, Skála, atd.)
    # Tato pravidla jsou příliš riskantní

    return obs

# =============== Regexy ===============

# Vylepšený ADDRESS_RE - zachytává adresy i bez prefixů
ADDRESS_RE = re.compile(
    r'(?<!\[)'
    r'(?:'
    r'(?:(?:trvale\s+)?bytem\s+|'
    r'(?:trvalé\s+)?bydlišt[eě]\s*:\s*|'
    r'(?:sídlo(?:\s+podnikání)?|se\s+sídlem)\s*:\s*|'
    r'(?:místo\s+podnikání)\s*:\s*|'
    r'(?:adresa|trvalý\s+pobyt)\s*:\s*|'
    r'(?:v\s+ulic[ií]|na\s+(?:adrese|ulici)|v\s+dom[eě])\s+)?'
    r')'
    r'[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ]'
    r'[a-záčďéěíňóřšťúůýž\s]{2,50}?'
    r'\s+\d{1,4}(?:/\d{1,4})?'
    r',\s*'
    r'\d{3}\s?\d{2}'
    r'\s+'
    r'[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž\s]{1,30}'
    r'(?:\s+\d{1,2})?'
    r'(?=\s|$|,|\.|;|:|\n|\r|Rodné|IČO|DIČ|Tel|E-mail|Kontakt|OP|Datum|Narozen)',
    re.IGNORECASE | re.UNICODE
)

# SPZ/RZ
LICENSE_PLATE_RE = re.compile(
    r'\b\d[A-Z]{2}\s?\d{4}\b',
    re.IGNORECASE
)

# IČO (8 číslic)
ICO_RE = re.compile(
    r'(?:IČO?\s*:?\s*)?(?<!\d)(\d{8})(?!\d)',
    re.IGNORECASE
)

# DIČ (CZ + 8-10 číslic)
DIC_RE = re.compile(
    r'\b(CZ\d{8,10})\b',
    re.IGNORECASE
)

# Rodné číslo (6 číslic / 3-4 číslice)
# DŮLEŽITÉ: Musí mít kontext (RČ, Rodné číslo, nar.) nebo negativní lookahead pro spisové značky
BIRTH_ID_RE = re.compile(
    r'(?:'
    r'(?:RČ|Rodné\s+číslo|nar\.|Narození)\s*:?\s*(\d{6}/\d{3,4})|'  # S kontextem
    r'(?<!FÚ-)(?<!KS-)(?<!VS-)(\d{6}/\d{3,4})'  # Bez kontextu, ale ne po FÚ-/KS-/VS-
    r')',
    re.IGNORECASE
)

# Číslo OP (formát: AB 123456 nebo OP: 123456789)
# DŮLEŽITÉ: Musí být před PHONE_RE!
ID_CARD_RE = re.compile(
    r'(?:'
    r'\b([A-Z]{2}\s?\d{6})\b|'  # Standardní formát: AB 123456
    r'(?:OP|pas|pas\.|pas\.č\.|č\.OP)\s*[:\-]?\s*(\d{6,9})'  # OP: 123456789 nebo Pas: 123456
    r')',
    re.IGNORECASE
)

# Email
EMAIL_RE = re.compile(
    r'\b([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b'
)

# Telefon (CZ formáty) - NESMÍ zachytit prefix do hodnoty!
# Prefix je mimo capture group, telefon je uvnitř
# DŮLEŽITÉ: Whitelist - NEchytej biometrické/technické prefixy!
PHONE_RE = re.compile(
    r'(?!'  # Negative lookahead - NEchytej pokud předchází:
    r'(?:IRIS_SCAN|VOICE_RK|HASH_BIO|FINGERPRINT|FACIAL_|RETINA_|PALM_|DNA_)_[A-Z0-9_]*'
    r')'
    r'(?:tel\.?|telefon|mobil|GSM)?\s*:?\s*'  # Volitelný prefix (MIMO capture group!)
    r'('  # START capture group - jen samotné číslo
    r'\+420\s?\d{3}\s?\d{3}\s?\d{3}|'  # +420 xxx xxx xxx
    r'\+420\s?\d{3}\s?\d{2}\s?\d{2}\s?\d{2}|'  # +420 xxx xx xx xx
    r'(?<!\d)\d{3}\s?\d{3}\s?\d{3}(?!\d)|'  # xxx xxx xxx (bez okolních číslic)
    r'(?<!\d)\d{3}\s?\d{2}\s?\d{2}\s?\d{2}(?!\d)'  # xxx xx xx xx
    r')',  # END capture group
    re.IGNORECASE
)

# Bankovní účet (formát: číslo/kód banky)
# DŮLEŽITÉ: IBAN je samostatný regex níže!
BANK_RE = re.compile(
    r'\b(\d{6,16}/\d{4})\b',
    re.IGNORECASE
)

# IBAN (mezinárodní formát bankovního účtu)
# CZ IBAN: CZ + 2 číslice + 20 číslic = 24 znaků celkem
# KRITICKÉ: Musí být před CARD_RE, protože obsahuje dlouhé sekvence číslic!
IBAN_RE = re.compile(
    r'\b(CZ\d{2}(?:\s?\d{4}){5})\b',
    re.IGNORECASE
)

# Datum (DD.MM.YYYY nebo DD. MM. YYYY)
DATE_RE = re.compile(
    r'\b(\d{1,2}\.\s?\d{1,2}\.\s?\d{4})\b'
)

# Datum slovně (např. "15. března 2024")
DATE_WORDS_RE = re.compile(
    r'\b(\d{1,2}\.\s?(?:ledna|února|března|dubna|května|června|července|srpna|září|října|listopadu|prosince)\s?\d{4})\b',
    re.IGNORECASE
)

# Hesla a credentials (KRITICKÉ - hodnotu neukládat!)
PASSWORD_RE = re.compile(
    r'(?:password|heslo|passwd|pwd|pass)\s*[:\-=]\s*([^\s,;\.]{3,50})',
    re.IGNORECASE
)

# Credentials pattern - "Credentials: username / password"
CREDENTIALS_RE = re.compile(
    r'(?i)\b(credentials?|login|přihlašovací\s+údaje)\s*:\s*([A-Za-z0-9._\-@]+)\s*/\s*(\S+)',
    re.IGNORECASE
)

# API klíče, Secrets, Tokens (KRITICKÉ - hodnotu neukládat!)
API_KEY_RE = re.compile(
    r'(?:AWS\s+)?(?:Access\s+)?(?:Key(?:\s+ID)?|Secret(?:\s+Access\s+Key)?|Token|API[_\s]?Key)\s*[:\-=]\s*([A-Za-z0-9+/=_\-]{16,})',
    re.IGNORECASE
)

SECRET_RE = re.compile(
    r'(?:Stripe|SendGrid|GitHub|Secret|Client[_\s]?Secret|Access\s+Token|Personal\s+Access\s+Token)\s*[:\-=]\s*([A-Za-z0-9+/=_\-]{16,})',
    re.IGNORECASE
)

SSH_KEY_RE = re.compile(
    r'(?:ssh-rsa|ssh-ed25519|ecdsa-sha2-nistp256)\s+([A-Za-z0-9+/=]{50,})',
    re.IGNORECASE
)

# Usernames, Account IDs, Hostnames
USERNAME_RE = re.compile(
    r'(?:Login|Username|Uživatel|User|GitHub|Jira|AWS\s+Console)\s*[:\-=]\s*([A-Za-z0-9._\-@]+)',
    re.IGNORECASE
)

ACCOUNT_ID_RE = re.compile(
    r'(?:Account\s+ID|AWS\s+Account)\s*[:\-=]\s*(\d{12})',
    re.IGNORECASE
)

HOSTNAME_RE = re.compile(
    r'(?:hostname|host|server|RDS|endpoint)\s*[:\-=]\s*([a-z0-9\-\.]+\.[a-z]{2,})',
    re.IGNORECASE
)

# IP adresy (IPv4)
IP_RE = re.compile(
    r'\b(?<!\d\.)(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})(?!\.\d)\b'
)

# Platební karty (13-19 číslic, Luhn-compatible)
# Rozšířený pattern: Visa/MC (16), AmEx (15), Diners (14), atd.
# DŮLEŽITÉ: IBAN se zpracovává PŘED tímto regexem!
CARD_RE = re.compile(
    r'(?:'
    # S prefixem: "Karta 1:", "Platební karta:", "Číslo karty:", etc.
    r'(?:Platební\s+)?(?:Karta|Card)(?:\s+\d+)?(?:\s+Number)?\s*[:\-=]?\s*'
    r'('
    r'\d{4}[\s\-]?\d{6}[\s\-]?\d{5}|'  # AmEx: 4-6-5 (15 číslic)
    r'\d{4}[\s\-]?\d{4}[\s\-]?\d{4}[\s\-]?\d{4}(?:[\s\-]?\d{2,3})?|'  # Visa/MC: 16-19
    r'\d{4}[\s\-]?\d{6}[\s\-]?\d{4}'  # Diners: 4-6-4 (14 číslic)
    r')'
    r')',
    re.IGNORECASE
)

# Čísla pojištěnce
INSURANCE_ID_RE = re.compile(
    r'(?:Číslo\s+pojištěnce|Pojišťovna|VZP|ČPZP|ZPŠ|OZP)[,\s:]+(?:číslo\s*:?\s*)?(\d{10})',
    re.IGNORECASE
)

# RFID/Badge čísla
RFID_RE = re.compile(
    r'(?:RFID\s+karta|badge|ID\s+karta)\s*[:#]?\s*([A-Za-z0-9\-_/]+)',
    re.IGNORECASE
)

# Částky (Kč, EUR, USD) - NESMÍ chytat telefony!
# KLÍČOVÉ: xxx xxx xxx může být telefon, takže vyžadujeme kontext nebo měnu
AMOUNT_RE = re.compile(
    r'(?:'
    # Částky s měnou: 1 234 Kč, 50 000 EUR (vyloučit phone/+ prefixy)
    r'(?<![\+])\b(\d{1,3}(?:\s+\d{3})+(?:,\d{2})?)\s*(?:Kč|EUR|USD|CZK)\b|'
    # Částky s kontextem: "částka: 50 000", "hodnota: 5 240 000"
    r'(?:částka|cena|hodnota|kapitál|invest(?:ice)?|fond|úrok|splátka|dluh|hodnot[ayě])\s*:?\s*(\d{1,3}(?:\s+\d{3})+)\b|'
    # Velké částky (10+ mil) POUZE pokud mají 4+ skupiny: "125 000 000 000"
    # DŮLEŽITÉ: xxx xxx xxx je telefon, takže 9 číslic vynecháváme!
    r'\b(\d{3}(?:\s+\d{3}){3,})\b(?!\s*[-/])'  # 4+ skupiny = 12+ číslic
    r')',
    re.IGNORECASE
)

# =============== Třída Anonymizer ===============
class Anonymizer:
    def __init__(self, verbose=False):
        self.verbose = verbose
        self.counter = defaultdict(int)
        self.canonical_persons = OrderedDict()  # canonical -> label
        self.entity_map = defaultdict(lambda: defaultdict(set))  # typ -> original -> varianty

    def _get_or_create_label(self, typ: str, original: str, store_value: bool = True) -> str:
        """Vrátí existující nebo vytvoří nový štítek pro entitu.

        Args:
            typ: Typ entity (PASSWORD, EMAIL, atd.)
            original: Původní hodnota
            store_value: False pro citlivá data (hesla, API klíče) - uloží jen "***REDACTED***"
        """
        # Normalizace
        orig_norm = original.strip()

        # Pro citlivá data: kontrola duplicit podle skutečné hodnoty
        if not store_value:
            for existing_orig, variants in self.entity_map[typ].items():
                if orig_norm in variants:
                    existing_idx = list(self.entity_map[typ].keys()).index(existing_orig) + 1
                    return f"[[{typ}_{existing_idx}]]"
        else:
            # Pro běžná data: standardní kontrola
            for existing_orig, variants in self.entity_map[typ].items():
                if orig_norm in variants or orig_norm == existing_orig:
                    existing_idx = list(self.entity_map[typ].keys()).index(existing_orig) + 1
                    return f"[[{typ}_{existing_idx}]]"

        # Vytvoř nový
        idx = len(self.entity_map[typ]) + 1

        # Pro citlivá data: unikátní placeholder pro každý item
        # Pro běžná data: ukládej skutečnou hodnotu
        if store_value:
            map_key = orig_norm
        else:
            # Každý citlivý item má unikátní klíč, ale zobrazí se jako ***REDACTED***
            map_key = f"***REDACTED_{idx}***"

        self.entity_map[typ][map_key].add(orig_norm)
        return f"[[{typ}_{idx}]]"

    def _apply_known_people(self, text: str) -> str:
        """Aplikuje známé osoby (již detekované)."""
        for canonical, label in self.canonical_persons.items():
            # Vygeneruj všechny varianty
            parts = canonical.split()
            if len(parts) == 2:
                first, last = parts
                first_vars = variants_for_first(first)
                last_vars = variants_for_surname(last)

                # Všechny kombinace
                for fv in first_vars:
                    for lv in last_vars:
                        if fv and lv:
                            pattern = rf'\b{re.escape(fv)}\s+{re.escape(lv)}\b'
                            text = re.sub(pattern, label, text, flags=re.IGNORECASE)

        return text

    def _replace_remaining_people(self, text: str) -> str:
        """Detekuje a nahradí zbývající osoby."""
        # Hledá vzory: Jméno Příjmení
        # Nejprve tituly
        titles = r'(?:Ing\.|Mgr\.|Bc\.|MUDr\.|JUDr\.|PhDr\.|RNDr\.|Ph\.D\.|MBA|CSc\.|DrSc\.)'

        # Pattern pro jméno příjmení s volitelným titulem
        person_pattern = re.compile(
            rf'(?:{titles}\s+)?'
            r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)'
            r'\s+'
            r'([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+)',
            re.UNICODE
        )

        def replace_person(match):
            first_obs = match.group(1)
            last_obs = match.group(2)

            # Rozšířený seznam slov k ignorování (firmy, produkty, instituce)
            ignore_words = {
                # Běžná slova ve smlouvách
                'místo', 'datum', 'částku', 'bytem', 'sídlo', 'adresa',
                'číslo', 'kontakt', 'telefon', 'email', 'rodné', 'narozena',
                'vydán', 'uzavřena', 'podepsána', 'smlouva', 'dohoda',
                # Místa
                'staré', 'město', 'nové', 'město', 'malá', 'strana',
                'václavské', 'náměstí', 'hlavní', 'nádraží',
                # Organizace/instituce klíčová slova
                'česká', 'spořitelna', 'komerční', 'banka', 'raiffeisen',
                'credit', 'bank', 'financial', 'global', 'senior',
                'junior', 'lead', 'chief', 'head', 'director',
                # Finance/Investment
                'capital', 'equity', 'value', 'crescendo', 'investment',
                'fund', 'holdings', 'partners', 'assets', 'portfolio',
                # Pozice/role
                'jednatel', 'jednatelka', 'ředitel', 'ředitelka',
                'auditor', 'manager', 'consultant', 'specialist',
                'assistant', 'coordinator', 'analyst', 'pacient',
                'scrum', 'master', 'developer', 'architect', 'engineer',
                'officer', 'professional', 'certified', 'advanced',
                'management', 'legal', 'counsel', 'executive',
                # Pozdravy/oslovení
                'ahoj', 'dobrý', 'den', 'vážený', 'vážená',
                # Značky aut
                'škoda', 'octavia', 'fabia', 'superb', 'kodiaq',
                'volkswagen', 'toyota', 'ford', 'bmw', 'audi',
                # Technologie a software
                'google', 'amazon', 'microsoft', 'apple', 'facebook',
                'cloud', 'web', 'tech', 'solutions', 'data', 'digital',
                'software', 'enterprise', 'premium', 'standard',
                'analytics', 'computer', 'vision', 'protection',
                'security', 'authenticator', 'repository', 'access',
                'personal', 'hub', 'book', 'pro', 'series', 'launch',
                'team', 'development', 'react', 'splunk', 'innovate',
                'ventures', 'credo', 'mayo', 'clinic', 'met', 'london',
                'avenue', 'contractual', 'plánovaná', 'diagno',
                'cisco', 'processing', 'notářská',
                # Zdravotnictví + Léky/Produkty
                'nemocnice', 'poliklinika', 'polikliniek', 'nemocniec',
                'healthcare', 'symbicort', 'turbuhaler', 'spirometr',
                'jaeger', 'medical', 'health', 'pharma', 'pharmaceutical',
                # Další
                'care', 'plus', 'minus', 'service', 'services',
                'group', 'company', 'corp', 'ltd', 'gmbh', 'inc'
            }

            # Kontrola proti ignore listu
            if first_obs.lower() in ignore_words or last_obs.lower() in ignore_words:
                return match.group(0)

            # Detekce firem, produktů, institucí (neměly by být PERSON)
            combined = f"{first_obs} {last_obs}".lower()
            non_person_patterns = [
                # Tech/Software
                r'\b(tech|cloud|web|solutions?|data|digital|software|analytics)\b',
                r'\b(team|hub|enterprise|premium|standard|professional)\b',
                r'\b(google|amazon|microsoft|apple|facebook|splunk|cisco)\b',
                r'\b(repository|authenticator|vision|protection|security)\b',
                # Finance/Investment
                r'\b(capital|equity|value|investment|fund|holdings|assets)\b',
                r'\b(crescendo|ventures|partners|portfolio)\b',
                # Management/Business
                r'\b(management|processing|executive|legal|counsel)\b',
                r'\b(clinic|series|launch|innovate|healthcare)\b',
                # Products/Medical
                r'\b(symbicort|turbuhaler|spirometr|jaeger)\b',
                r'\b(pharma|pharmaceutical|medical)\b',
                # Company suffixes když jsou uprostřed
                r'\b(group|company|corp|ltd|gmbh|inc|services?)\b'
            ]
            for pattern in non_person_patterns:
                if re.search(pattern, combined):
                    return match.group(0)

            # Detekce názvů firem (končí na s.r.o., a.s., spol., Ltd. atd.)
            context_after = text[match.end():match.end()+20]
            if re.search(r'^\s*(s\.r\.o\.|a\.s\.|spol\.|k\.s\.|v\.o\.s\.|ltd\.?|inc\.?)', context_after, re.IGNORECASE):
                return match.group(0)

            # Nejdřív inference příjmení
            last_nom = infer_surname_nominative(last_obs)

            # Určení rodu podle příjmení
            is_female_surname = last_nom.lower().endswith(('ová', 'á'))

            # Inference křestního jména podle rodu příjmení
            first_lo = first_obs.lower()

            # Pokud příjmení je ženské, jméno musí být ženské
            if is_female_surname:
                # Han → Hana, Martin → Martina
                # Pravidlo: pokud jméno končí na souhlásku, přidej 'a'
                if not first_lo.endswith(('a', 'e', 'i', 'o', 'u', 'y')):
                    # Jméno končí na souhlásku → přidej 'a'
                    first_nom = (first_obs + 'a').capitalize()
                elif first_lo.endswith('a'):
                    # Jméno už končí na 'a' → je to pravděpodobně nominativ ženského jména, ponech
                    first_nom = first_obs.capitalize()
                else:
                    # Jiné koncovky → zkus inference
                    first_nom = infer_first_name_nominative(first_obs)
            else:
                # Příjmení je mužské, jméno musí být mužské
                # Jana → Jan, Petra → Petr (odstraň 'a' pokud je to genitiv)
                if first_lo.endswith('a') and len(first_lo) > 2:
                    # Výjimky - skutečná mužská jména končící na 'a'
                    male_names_with_a = {'kuba', 'míla', 'nikola', 'saša', 'jirka', 'honza'}
                    if first_lo in male_names_with_a:
                        first_nom = first_obs.capitalize()
                    else:
                        # Odstraň koncové 'a'
                        first_nom = first_obs[:-1].capitalize()
                elif first_lo.endswith(('u', 'e', 'em', 'ovi', 'ům')):
                    # Typické pádové koncovky → použij inference
                    first_nom = infer_first_name_nominative(first_obs)
                else:
                    # Jiné (pravděpodobně nominativ) → ponech jak je
                    first_nom = first_obs.capitalize()

            canonical = f"{first_nom} {last_nom}"

            # Získej nebo vytvoř label
            if canonical not in self.canonical_persons:
                idx = len(self.canonical_persons) + 1
                label = f"[[PERSON_{idx}]]"
                self.canonical_persons[canonical] = label
            else:
                label = self.canonical_persons[canonical]

            return label

        text = person_pattern.sub(replace_person, text)
        return text

    def anonymize_entities(self, text: str) -> str:
        """Anonymizuje všechny entity (adresy, kontakty, IČO, atd.)."""

        # DŮLEŽITÉ: Pořadí je klíčové! Od nejvíce specifických po nejméně specifické
        # KRITICKÉ: Credentials (hesla, API klíče) PRVNÍ s store_value=False!

        # 1. CREDENTIALS (username / password) - NEJPRVE!
        def replace_credentials(match):
            username = match.group(2)
            password = match.group(3)
            username_tag = self._get_or_create_label('USERNAME', username)
            password_tag = self._get_or_create_label('PASSWORD', password, store_value=False)
            return f"{match.group(1)}: {username_tag} / {password_tag}"
        text = CREDENTIALS_RE.sub(replace_credentials, text)

        # 2. HESLA (KRITICKÉ - hodnotu neukládat!)
        def replace_password(match):
            return self._get_or_create_label('PASSWORD', match.group(1), store_value=False)
        text = PASSWORD_RE.sub(replace_password, text)

        # 2. API KLÍČE, SECRETS (KRITICKÉ - hodnotu neukládat!)
        def replace_api_key(match):
            return self._get_or_create_label('API_KEY', match.group(1), store_value=False)
        text = API_KEY_RE.sub(replace_api_key, text)

        def replace_secret(match):
            return self._get_or_create_label('SECRET', match.group(1), store_value=False)
        text = SECRET_RE.sub(replace_secret, text)

        # 3. SSH KLÍČE (KRITICKÉ - hodnotu neukládat!)
        def replace_ssh_key(match):
            return self._get_or_create_label('SSH_KEY', match.group(1), store_value=False)
        text = SSH_KEY_RE.sub(replace_ssh_key, text)

        # 3.5. IBAN (PŘED kartami! IBAN má dlouhé číselné sekvence)
        def replace_iban(match):
            return self._get_or_create_label('IBAN', match.group(1), store_value=False)
        text = IBAN_RE.sub(replace_iban, text)

        # 4. PLATEBNÍ KARTY (KRITICKÉ - hodnotu neukládat!)
        def replace_card(match):
            # CARD_RE má 2 capture groups - získej první non-None
            card = match.group(1) if match.group(1) else match.group(2)
            if card:
                return self._get_or_create_label('CARD', card, store_value=False)
            return match.group(0)
        text = CARD_RE.sub(replace_card, text)

        # 5. USERNAMES, ACCOUNTS, HOSTNAMES
        def replace_username(match):
            return self._get_or_create_label('USERNAME', match.group(1))
        text = USERNAME_RE.sub(replace_username, text)

        def replace_account_id(match):
            return self._get_or_create_label('ACCOUNT_ID', match.group(1))
        text = ACCOUNT_ID_RE.sub(replace_account_id, text)

        def replace_hostname(match):
            return self._get_or_create_label('HOST', match.group(1))
        text = HOSTNAME_RE.sub(replace_hostname, text)

        # 6. IP ADRESY
        def replace_ip(match):
            return self._get_or_create_label('IP', match.group(1))
        text = IP_RE.sub(replace_ip, text)

        # 7. ČÍSLA POJIŠTĚNCE
        def replace_insurance_id(match):
            return self._get_or_create_label('INSURANCE_ID', match.group(1))
        text = INSURANCE_ID_RE.sub(replace_insurance_id, text)

        # 8. RFID/BADGE
        def replace_rfid(match):
            return self._get_or_create_label('RFID', match.group(1))
        text = RFID_RE.sub(replace_rfid, text)

        # 9. ADRESY (před jmény, aby "Novákova 45" nebylo osobou)
        def replace_address(match):
            return self._get_or_create_label('ADDRESS', match.group(0))
        text = ADDRESS_RE.sub(replace_address, text)

        # 10. EMAILY (před ostatními, protože obsahují speciální znaky)
        def replace_email(match):
            return self._get_or_create_label('EMAIL', match.group(1))
        text = EMAIL_RE.sub(replace_email, text)

        # 11. RODNÁ ČÍSLA (před čísly OP a telefony)
        def replace_birth_id(match):
            birth_id = match.group(1) if match.group(1) else match.group(2)
            if birth_id:
                return self._get_or_create_label('BIRTH_ID', birth_id)
            return match.group(0)
        text = BIRTH_ID_RE.sub(replace_birth_id, text)

        # 12. ČÍSLA OP (KRITICKÉ: MUSÍ být PŘED telefony!)
        def replace_id_card(match):
            id_card = match.group(1) if match.group(1) else match.group(2)
            if id_card:
                return self._get_or_create_label('ID_CARD', id_card)
            return match.group(0)
        text = ID_CARD_RE.sub(replace_id_card, text)

        # 13. TELEFONY (po ID_CARD, ale PŘED částkami!)
        def replace_phone(match):
            # PHONE_RE má capture group (1) pro samotné číslo (bez prefixu!)
            return self._get_or_create_label('PHONE', match.group(1))
        text = PHONE_RE.sub(replace_phone, text)

        # 14. BANKOVNÍ ÚČTY (před telefony!)
        def replace_bank(match):
            account = match.group(1) if match.group(1) else match.group(2)
            if account:
                return self._get_or_create_label('BANK', account)
            return match.group(0)
        text = BANK_RE.sub(replace_bank, text)

        # 15. DIČ (před IČO)
        def replace_dic(match):
            return self._get_or_create_label('DIC', match.group(1))
        text = DIC_RE.sub(replace_dic, text)

        # 16. IČO
        def replace_ico(match):
            full = match.group(0)
            # Ale ne pokud je to DIČ (CZ prefix)
            if 'CZ' in full.upper():
                return full
            # A ne pokud je to rodné číslo
            if '/' in full:
                return full
            # A ne pokud je to číslo OP
            if match.group(1):
                return self._get_or_create_label('ICO', match.group(1))
            return full
        text = ICO_RE.sub(replace_ico, text)

        # 17. SPZ
        def replace_license_plate(match):
            return self._get_or_create_label('SPZ', match.group(0))
        text = LICENSE_PLATE_RE.sub(replace_license_plate, text)

        # 18. ČÁSTKY (AŽ NAKONEC! Po telefonech a všech číselných identifikátorech)
        def replace_amount(match):
            # AMOUNT_RE má 3 capture groups - získej první non-None
            amount = match.group(1) or match.group(2) or match.group(3)
            if amount:
                return self._get_or_create_label('AMOUNT', amount)
            return match.group(0)
        text = AMOUNT_RE.sub(replace_amount, text)

        # 19. MASKOVÁNÍ CVV/EXPIRACE u karet
        # Nahradí "CVV: 123" → "CVV: ***" a "exp: 12/26" → "exp: **/**"
        text = re.sub(r'(CVV|CVC)\s*:\s*\d{3,4}', r'\1: ***', text, flags=re.IGNORECASE)
        text = re.sub(r'exp(?:\.|\s+|iration)?\s*:?\s*\d{2}/\d{2,4}', r'exp: **/**', text, flags=re.IGNORECASE)

        # 19.5. CLEANUP biometrických identifikátorů - odstraň [[PHONE_*]] z bio prefixů
        # "IRIS_SCAN_PD_[[PHONE_10]]" → "IRIS_SCAN_PD_10"
        text = re.sub(
            r'(IRIS_SCAN|VOICE_RK|HASH_BIO|FINGERPRINT|FACIAL|RETINA|PALM|DNA)_([A-Z0-9_]+)\[\[PHONE_\d+\]\]',
            r'\1_\2',
            text
        )

        # 20. END-SCAN - finální kontrola citlivých dat (chytá zbytky nalepené na ]])
        text = self._end_scan(text)

        return text

    def _luhn_check(self, card_number: str) -> bool:
        """Validace Luhn algoritmem pro platební karty."""
        digits = [int(d) for d in card_number if d.isdigit()]
        if len(digits) < 13 or len(digits) > 19:
            return False

        checksum = 0
        for i, digit in enumerate(reversed(digits)):
            if i % 2 == 1:
                digit *= 2
                if digit > 9:
                    digit -= 9
            checksum += digit

        return checksum % 10 == 0

    def _end_scan(self, text: str) -> str:
        """Finální sken po všech náhradách - chytá případné zbytky citlivých dat."""

        # LUHN END-SCAN - zachytí všechny Luhn-validní karty (13-19 číslic)
        luhn_pattern = re.compile(r'\b(\d[\s\-]?){12,18}\d\b')
        def final_luhn_card(match):
            candidate = match.group(0)
            # Přeskoč pokud už je v [[CARD_*]]
            if '[[CARD_' in text[max(0, match.start()-10):match.end()+10]:
                return match.group(0)
            # Validuj Luhn
            if self._luhn_check(candidate):
                return self._get_or_create_label('CARD', candidate, store_value=False)
            return match.group(0)
        text = luhn_pattern.sub(final_luhn_card, text)

        # IBAN (pokud unikl)
        def final_iban(match):
            if not '[[IBAN_' in text[max(0, match.start()-10):match.start()+30]:
                return self._get_or_create_label('IBAN', match.group(1), store_value=False)
            return match.group(0)
        text = IBAN_RE.sub(final_iban, text)

        # Platební karty (pokud unikly nebo mají CVV/exp. datum)
        def final_card(match):
            card = match.group(1) if match.group(1) else match.group(2)
            if card and not '[[CARD_' in text[max(0, match.start()-10):match.start()+len(card)+10]:
                return self._get_or_create_label('CARD', card, store_value=False)
            return match.group(0)
        text = CARD_RE.sub(final_card, text)

        # Hesla (pokud unikla nebo jsou nalepená na jiných entitách)
        def final_password(match):
            return self._get_or_create_label('PASSWORD', match.group(1), store_value=False)
        text = PASSWORD_RE.sub(final_password, text)

        # IP adresy (pokud unikly)
        def final_ip(match):
            # Přeskoč pokud už je součástí URL/hostname
            if not re.search(r'[a-z]', text[max(0, match.start()-10):match.start()], re.IGNORECASE):
                return self._get_or_create_label('IP', match.group(1))
            return match.group(0)
        text = IP_RE.sub(final_ip, text)

        # Usernames (pokud unikly)
        def final_username(match):
            return self._get_or_create_label('USERNAME', match.group(1))
        text = USERNAME_RE.sub(final_username, text)

        # API klíče, secrets (pokud unikly)
        def final_api(match):
            return self._get_or_create_label('API_KEY', match.group(1), store_value=False)
        text = API_KEY_RE.sub(final_api, text)

        def final_secret(match):
            return self._get_or_create_label('SECRET', match.group(1), store_value=False)
        text = SECRET_RE.sub(final_secret, text)

        # Pojištěnce (pokud unikla)
        def final_insurance(match):
            return self._get_or_create_label('INSURANCE_ID', match.group(1))
        text = INSURANCE_ID_RE.sub(final_insurance, text)

        # RFID (pokud uniklo)
        def final_rfid(match):
            return self._get_or_create_label('RFID', match.group(1))
        text = RFID_RE.sub(final_rfid, text)

        return text

    def anonymize_docx(self, input_path: str, output_path: str, json_map: str, txt_map: str):
        """Hlavní metoda pro anonymizaci DOCX dokumentu."""
        print(f"\n🔍 Zpracovávám: {Path(input_path).name}")

        # Načti dokument
        doc = Document(input_path)

        # Zpracuj všechny odstavce
        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            original = para.text

            # POŘADÍ JE KLÍČOVÉ!
            # 1. Nejprve anonymizuj entity (adresy, IČO, telefony...)
            text = self.anonymize_entities(original)

            # 2. Potom aplikuj známé osoby
            text = self._apply_known_people(text)

            # 3. Nakonec detekuj nové osoby
            text = self._replace_remaining_people(text)

            if text != original:
                para.text = text

        # Zpracuj tabulky
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.text.strip():
                            continue

                        original = para.text
                        text = self.anonymize_entities(original)
                        text = self._apply_known_people(text)
                        text = self._replace_remaining_people(text)

                        if text != original:
                            para.text = text

        # Ulož dokument
        doc.save(output_path)

        # Vytvoř mapy
        self._create_maps(json_map, txt_map, input_path)

        print(f"✅ Hotovo! Nalezeno {len(self.canonical_persons)} osob")

    def _create_maps(self, json_path: str, txt_path: str, source_file: str):
        """Vytvoří JSON a TXT mapy náhrad."""

        # JSON mapa
        json_data = {
            "version": "1.0",
            "generated_at": datetime.now().isoformat(),
            "source_file": Path(source_file).name,
            "entities": []
        }

        # Osoby
        for canonical, label in self.canonical_persons.items():
            json_data["entities"].append({
                "type": "PERSON",
                "label": label,
                "original": canonical,
                "occurrences": 1
            })

        # Ostatní entity
        for typ, entities in self.entity_map.items():
            for idx, (original, variants) in enumerate(entities.items(), 1):
                json_data["entities"].append({
                    "type": typ,
                    "label": f"[[{typ}_{idx}]]",
                    "original": original,
                    "occurrences": len(variants)
                })

        # Ulož JSON
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)

        # TXT mapa
        with open(txt_path, 'w', encoding='utf-8') as f:
            # Osoby
            if self.canonical_persons:
                f.write("OSOBY\n")
                for canonical, label in self.canonical_persons.items():
                    f.write(f"{label}: {canonical}\n")
                f.write("\n")

            # Ostatní entity
            for typ, entities in sorted(self.entity_map.items()):
                if entities:
                    f.write(f"{typ}\n")
                    for idx, (original, variants) in enumerate(entities.items(), 1):
                        label = f"[[{typ}_{idx}]]"
                        # Pro citlivá data zobraz jen ***REDACTED*** bez čísla
                        display_value = "***REDACTED***" if original.startswith("***REDACTED_") else original
                        f.write(f"{label}: {display_value}\n")
                    f.write("\n")

# =============== Batch processing ===============
def batch_anonymize(folder_path, names_json="cz_names.v1.json"):
    """Zpracuje všechny DOCX soubory v adresáři."""
    folder = Path(folder_path)
    docx_files = sorted([f for f in folder.glob("*.docx") if not f.name.startswith('~') and '_anon' not in f.name])

    if not docx_files:
        print(f"Nebyly nalezeny žádné .docx soubory v adresáři {folder_path}")
        return

    print(f"\n📁 Zpracovávám {len(docx_files)} souborů v adresáři {folder_path}\n")

    global CZECH_FIRST_NAMES
    CZECH_FIRST_NAMES = load_names_library(names_json)

    for path in docx_files:
        print(f"\n{'='*60}")
        base = path.stem
        out_docx = path.parent / f"{base}_anon.docx"
        out_json = path.parent / f"{base}_map.json"
        out_txt = path.parent / f"{base}_map.txt"

        try:
            a = Anonymizer(verbose=False)
            a.anonymize_docx(str(path), str(out_docx), str(out_json), str(out_txt))
            print(f"✅ Výstupy: {out_docx.name}, {out_json.name}, {out_txt.name}")
        except Exception as e:
            print(f"❌ CHYBA při zpracování {path.name}: {e}")
            import traceback
            traceback.print_exc()

# =============== Main ===============
def main():
    import argparse
    ap = argparse.ArgumentParser(description="Anonymizace českých DOCX s JSON knihovnou jmen")
    ap.add_argument("docx_path", nargs='?', help="Cesta k .docx souboru nebo adresáři")
    ap.add_argument("--names-json", default="cz_names.v1.json", help="Cesta k JSON knihovně jmen")
    ap.add_argument("--batch", action="store_true", help="Zpracovat všechny .docx v adresáři")
    args = ap.parse_args()

    try:
        global CZECH_FIRST_NAMES
        CZECH_FIRST_NAMES = load_names_library(args.names_json)

        if args.batch and args.docx_path:
            batch_anonymize(args.docx_path, args.names_json)
            return 0

        if args.batch and not args.docx_path:
            # Batch mode v aktuálním adresáři
            batch_anonymize(".", args.names_json)
            return 0

        # Single file mode
        if not args.docx_path:
            print("❌ Chybí cesta k souboru. Použij: python script.py <soubor.docx>")
            print("   Nebo: python script.py --batch <adresář>")
            return 2

        path = Path(args.docx_path)
        if not path.exists():
            print(f"❌ Soubor nenalezen: {path}")
            return 2

        base = path.stem
        out_docx = path.parent / f"{base}_anon.docx"
        out_json = path.parent / f"{base}_map.json"
        out_txt = path.parent / f"{base}_map.txt"

        # Kontrola zamčených souborů
        files_locked = False
        for out_file in [out_docx, out_json, out_txt]:
            if out_file.exists():
                try:
                    with open(out_file, 'a'):
                        pass
                except PermissionError:
                    files_locked = True
                    break

        if files_locked:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            out_docx = path.parent / f"{base}_anon_{timestamp}.docx"
            out_json = path.parent / f"{base}_map_{timestamp}.json"
            out_txt = path.parent / f"{base}_map_{timestamp}.txt"
            print(f"\n⚠️  Výstupní soubory jsou otevřené v jiné aplikaci!")
            print(f"   Vytvářím nové soubory s časovým razítkem: {timestamp}\n")

        a = Anonymizer(verbose=False)
        a.anonymize_docx(str(path), str(out_docx), str(out_json), str(out_txt))

        print(f"\n✅ Výstupy:")
        print(f" - {out_docx}")
        print(f" - {out_json}")
        print(f" - {out_txt}")
        print(f"\n📊 Statistiky:")
        print(f" - Nalezeno osob: {len(a.canonical_persons)}")
        print(f" - Celkem entit: {sum(len(e) for e in a.entity_map.values())}")

        return 0

    except Exception as e:
        print(f"\n❌ CHYBA: {e}")
        import traceback
        traceback.print_exc()
        return 1

if __name__ == "__main__":
    sys.exit(main())
