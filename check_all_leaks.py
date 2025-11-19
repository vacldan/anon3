#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Comprehensive leak check for contracts 13-16"""

import re
from docx import Document
from pathlib import Path

def check_document(doc_path):
    """Check document for plain text PII leaks"""
    doc = Document(doc_path)

    # Get all text
    all_text = []
    for para in doc.paragraphs:
        all_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_text.append(para.text)

    full_text = '\n'.join(all_text)

    print(f"\n{'='*80}")
    print(f"CHECKING: {doc_path.name}")
    print(f"{'='*80}")

    leaks = []

    # 1. Plain IBANs (CZ + 2 digits + 20 digits)
    iban_pattern = re.compile(r'\b(CZ\d{22})\b')
    plain_ibans = iban_pattern.findall(full_text)
    if plain_ibans:
        leaks.append(f"❌ HARD FAIL: {len(plain_ibans)} plain IBAN(s): {plain_ibans[:3]}")
    else:
        print("✅ No plain IBANs")

    # 2. Plain bank accounts (XXXXXXXX/XXXX)
    bank_pattern = re.compile(r'\b(\d{6,16}/\d{4})\b')
    # Filter out those already tagged
    potential_banks = bank_pattern.findall(full_text)
    plain_banks = [b for b in potential_banks if f"[[BANK_" not in full_text[max(0, full_text.find(b)-10):full_text.find(b)+len(b)+20]]

    if plain_banks:
        leaks.append(f"❌ HARD FAIL: {len(plain_banks)} plain BANK account(s): {plain_banks[:3]}")
    else:
        print("✅ No plain bank accounts")

    # 3. Plain cards (16-19 digits)
    card_pattern = re.compile(r'\b(\d{4}[\s\-]?\d{4}[\s\-]?\d{4}[\s\-]?\d{4}(?:[\s\-]?\d{2,3})?)\b')
    potential_cards = card_pattern.findall(full_text)
    plain_cards = [c for c in potential_cards if len(c.replace(' ', '').replace('-', '')) >= 13
                   and f"[[CARD_" not in full_text[max(0, full_text.find(c)-10):full_text.find(c)+len(c)+20]]

    if plain_cards:
        leaks.append(f"❌ HARD FAIL: {len(plain_cards)} plain CARD(s): {plain_cards[:3]}")
    else:
        print("✅ No plain cards")

    # 4. Plain emails
    email_pattern = re.compile(r'\b([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b')
    potential_emails = email_pattern.findall(full_text)
    plain_emails = [e for e in potential_emails if f"[[EMAIL_" not in full_text[max(0, full_text.find(e)-10):full_text.find(e)+len(e)+20]]

    if plain_emails:
        leaks.append(f"❌ HARD FAIL: {len(plain_emails)} plain EMAIL(s): {plain_emails[:3]}")
    else:
        print("✅ No plain emails")

    # 5. Plain birth IDs (XXXXXX/XXXX)
    birth_id_pattern = re.compile(r'\b(\d{6}/\d{3,4})\b')
    potential_births = birth_id_pattern.findall(full_text)
    plain_births = [b for b in potential_births if f"[[BIRTH_ID_" not in full_text[max(0, full_text.find(b)-10):full_text.find(b)+len(b)+20]]

    if plain_births:
        leaks.append(f"❌ HARD FAIL: {len(plain_births)} plain BIRTH_ID(s): {plain_births[:3]}")
    else:
        print("✅ No plain birth IDs")

    # 6. Plain phones (+420 XXX XXX XXX or XXX XXX XXX)
    phone_pattern = re.compile(r'\b(\+?420\s?\d{3}\s?\d{3}\s?\d{3}|\d{3}\s?\d{3}\s?\d{3})\b')
    potential_phones = phone_pattern.findall(full_text)
    plain_phones = [p for p in potential_phones if f"[[PHONE_" not in full_text[max(0, full_text.find(p)-10):full_text.find(p)+len(p)+20]]

    if plain_phones:
        leaks.append(f"⚠️  MAJOR: {len(plain_phones)} plain PHONE(s): {plain_phones[:3]}")
    else:
        print("✅ No plain phones")

    # 7. Plain SPZ
    spz_pattern = re.compile(r'\b(\d[A-Z]\d\s+\d{4}|\d[A-Z]{2}\s+\d{4})\b', re.IGNORECASE)
    potential_spz = spz_pattern.findall(full_text)
    plain_spz = [s for s in potential_spz if f"[[LICENSE_PLATE_" not in full_text[max(0, full_text.find(s)-10):full_text.find(s)+len(s)+20]]

    if plain_spz:
        leaks.append(f"❌ HARD FAIL: {len(plain_spz)} plain SPZ: {plain_spz[:3]}")
    else:
        print("✅ No plain SPZ")

    # Print summary
    if leaks:
        print(f"\n🚨 TOTAL LEAKS: {len(leaks)}")
        for leak in leaks:
            print(f"  {leak}")
        return False
    else:
        print(f"\n✅ NO LEAKS FOUND - Document appears clean")
        return True

# Check all contracts
contracts = ['smlouva13_anon.docx', 'smlouva14_anon.docx', 'smlouva15_anon.docx', 'smlouva16_anon.docx']
all_clean = True

for contract in contracts:
    path = Path(contract)
    if path.exists():
        clean = check_document(path)
        all_clean = all_clean and clean
    else:
        print(f"\n⚠️  {contract} not found")

print(f"\n{'='*80}")
if all_clean:
    print("✅ ALL CONTRACTS CLEAN - NO PII LEAKS DETECTED")
else:
    print("❌ PII LEAKS DETECTED - REQUIRES IMMEDIATE FIX")
print(f"{'='*80}")
