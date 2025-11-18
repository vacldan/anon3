#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Test zpracování prvního odstavce smlouvy 14"""

from docx import Document
import sys

# Načti dokument
doc = Document('smlouva14.docx')

print("Načítám první 5 odstavců...")
for i, para in enumerate(doc.paragraphs[:5]):
    print(f"\n[{i}] Text: {para.text[:100]}...")
    print(f"    Délka: {len(para.text)}")

print("\n\nNačítám první neprázdný odstavec...")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"\n[{i}] První neprázdný:")
        print(f"Text: {para.text}")
        print(f"Délka: {len(para.text)}")
        break
