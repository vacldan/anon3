# -*- coding: utf-8 -*-
"""
Czech DOCX Anonymizer - Simplified Version
Správně sjednocuje české pády jmen pod jeden tag
"""

import sys, re, json, unicodedata
from pathlib import Path
from collections import defaultdict
from docx import Document

# Globální proměnná pro knihovnu jmen
CZECH_FIRST_NAMES = set()

def load_names_library(json_path="cz_names.v1.json"):
    """Načte českou knihovnu jmen z JSON"""
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    names = set()
    for gender in ['M', 'F']:
        for name in data.get('firstnames', {}).get(gender, []):
            names.add(name.lower())
    return names

def remove_diacritics(text):
    """Odstraní diakritiku z textu"""
    return ''.join(
        c for c in unicodedata.normalize('NFD', text)
        if unicodedata.category(c) != 'Mn'
    )

def infer_first_name_nominative(name):
    """Převede křestní jméno do nominativu (1. pádu)"""
    if not name or len(name) < 2:
        return name

    obs = name.strip()
    lo = obs.lower()

    # SPECIÁLNÍ PŘÍPAD PRVNÍ: Roberta může být genitiv od Robert
    # Musí být PŘED kontrolou knihovny, protože Roberta je v knihovně jako ženské jméno!
    if lo == 'roberta':
        # Preferujeme Robert (mužské jméno), protože Roberta je častěji genitiv než samostatné jméno
        return 'Robert'

    # Pokud je jméno přímo v knihovně, vrátíme ho
    if lo in CZECH_FIRST_NAMES:
        return obs.capitalize()

    # Ochrana běžných jmen která končí na typické deklinační koncovky
    protected_female_names = {
        'martina', 'kristina', 'pavlína', 'karolína', 'jana', 'hana',
        'eva', 'anna', 'petra', 'daniela', 'michaela', 'andrea',
        'lenka', 'tereza', 'barbora', 'veronika', 'nikola', 'viktorie',
        'irena', 'sylva', 'šárka', 'alžběta', 'adéla'
        # POZOR: 'roberta' NENÍ v protected - může být genitiv od Robert!
    }
    protected_male_names = {
        'david', 'martin', 'jakub', 'tomáš', 'jan', 'petr', 'pavel',
        'ivan', 'robert', 'ondřej', 'hynek', 'luděk', 'radovan', 'marek',
        'tobiáš', 'mark', 'radek'
    }

    if lo in protected_female_names:
        return obs

    # Genitiv/Akuzativ mužských jmen: -a → odstranit (Ivana → Ivan, Roberta → Robert)
    # ALE: ne pokud je to ženské jméno (Martina, Andrea, Petra...)
    if lo.endswith('a') and len(obs) > 3:
        base = obs[:-1]
        base_lo = base.lower()

        # Tato část už není potřeba - Roberta se řeší nahoře
        # (ponechána jako komentář pro dokumentaci)

        # Zkontroluj, zda základ je mužské jméno
        if base_lo in protected_male_names:
            return base.capitalize()
        if base_lo in CZECH_FIRST_NAMES:
            return base.capitalize()
        if remove_diacritics(base_lo) in {remove_diacritics(n) for n in CZECH_FIRST_NAMES}:
            return base.capitalize()

    # Genitiv ženských jmen: -y → -a (Pavlíny → Pavlína)
    if lo.endswith('y') and len(obs) > 3:
        base = obs[:-1] + 'a'
        if base.lower() in CZECH_FIRST_NAMES or base.lower() in protected_female_names:
            return base

    # Instrumentál ženských jmen: -ou → -a (Pavlínou → Pavlína)
    if lo.endswith('ou') and len(obs) > 4:
        base = obs[:-2] + 'a'
        if base.lower() in CZECH_FIRST_NAMES or base.lower() in protected_female_names:
            return base

    # Dativ/Lokál: -ovi, -emu → odstranit (Ivanovi → Ivan)
    if lo.endswith('ovi') and len(obs) > 5:
        base = obs[:-3]
        if base.lower() in CZECH_FIRST_NAMES or remove_diacritics(base.lower()) in {remove_diacritics(n) for n in CZECH_FIRST_NAMES}:
            return base.capitalize()

    # Instrumentál mužských jmen: -em → odstranit (Ivanem → Ivan, Tomášem → Tomáš)
    if lo.endswith('em') and len(obs) > 4:
        base = obs[:-2]
        if base.lower() in CZECH_FIRST_NAMES or base.lower() in protected_male_names or remove_diacritics(base.lower()) in {remove_diacritics(n) for n in CZECH_FIRST_NAMES}:
            return base.capitalize()

    return obs

def infer_surname_nominative(surname):
    """Převede příjmení do nominativu (1. pádu)"""
    if not surname or len(surname) < 3:
        return surname

    obs = surname.strip()
    lo = obs.lower()

    # Ochrana příjmení, která končí na deklinační vzory ale jsou už v nominativu
    protected_surnames = {
        'procházka', 'němec', 'sedláček', 'kučera', 'fiala', 'dušek',
        'kozel', 'šembera', 'havel', 'pavel', 'klíma', 'svoboda',
        'valach', 'štrunc', 'jůza'
    }
    if lo in protected_surnames:
        return obs

    # Genitiv/Dativ/Lokál žen: -é → -á (Pokorné → Pokorná, Houfové → Houfová)
    if lo.endswith('é') and len(obs) > 3:
        return obs[:-1] + 'á'

    # Instrumentál: -ou → může být -á (žena) nebo -ý (muž)
    if lo.endswith('ou') and len(obs) > 4:
        # Heuristika: pokud základ končí na souhlásku + typický vzor
        base = obs[:-2]
        # Pro příjmení jako "Vránou" → může být "Vráný" (muž) nebo "Vráná" (žena)
        # Zkusíme nejprve mužský tvar
        if base.lower().endswith(('vrán', 'novot', 'malý', 'černý')):
            return base + 'ý'
        # Jinak ženský tvar
        return obs[:-2] + 'á'

    # Genitiv mužů: -y → -a (Klímy → Klíma, ale ne Nováky → Novák)
    if lo.endswith('y') and len(obs) > 3:
        # Kontrola: příjmení na -a v nominativu (Klíma, Procházka)
        base_a = obs[:-1] + 'a'
        if base_a.lower() in protected_surnames:
            return base_a
        # Obecná heuristika: -y → -a pro příjmení jako Klíma
        if obs[:-1].lower().endswith(('klím', 'dvořák', 'svobod')):
            return base_a
        # Jinak jen odstraň -y (Nováky → Novák)
        return obs[:-1]

    # Genitiv mužů: -a → odstranit (Nováka → Novák)
    if lo.endswith('a') and len(obs) > 3 and lo not in protected_surnames:
        base = obs[:-1]
        return base

    # Dativ/Lokál: -ovi → odstranit (Novákovi → Novák)
    if lo.endswith('ovi') and len(obs) > 5:
        return obs[:-3]

    # Instrumentál mužů: -em → odstranit
    if lo.endswith('em') and len(obs) > 4:
        # Speciální: -alem, -elem, -olem (Doležalem → Doležal)
        if lo.endswith(('alem', 'elem', 'olem', 'ilem')):
            return obs[:-2]  # odstraň -em, nechej -al/-el/-ol/-il
        # Kontrola: -kem → -ek (Práškem → Prášek, Štefánkem → Štefánek)
        # Musíme přidat 'e' zpět: -kem → -ek (odstraň -em, nechej -k, přidej e před k)
        if lo.endswith('kem') and len(obs) > 5:
            return obs[:-3] + 'ek'  # např. Práškem → Prášek
        # Běžný instrumentál (Novákem → Novák)
        elif not lo.endswith(('lem', 'rem', 'sem', 'šem', 'cem', 'dem', 'nem', 'bem', 'gem', 'chem')):
            return obs[:-2]

    # Genitiv -ka → -ek (Hájka → Hájek, Štefánka → Štefánek)
    if lo.endswith('ka') and len(obs) > 4:
        # Kontrola: ne pro ženská příjmení jako "Procházková"
        if not lo.endswith(('ková', 'ská', 'ná')):
            base = obs[:-1] + 'ek'
            return base

    return obs

def normalize_person_name(first, last):
    """Normalizuje celé jméno osoby do nominativu"""
    first_nom = infer_first_name_nominative(first) if first else ""
    last_nom = infer_surname_nominative(last) if last else ""
    return first_nom, last_nom

def is_likely_person_name(first, last):
    """Zjistí, zda se pravděpodobně jedná o skutečné jméno osoby"""
    if not first or not last:
        return False

    first_lo = first.lower()
    last_lo = last.lower()
    first_no_dia = remove_diacritics(first_lo)

    # Blacklist - známé ne-osoby
    non_persons = {
        ('ford', 'transit'), ('škoda', 'karoq'), ('škoda', 'octavia'),
        ('nemocnice', 'sv'), ('hospitalizace', 'pacient'),
        ('kardiologie', 'pacient'), ('neurologie', 'neurologická'),
        ('pacient', 'ing'), ('pacienta', 'pacient')
    }
    if (first_lo, last_lo) in non_persons:
        return False

    # Kontrola v knihovně jmen
    if first_lo in CZECH_FIRST_NAMES:
        return True
    if first_no_dia in {remove_diacritics(n) for n in CZECH_FIRST_NAMES}:
        return True

    # Hardcoded seznam běžných jmen
    common_names = {
        'ivan', 'jan', 'petr', 'pavel', 'tomáš', 'jiří', 'martin', 'david',
        'jakub', 'lukáš', 'michal', 'ondřej', 'václav', 'josef', 'karel',
        'františek', 'jaroslav', 'miroslav', 'stanislav', 'zdeněk',
        'eva', 'marie', 'jana', 'anna', 'hana', 'věra', 'alena', 'lenka',
        'petra', 'martina', 'marcela', 'jitka', 'lucie', 'kateřina',
        'barbora', 'veronika', 'tereza', 'nikola', 'michaela', 'andrea',
        'radovan', 'robert', 'hynek', 'luděk', 'radek', 'marek',
        'pavlína', 'irena', 'sylva', 'šárka', 'alžběta', 'adéla', 'viktorie',
        'tobiáš', 'mark', 'ivanem', 'tomášem', 'markem', 'pavlínou',
        'pavlíny', 'roberta', 'terezy', 'lenkou', 'šárkou', 'alžbětou'
    }

    if first_lo in common_names or first_no_dia in {remove_diacritics(n) for n in common_names}:
        return True

    return False

class CzechAnonymizer:
    def __init__(self):
        self.persons = {}  # canonical_name -> {'tag': 'PERSON_X', 'variants': set()}
        self.person_counter = 0
        self.entity_map = defaultdict(dict)  # typ -> {canonical -> [variants]}
        self.all_replacements = []  # [(original_text, tag, start, end), ...]

    def _get_or_create_person_tag(self, first_nom, last_nom, original_form):
        """Získá nebo vytvoří tag pro osobu"""
        canonical = f"{first_nom} {last_nom}"

        if canonical not in self.persons:
            self.person_counter += 1
            tag = f"[[PERSON_{self.person_counter}]]"
            self.persons[canonical] = {
                'tag': tag,
                'variants': set()
            }

        # Přidej variantu (pokud je jiná než kanonická)
        if original_form.lower() != canonical.lower():
            self.persons[canonical]['variants'].add(original_form)

        return self.persons[canonical]['tag']

    def _find_and_tag_persons(self, text):
        """Najde všechny osoby v textu a přiřadí jim tagy"""
        # Regex pro detekci jmen: Dvě slova začínající velkým písmenem
        pattern = r'\b([A-ZČĎŇŘŠŤŽÁ-Ž][a-zčďěňřšťůúýžá-ž]+)\s+([A-ZČĎŇŘŠŤŽÁ-Ž][a-zčďěňřšťůúýžá-ž]+)\b'

        matches = list(re.finditer(pattern, text))

        for match in matches:
            first = match.group(1)
            last = match.group(2)

            # Zkontroluj, zda je to pravděpodobně osoba
            if not is_likely_person_name(first, last):
                continue

            # Normalizuj do nominativu
            first_nom, last_nom = normalize_person_name(first, last)

            # Získej tag
            original_form = f"{first} {last}"
            tag = self._get_or_create_person_tag(first_nom, last_nom, original_form)

            # Zapamatuj si nahrazení
            self.all_replacements.append({
                'original': original_form,
                'tag': tag,
                'start': match.start(),
                'end': match.end(),
                'type': 'PERSON'
            })

    def _replace_in_text(self, text):
        """Nahradí všechny nalezené entity v textu"""
        # Seřaď nahrazení od konce k začátku (aby indexy zůstaly platné)
        sorted_replacements = sorted(self.all_replacements, key=lambda x: x['start'], reverse=True)

        for repl in sorted_replacements:
            text = text[:repl['start']] + repl['tag'] + text[repl['end']:]

        return text

    def _create_map_txt(self, output_path):
        """Vytvoří textovou mapu s variantami"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("OSOBY\n-----\n")

            # Seřaď osoby podle čísla tagu
            sorted_persons = sorted(
                self.persons.items(),
                key=lambda x: int(x[1]['tag'].split('_')[1].rstrip(']'))
            )

            for canonical, info in sorted_persons:
                f.write(f"{info['tag']}: {canonical}\n")
                # Vypiš varianty, pokud existují
                for variant in sorted(info['variants']):
                    f.write(f"  - {variant}\n")

            f.write("\n")

    def anonymize_docx(self, input_path, output_docx, output_map_txt):
        """Hlavní funkce pro anonymizaci DOCX"""
        print(f"📖 Načítám: {input_path}")
        doc = Document(input_path)

        # Extrahuj text z dokumentu
        full_text = '\n'.join([para.text for para in doc.paragraphs])

        print(f"🔍 Hledám osoby...")
        self._find_and_tag_persons(full_text)

        print(f"   Nalezeno osob: {len(self.persons)}")

        # Nahraď text v dokumentu
        print(f"🔄 Nahrazuji v dokumentu...")
        for para in doc.paragraphs:
            if para.text.strip():
                # Pro každý odstavec najdeme nahrazení
                para_start = full_text.find(para.text)
                if para_start == -1:
                    continue

                para_end = para_start + len(para.text)

                # Najdi všechna nahrazení v tomto odstavci
                para_replacements = [
                    r for r in self.all_replacements
                    if para_start <= r['start'] < para_end
                ]

                # Seřaď od konce
                para_replacements.sort(key=lambda x: x['start'], reverse=True)

                # Aplikuj nahrazení
                new_text = para.text
                for repl in para_replacements:
                    local_start = repl['start'] - para_start
                    local_end = repl['end'] - para_start
                    new_text = new_text[:local_start] + repl['tag'] + new_text[local_end:]

                para.text = new_text

        # Ulož výstupy
        print(f"💾 Ukládám výstupy...")
        doc.save(output_docx)
        self._create_map_txt(output_map_txt)

        print(f"✅ Hotovo!")
        return len(self.persons)

def main():
    if len(sys.argv) < 2:
        print("Použití: python anonymizer.py <cesta_k_docx>")
        return 1

    input_file = sys.argv[1]
    input_path = Path(input_file)

    if not input_path.exists():
        print(f"❌ Soubor nenalezen: {input_file}")
        return 1

    # Načti knihovnu jmen
    global CZECH_FIRST_NAMES
    print("📚 Načítám knihovnu jmen...")
    CZECH_FIRST_NAMES = load_names_library("cz_names.v1.json")
    print(f"   Načteno {len(CZECH_FIRST_NAMES)} jmen")

    # Vytvoř výstupní cesty
    base = input_path.stem
    output_docx = input_path.parent / f"{base}_anon.docx"
    output_map = input_path.parent / f"{base}_map.txt"

    # Anonymizuj
    anon = CzechAnonymizer()
    person_count = anon.anonymize_docx(str(input_path), str(output_docx), str(output_map))

    print(f"\n📊 Statistiky:")
    print(f"   Nalezeno osob: {person_count}")
    print(f"\n📁 Výstupy:")
    print(f"   {output_docx}")
    print(f"   {output_map}")

    return 0

if __name__ == "__main__":
    sys.exit(main())
